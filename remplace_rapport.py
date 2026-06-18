import argparse
import re
import io
import tempfile
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from excel_processor import strip_html_tags

# Configuration de compression des images pour réduire la taille du fichier Word
IMAGE_MAX_WIDTH = 1600  # Largeur max en pixels (réduit de 1920)
IMAGE_QUALITY = 80  # Qualité JPEG (0-100, 80 = bon compromis qualité/taille)


# Mise a l'echelle Word des bulles et composites de miniatures.
# Probleme: width=Cm(12) fixe etire les petites bulles (effet zoom incoherent
# entre bulles courtes et longues). On scale en cm proportionnellement a la
# largeur reelle du PNG pour conserver un DPI uniforme.
BUBBLE_PX_PER_CM = 80.0   # 1120 px -> 14 cm, 640 px -> 8 cm, 320 px -> 4 cm
BUBBLE_MIN_CM = 5.0
BUBBLE_MAX_CM = 14.0

# Pour les bulles CLEAN (nom contenant "_clean") : largeur Word FIXE.
# PNG genere a haute densite (~840 px) pour avoir ~305 DPI a 7 cm -> tres net.
# Largeur reduite a 7 cm pour que la bulle soit moins large (police inchangee).
CLEAN_BUBBLE_CM = 7.0

def _compute_bubble_cm(image_path):
    """Renvoie la largeur cm a utiliser pour add_picture.
    - Bulles CLEAN (_clean.png) : largeur FIXE = CLEAN_BUBBLE_CM (haute densite).
    - Autres : proportionnelle a la largeur reelle (px) du PNG, capee."""
    try:
        with Image.open(image_path) as im:
            px_w = im.width
    except Exception:
        return 12.0
    # Largeur fixe pour les bulles clean (la generation cote app.py fixe leur px)
    try:
        if "_clean" in Path(image_path).name.lower():
            return CLEAN_BUBBLE_CM
    except Exception:
        pass
    target = px_w / BUBBLE_PX_PER_CM
    if target < BUBBLE_MIN_CM:
        return BUBBLE_MIN_CM
    if target > BUBBLE_MAX_CM:
        return BUBBLE_MAX_CM
    return target


def compress_image_for_docx(image_path, max_width=IMAGE_MAX_WIDTH, quality=IMAGE_QUALITY, preserve_transparency=False):
    """
    Compresse une image pour réduire la taille du fichier Word.
    - Redimensionne si trop grande
    - Convertit en JPEG (plus petit que PNG) sauf si preserve_transparency=True
    - Retourne un BytesIO avec l'image compressée
    """
    try:
        with Image.open(image_path) as img:
            # Vérifier si l'image a de la transparence à préserver
            # (bulles de conversation avec coins arrondis)
            has_transparency = img.mode in ('RGBA', 'LA', 'P')

            # Détecter automatiquement les bulles de message (fichiers msg_*.png)
            path_str = str(image_path).lower()
            is_bubble_image = 'msg_' in path_str and path_str.endswith('.png')

            if (preserve_transparency or is_bubble_image) and has_transparency:
                # Préserver la transparence - sauvegarder en PNG
                if img.mode == 'P':
                    img = img.convert('RGBA')
                elif img.mode == 'LA':
                    img = img.convert('RGBA')

                # Redimensionner si trop grand
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

                # Sauvegarder en PNG avec transparence
                output = io.BytesIO()
                img.save(output, format='PNG', optimize=True)
                output.seek(0)
                return output
            else:
                # Convertir RGBA en RGB (JPEG ne supporte pas la transparence)
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Créer un fond blanc
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')

                # Redimensionner si trop grand
                if img.width > max_width:
                    ratio = max_width / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

                # Compresser en JPEG
                output = io.BytesIO()
                img.save(output, format='JPEG', quality=quality, optimize=True)
                output.seek(0)
                return output
    except Exception as e:
        print(f"[COMPRESS] Erreur compression {image_path}: {e}")
        return None

PLACEHOLDER_PATTERN = re.compile(r"\{[^{}]+\}")
DEFAULT_ANALYSIS_TEMPLATE = ""
BACK_TOKEN = "__BACK__"
IMAGE_MARKER = re.compile(r"\[\[\s*IMG\s*:\s*([^\]]+?)\s*\]\]")


def disable_clickable_emails(doc):
    """
    Désactive les emails/URLs cliquables dans le document Word.
    1. Supprime tous les hyperliens explicites
    2. Remplace @ par @ + caractère invisible dans TOUT le XML du document
    """
    from docx.oxml.ns import qn

    # Zero-width space - caractère invisible qui casse le pattern email
    ZERO_WIDTH_SPACE = '\u200B'

    # --- ÉTAPE 1: Supprimer les hyperliens explicites ---
    body = doc._element.body

    for hyperlink in list(body.iter(qn('w:hyperlink'))):
        parent = hyperlink.getparent()
        if parent is None:
            continue
        index = list(parent).index(hyperlink)
        runs = list(hyperlink.iter(qn('w:r')))
        for run in runs:
            parent.insert(index, run)
            index += 1
        parent.remove(hyperlink)

    # Headers et footers - supprimer hyperliens
    for section in doc.sections:
        for header in [section.header, section.footer]:
            if header is None or header._element is None:
                continue
            for hyperlink in list(header._element.iter(qn('w:hyperlink'))):
                parent = hyperlink.getparent()
                if parent is None:
                    continue
                index = list(parent).index(hyperlink)
                runs = list(hyperlink.iter(qn('w:r')))
                for run in runs:
                    parent.insert(index, run)
                    index += 1
                parent.remove(hyperlink)

    # --- ÉTAPE 2: Parcourir TOUT le XML et remplacer @ dans les éléments <w:t> ---
    # Cela capture TOUS les textes, même dans les tableaux imbriqués
    count = 0
    for t_elem in doc._element.body.iter(qn('w:t')):
        if t_elem.text and '@' in t_elem.text:
            t_elem.text = t_elem.text.replace('@', '@' + ZERO_WIDTH_SPACE)
            count += 1

    # Traiter tous les headers et footers
    for section in doc.sections:
        if section.header and section.header._element is not None:
            for t_elem in section.header._element.iter(qn('w:t')):
                if t_elem.text and '@' in t_elem.text:
                    t_elem.text = t_elem.text.replace('@', '@' + ZERO_WIDTH_SPACE)
                    count += 1
        if section.footer and section.footer._element is not None:
            for t_elem in section.footer._element.iter(qn('w:t')):
                if t_elem.text and '@' in t_elem.text:
                    t_elem.text = t_elem.text.replace('@', '@' + ZERO_WIDTH_SPACE)
                    count += 1

    print(f"[INFO] {count} éléments avec @ modifiés pour désactiver les liens cliquables")


def update_table_of_contents(doc):
    """
    Force la mise à jour automatique de la table des matières à l'ouverture du document.
    Ajoute le flag updateFields dans les settings du document.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Accéder aux settings du document
    settings = doc.settings._element

    # Créer l'élément updateFields s'il n'existe pas
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)

    # Définir la valeur à true pour forcer la mise à jour
    update_fields.set(qn('w:val'), 'true')

    print("[INFO] Table des matières configurée pour mise à jour automatique à l'ouverture")


def fix_heading_numbering_tabs(doc):
    """
    Corrige les tabulations dans la numérotation des titres pour éviter
    les écarts quand les numéros ont plus de chiffres (ex: 2.2.10 vs 2.2.9).
    Change les tabulations en espaces simples.
    """
    from docx.oxml.ns import qn

    # Accéder au numbering part si disponible
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return
    except:
        return

    numbering_element = numbering_part._element

    # Parcourir tous les niveaux de numérotation et modifier le suffixe de tab à space
    for lvl in numbering_element.iter(qn('w:lvl')):
        # Trouver l'élément suff (suffix) qui définit ce qui suit le numéro
        suff = lvl.find(qn('w:suff'))
        if suff is not None:
            # Changer de 'tab' à 'space' pour avoir un simple espace après le numéro
            current_val = suff.get(qn('w:val'))
            if current_val == 'tab':
                suff.set(qn('w:val'), 'space')

    print("[INFO] Tabulations des numéros de titres corrigées")


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def find_placeholders_in_order(doc):
    seen = set()
    ordered = []

    # Chercher dans les en-têtes de toutes les sections
    for section in doc.sections:
        header = section.header
        # Paragraphes dans l'en-tête
        for p in header.paragraphs:
            for ph in PLACEHOLDER_PATTERN.findall(p.text):
                if ph not in seen:
                    ordered.append(ph)
                    seen.add(ph)
        # Tableaux dans l'en-tête
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for ph in PLACEHOLDER_PATTERN.findall(p.text):
                            if ph not in seen:
                                ordered.append(ph)
                                seen.add(ph)

    # Chercher dans le corps du document
    for p in iter_all_paragraphs(doc):
        for ph in PLACEHOLDER_PATTERN.findall(p.text):
            if ph not in seen:
                ordered.append(ph)
                seen.add(ph)
    return ordered


def find_image_markers_in_order(doc):
    """Extract all [[IMG:key]] markers from the document in order."""
    seen = set()
    ordered = []

    # Search in headers of all sections
    for section in doc.sections:
        header = section.header
        # Paragraphs in header
        for p in header.paragraphs:
            for marker in IMAGE_MARKER.findall(p.text):
                if marker not in seen:
                    ordered.append(marker)
                    seen.add(marker)
        # Tables in header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for marker in IMAGE_MARKER.findall(p.text):
                            if marker not in seen:
                                ordered.append(marker)
                                seen.add(marker)

    # Search in document body
    for p in iter_all_paragraphs(doc):
        for marker in IMAGE_MARKER.findall(p.text):
            if marker not in seen:
                ordered.append(marker)
                seen.add(marker)
    return ordered


def is_in_table_cell(paragraph):
    """Vérifie si le paragraphe est à l'intérieur d'une cellule de tableau."""
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith('}tc'):  # tc = table cell
            return True
        parent = parent.getparent()
    return False


def replace_in_runs(paragraph, mapping, allow_delete=True):
    """Remplace placeholders coupes en runs. Valeur vide -> supprime le paragraphe entier (sauf dans les cellules de tableau ou les titres)."""
    if not mapping:
        return False
    original = paragraph.text
    new_text = original
    remove_entire = False

    for old, new in mapping.items():
        if old not in new_text:
            continue
        if new == "":
            # Ne supprimer le paragraphe que si on est HORS d'une cellule de tableau ET que ce n'est pas un titre
            if allow_delete and not is_in_table_cell(paragraph) and not is_heading(paragraph):
                remove_entire = True
                break
            else:
                # Dans une cellule de tableau ou un titre, on remplace juste par une chaîne vide
                new_text = new_text.replace(old, new)
        else:
            new_text = new_text.replace(old, new)

    if remove_entire:
        remove_paragraph(paragraph)
        return True
    if new_text == original:
        return False
    for run in list(paragraph.runs):
        run.text = ""
    paragraph.add_run(new_text)
    return False


def fill_with_mapping(text, mapping):
    out = text
    for old, new in mapping.items():
        out = out.replace(old, new)
    return out


def is_heading(paragraph):
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.startswith("Heading") or style_name.startswith("Titre")


def get_heading_level(paragraph):
    """Retourne le niveau du heading (1, 2, 3...) ou 1 par défaut."""
    style_name = paragraph.style.name if paragraph.style else ""
    # Extraire le numéro du style (Heading 1, Heading 2, Titre 1, etc.)
    import re
    match = re.search(r'(\d+)$', style_name)
    if match:
        return int(match.group(1))
    return 1  # Par défaut niveau 1


def insert_after(paragraph, text, style=None, space_after_pt=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    # Ajouter de l'espacement après le paragraphe si spécifié
    if space_after_pt is not None:
        new_para.paragraph_format.space_after = Pt(space_after_pt)
    return new_para


def insert_after_in_cell(paragraph, text, style=None):
    """Insère un nouveau paragraphe après le paragraphe donné, même dans une cellule de tableau."""
    # Utiliser toujours insert_after qui fonctionne aussi dans les cellules de tableau
    return insert_after(paragraph, text, style)


def remove_table(table):
    tbl_element = table._element
    parent = tbl_element.getparent()
    if parent is not None:
        parent.remove(tbl_element)
    if table in table._parent.tables:
        try:
            table._parent.tables.remove(table)
        except ValueError:
            pass


def remove_empty_paragraphs(doc):
    for p in list(iter_all_paragraphs(doc)):
        if not p.text or not p.text.strip():
            # Ne pas supprimer les paragraphes vides dans les cellules de tableau
            if not is_in_table_cell(p):
                remove_paragraph(p)


def remove_empty_sim_tables(doc, mapping):
    """Supprime les tableaux SIM dont tous les placeholders sont vides, ainsi que leurs sous-titres associés."""
    tables_to_remove = []
    sim_indices_to_remove = []

    print(f"[DEBUG remove_empty_sim_tables] Nombre de tableaux dans le document: {len(doc.tables)}")
    print(f"[DEBUG remove_empty_sim_tables] Clés SIM dans le mapping:")
    for i in range(1, 9):
        sim_keys = [f"{{operateur{i}}}", f"{{iccid{i}}}", f"{{imsi{i}}}", f"{{msisdn{i}}}", f"{{datesync{i}}}"]
        for key in sim_keys:
            if key in mapping:
                print(f"  {key} = '{mapping[key]}'")

    for table_idx, table in enumerate(doc.tables):
        # Vérifier si c'est un tableau SIM (a une ligne avec des placeholders operateur/iccid/imsi/msisdn/datesync)
        is_sim_table = False
        sim_index = None
        has_at_least_one_value = False

        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            # Chercher des placeholders SIM indexés
            for i in range(1, 9):
                sim_keys = [f"{{operateur{i}}}", f"{{iccid{i}}}", f"{{imsi{i}}}", f"{{msisdn{i}}}", f"{{datesync{i}}}"]
                if any(key in row_text for key in sim_keys):
                    is_sim_table = True
                    sim_index = i
                    print(f"[DEBUG] Tableau {table_idx} identifié comme SIM {i}, row_text contient: {[k for k in sim_keys if k in row_text]}")
                    # Vérifier si AU MOINS UNE valeur est remplie dans le mapping pour cette carte SIM
                    for key in sim_keys:
                        value = mapping.get(key, "").strip()
                        print(f"[DEBUG]   Vérification {key} -> '{value}'")
                        if value:
                            has_at_least_one_value = True
                            print(f"[DEBUG]   => Valeur trouvée pour {key}!")
                            break
                    break
            if is_sim_table:
                break

        # Ne supprimer que si c'est un tableau SIM ET qu'aucune valeur n'est remplie
        if is_sim_table and not has_at_least_one_value:
            print(f"[DEBUG] Suppression du tableau SIM {sim_index} - aucune valeur remplie")
            tables_to_remove.append(table)
            sim_indices_to_remove.append(sim_index)
        elif is_sim_table:
            print(f"[DEBUG] Conservation du tableau SIM {sim_index} - au moins une valeur remplie")

    # Supprimer les sous-titres "Carte SIM n°X" associés aux tableaux supprimés
    paragraphs_to_remove = []
    for p in doc.paragraphs:
        p_text = p.text.strip()
        for sim_idx in sim_indices_to_remove:
            # Chercher les variantes du titre de carte SIM
            patterns = [
                f"Carte SIM n°{sim_idx}",
                f"Carte SIM n° {sim_idx}",
                f"Carte SIM {sim_idx}",
                f"SIM n°{sim_idx}",
                f"SIM {sim_idx}",
            ]
            for pattern in patterns:
                if pattern in p_text:
                    print(f"[DEBUG] Suppression du sous-titre: '{p_text}' (SIM {sim_idx})")
                    paragraphs_to_remove.append(p)
                    break

    # Supprimer les paragraphes marqués
    for p in paragraphs_to_remove:
        remove_paragraph(p)

    # Supprimer les tableaux marqués
    for table in tables_to_remove:
        remove_table(table)


def add_spacing_between_tables(doc, spacing_pt=12):
    """
    Ajoute un espacement entre les tableaux pour améliorer la lisibilité.
    Insère un paragraphe vide avec espacement après chaque tableau.

    Args:
        doc: Document python-docx
        spacing_pt: Espacement en points (défaut: 12pt)
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    from docx.shared import Pt

    # Parcourir tous les tableaux du document
    tables = doc.tables

    if len(tables) <= 1:
        return  # Pas besoin d'espacement avec 0 ou 1 tableau

    # Traiter chaque tableau (sauf le dernier)
    for i in range(len(tables) - 1):
        table = tables[i]
        table_element = table._element

        # Vérifier s'il y a déjà un paragraphe après ce tableau
        next_sibling = table_element.getnext()

        # Si le prochain élément est un autre tableau, insérer un paragraphe
        if next_sibling is not None and next_sibling.tag.endswith('}tbl'):
            # Créer un nouveau paragraphe vide
            new_para_element = OxmlElement('w:p')

            # Insérer le paragraphe après le tableau
            table_element.addnext(new_para_element)

            # Créer un objet Paragraph pour accéder aux propriétés de format
            new_para = Paragraph(new_para_element, doc)

            # Appliquer l'espacement
            new_para.paragraph_format.space_after = Pt(spacing_pt)
            new_para.paragraph_format.space_before = Pt(0)

            print(f"[DEBUG] Espacement de {spacing_pt}pt ajouté après le tableau {i+1}")


def reduce_spacing_for_first_tables(doc, num_tables=3):
    """
    Réduit les espacements entre les premiers tableaux pour qu'ils tiennent sur une page.
    Spécifiquement pour test2.docx où les 3 premiers tableaux doivent rester sur la page 1.

    Args:
        doc: Document python-docx
        num_tables: Nombre de tableaux à traiter (défaut: 3)
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tables = doc.tables
    if len(tables) < num_tables:
        return

    body = doc._body._body
    tables_found = 0

    for child in body:
        tag = child.tag.split('}')[-1]

        if tag == 'tbl':
            tables_found += 1

            if tables_found <= num_tables:
                # Réduire les marges internes du tableau
                tblPr = child.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    child.insert(0, tblPr)

                tblCellMar = tblPr.find(qn('w:tblCellMar'))
                if tblCellMar is None:
                    tblCellMar = OxmlElement('w:tblCellMar')
                    tblPr.append(tblCellMar)

                # Marges top/bottom réduites à 1pt
                for margin_name in ['top', 'bottom']:
                    margin = tblCellMar.find(qn(f'w:{margin_name}'))
                    if margin is None:
                        margin = OxmlElement(f'w:{margin_name}')
                        tblCellMar.append(margin)
                    margin.set(qn('w:w'), '20')  # ~1pt
                    margin.set(qn('w:type'), 'dxa')

                # Réduire l'espacement du paragraphe suivant
                next_elem = child.getnext()
                if next_elem is not None:
                    next_tag = next_elem.tag.split('}')[-1]
                    if next_tag == 'p':
                        pPr = next_elem.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            next_elem.insert(0, pPr)

                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)

                        # Espacement minimal
                        spacing.set(qn('w:before'), '0')
                        spacing.set(qn('w:after'), '0')
                        spacing.set(qn('w:line'), '240')
                        spacing.set(qn('w:lineRule'), 'auto')

    print(f"[DEBUG] Espacements réduits pour les {num_tables} premiers tableaux (test2)")


def add_page_break_before_heading(doc, heading_text_contains):
    """
    Ajoute un saut de page avant un titre spécifique dans le document.
    Cherche d'abord dans les paragraphes normaux, puis dans les tableaux.

    Args:
        doc: Document python-docx
        heading_text_contains: Texte partiel à rechercher dans le titre
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 1. Chercher dans les paragraphes normaux (headings)
    for para in doc.paragraphs:
        if is_heading(para) and heading_text_contains.lower() in para.text.lower():
            # Créer l'élément de saut de page
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')

            # L'insérer au tout début du paragraphe
            first_run_element = para._element.find('.//' + qn('w:r'))
            if first_run_element is not None:
                first_run_element.insert(0, br)
            else:
                new_run_elem = OxmlElement('w:r')
                new_run_elem.append(br)
                para._element.insert(0, new_run_elem)

            print(f"[DEBUG] Saut de page ajouté avant le titre: '{para.text[:50]}...'")
            return True

    # 2. Chercher dans les tableaux (titres encadrés)
    for table in doc.tables:
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "

        if heading_text_contains.lower() in table_text.lower():
            # Créer un paragraphe avec saut de page avant le tableau
            new_para = OxmlElement('w:p')
            new_run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            new_run.append(br)
            new_para.append(new_run)

            # Insérer le paragraphe avec saut de page AVANT le tableau
            table._element.addprevious(new_para)

            print(f"[DEBUG] Saut de page ajouté avant le tableau contenant: '{heading_text_contains}'")
            return True

    print(f"[ATTENTION] Titre contenant '{heading_text_contains}' non trouvé pour le saut de page")
    return False


def insert_extra_support_images(doc, extra_support_images, section_text="Photographies des supports numériques"):
    """
    Insère des images supplémentaires dans la section "Photographies des supports numériques".
    Les images sont ajoutées dans la même cellule que les photos existantes (pas dans Commentaires).

    Args:
        doc: Document python-docx
        extra_support_images: Liste de dicts avec {path, width_inches}
        section_text: Texte à rechercher pour identifier la section
    """
    if not extra_support_images:
        print("[DEBUG] Pas d'images supplémentaires à insérer")
        return

    print(f"[DEBUG] Insertion de {len(extra_support_images)} images supplémentaires")
    for i, img in enumerate(extra_support_images):
        print(f"[DEBUG]   Image {i}: path={img.get('path')}, width={img.get('width_inches')}")

    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Trouver le tableau contenant "Photographies des supports numériques"
    target_table = None
    header_row_idx = -1

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip().lower()
                if "photographies des supports" in cell_text or "photographie des supports" in cell_text:
                    target_table = table
                    header_row_idx = row_idx
                    print(f"[DEBUG] Trouvé tableau Photographies, row={row_idx}, cell={cell_idx}")
                    break
            if target_table:
                break
        if target_table:
            break

    if not target_table:
        print(f"[ATTENTION] Tableau '{section_text}' non trouvé, insertion à la fin du document")
        for idx, img_data in enumerate(extra_support_images):
            img_path = img_data.get("path")
            width_inches = img_data.get("width_inches", 1.5)
            if not img_path or not Path(img_path).exists():
                continue
            try:
                img_para = doc.add_paragraph()
                img_para.paragraph_format.space_before = Pt(12)
                img_run = img_para.add_run()
                # Compresser l'image pour réduire la taille du fichier
                compressed = compress_image_for_docx(img_path)
                if compressed:
                    img_run.add_picture(compressed, width=Inches(width_inches))
                else:
                    img_run.add_picture(img_path, width=Inches(width_inches))
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"[ERREUR] {e}")
        return

    # Trouver la cellule avec les photos (après le header, AVANT "Commentaires")
    # Structure: Row 0 = header, Row 1 = photos, Row 2 = Commentaires
    target_cell = None

    # Chercher la cellule juste après le header qui n'est PAS "Commentaires"
    if len(target_table.rows) > header_row_idx + 1:
        for row_idx in range(header_row_idx + 1, len(target_table.rows)):
            row = target_table.rows[row_idx]
            for cell in row.cells:
                cell_text = cell.text.strip().lower()
                # Ignorer la cellule "Commentaires"
                if "commentaire" in cell_text:
                    print(f"[DEBUG] Row {row_idx}: cellule Commentaires ignorée")
                    continue
                # Prendre la première cellule qui n'est pas "Commentaires"
                target_cell = cell
                print(f"[DEBUG] Cellule photos trouvée à row {row_idx}")
                break
            if target_cell:
                break

    if not target_cell:
        print(f"[ATTENTION] Impossible de trouver la cellule photos")
        return

    # Ajouter les images dans la cellule des photos (sans label "Support X")
    images_added = 0
    for idx, img_data in enumerate(extra_support_images):
        img_path = img_data.get("path")
        width_inches = img_data.get("width_inches", 1.5)

        if not img_path:
            print(f"[ATTENTION] Image {idx}: pas de chemin")
            continue

        if not Path(img_path).exists():
            print(f"[ATTENTION] Image non trouvée: {img_path}")
            continue

        try:
            # Ajouter l'image directement (sans titre "Support X")
            img_para = target_cell.add_paragraph()
            img_para.paragraph_format.space_before = Pt(12)
            img_para.paragraph_format.space_after = Pt(6)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = img_para.add_run()
            # Compresser l'image pour réduire la taille du fichier
            compressed = compress_image_for_docx(img_path)
            if compressed:
                run.add_picture(compressed, width=Inches(width_inches))
            else:
                run.add_picture(img_path, width=Inches(width_inches))

            images_added += 1
            print(f"[DEBUG] Image ajoutée ({width_inches} pouces) - {img_path}")
        except Exception as e:
            print(f"[ERREUR] {e}")
            import traceback
            traceback.print_exc()

    print(f"[DEBUG] {images_added}/{len(extra_support_images)} images insérées")


def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("tbl"):
            yield Table(child, doc)


def prompt_placeholders(placeholders):
    mapping = {}
    if not placeholders:
        print("Aucun placeholder { } trouve dans le document.")
        return mapping
    print("Placeholders trouves dans l'ordre :", ", ".join(placeholders))
    for ph in placeholders:
        value = input(f"Valeur pour {ph} (laisser vide = supprime le paragraphe contenant ce placeholder) : ")
        mapping[ph] = value
    return mapping 


def prompt_phrase_for_heading(title_text, mapping):
    print(f"Titre detecte : {title_text}")
    default_phrase = fill_with_mapping(DEFAULT_ANALYSIS_TEMPLATE, mapping)
    print(" 1) Saisir une phrase personnalisee")
    print(f" 2) {default_phrase}")
    print(" 3) Rien (supprime le titre et les tableaux qui suivent)")
    print(" b) Retour au titre precedent")
    choice = input("Choix (1/2/3/b, Enter=2) : ").strip().lower()
    if choice == "b":
        return BACK_TOKEN
    if choice == "1":
        return input("Votre phrase (laisser vide = suppression du titre/tableaux) : ").strip()
    if choice == "3":
        return ""
    if choice == "" or choice == "2":
        return default_phrase
    return default_phrase


def collect_headings_in_order(doc):
    return [p for p in iter_all_paragraphs(doc) if isinstance(p, Paragraph) and is_heading(p)]


def collect_headings_with_levels(doc):
    """Retourne une liste de dict {text, level} pour chaque heading."""
    headings = []
    for p in iter_all_paragraphs(doc):
        if isinstance(p, Paragraph) and is_heading(p):
            headings.append({
                "text": p.text.strip(),
                "level": get_heading_level(p)
            })
    return headings


def collect_heading_decisions(headings, mapping):
    decisions = []
    idx = 0
    while idx < len(headings):
        title_text = headings[idx].text.strip()
        res = prompt_phrase_for_heading(title_text, mapping)
        if res == BACK_TOKEN:
            if idx > 0:
                decisions.pop()
                idx -= 1
            else:
                print("Deja au premier titre, impossible de revenir en arriere.")
            continue
        if idx < len(decisions):
            decisions[idx] = res
        else:
            decisions.append(res)
        idx += 1
    return decisions


def normalize_heading_title(title: str) -> str:
    """Normalise un titre pour la comparaison (supprime numérotation, espaces, accents)."""
    import unicodedata
    t = title.strip()
    # Supprimer la numérotation au début (ex: "2.3.1 " ou "2.3.1. ")
    t = re.sub(r'^[\d.]+\s*', '', t)
    # Normaliser les accents et mettre en minuscules
    t = unicodedata.normalize('NFD', t).encode('ascii', 'ignore').decode('ascii').lower()
    # Supprimer les caractères spéciaux
    t = re.sub(r'[^a-z0-9\s]', '', t)
    # Normaliser les espaces
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def apply_heading_decisions(doc, decisions, decisions_by_title: Optional[Dict[str, str]] = None):
    """
    Applique les décisions aux headings du document.
    Si decisions_by_title est fourni, utilise le mapping par titre (plus fiable).
    Sinon, utilise l'ancien système par index (pour rétrocompatibilité).
    """
    blocks = list(iter_block_items(doc))
    heading_idx = 0
    idx = 0

    # Debug: afficher les décisions et headings
    all_headings = [b.text.strip() for b in blocks if isinstance(b, Paragraph) and is_heading(b)]
    print(f"[DEBUG apply_heading_decisions] Headings dans doc: {all_headings}")
    print(f"[DEBUG apply_heading_decisions] Nb decisions: {len(decisions)}, Nb headings: {len(all_headings)}")
    if decisions_by_title:
        print(f"[DEBUG apply_heading_decisions] Utilisation du mapping par titre ({len(decisions_by_title)} entrées)")

    while idx < len(blocks):
        blk = blocks[idx]
        if isinstance(blk, Paragraph) and is_heading(blk):
            heading_title = blk.text.strip()

            # Chercher la décision par titre (prioritaire) ou par index (fallback)
            decision = None
            if decisions_by_title:
                # Essayer d'abord le titre exact
                decision = decisions_by_title.get(heading_title)
                # Sinon essayer avec le titre normalisé
                if decision is None:
                    normalized = normalize_heading_title(heading_title)
                    for key, val in decisions_by_title.items():
                        if normalize_heading_title(key) == normalized:
                            decision = val
                            break

            # Fallback sur l'index si pas trouvé par titre
            if decision is None:
                decision = decisions[heading_idx] if heading_idx < len(decisions) else ""

            print(f"[DEBUG] Heading {heading_idx}: '{heading_title}' -> decision: '{decision[:50] if decision else '(vide)'}...'")
            heading_idx += 1

            if decision == "__KEEP_TITLE_ONLY__":
                # Garder le titre mais ne rien ajouter en dessous
                pass
            elif decision:
                # Ajouter la phrase sous le titre
                insert_after(blk, decision, style=None)
            else:
                # Supprimer le titre et les tableaux qui suivent
                remove_paragraph(blk)
                blocks.pop(idx)
                j = idx
                while j < len(blocks):
                    nxt = blocks[j]
                    if isinstance(nxt, Paragraph) and is_heading(nxt):
                        break
                    if isinstance(nxt, Table):
                        remove_table(nxt)
                        blocks.pop(j)
                        continue
                    j += 1
                idx = j
                continue
        idx += 1


def default_heading_decisions(headings: List[Paragraph], mapping: Dict[str, str]) -> List[str]:
    """Genere une phrase par defaut sous chaque titre (non interactif)."""
    return [fill_with_mapping(DEFAULT_ANALYSIS_TEMPLATE, mapping) for _ in headings]

def insert_image_after(paragraph: Paragraph, image_path: str, width_inches: float = 1.5, text_before: str = "", text_after: str = ""):
    """Insere une image juste apres le paragraphe donne, avec optionnellement du texte avant/après."""
    # Vérifier que le fichier image existe
    img_path = Path(image_path)
    if not img_path.exists():
        print(f"ATTENTION: Image introuvable: {image_path}")
        return

    if text_before:
        new_p = insert_after(paragraph, text_before)
        paragraph = new_p
    new_p = insert_after(paragraph, "")
    run = new_p.add_run()
    try:
        # Compresser l'image pour réduire la taille du fichier
        compressed = compress_image_for_docx(str(img_path))
        if compressed:
            run.add_picture(compressed, width=Inches(width_inches))
        else:
            run.add_picture(str(img_path), width=Inches(width_inches))
    except Exception as e:
        print(f"ERREUR lors de l'insertion de l'image {image_path}: {type(e).__name__}: {str(e)}")
        return
    if text_after:
        insert_after(new_p, text_after)


def apply_images_after_headings(doc: Document, images_after: Dict[str, str], width_inches: float = 1.5,
                                per_image_widths: Optional[Dict[str, float]] = None,
                                image_texts: Optional[Dict[str, Dict[str, str]]] = None):
    """Insere des images apres les titres dont le texte matche exactement la cle du mapping."""
    if not images_after:
        return

    # Parcourir tous les paragraphes et identifier les headings
    all_paras = list(doc.paragraphs)
    for idx, para in enumerate(all_paras):
        if is_heading(para):
            heading_text = para.text.strip()
            img = images_after.get(heading_text)
            if img:
                w = per_image_widths.get(heading_text) if per_image_widths else None
                text_data = image_texts.get(heading_text, {}) if image_texts else {}
                text_before = text_data.get("before", "") if text_data.get("position") == "before" else ""
                text_after = text_data.get("after", "") if text_data.get("position") == "after" else ""

                # Trouver le paragraphe suivant qui n'est pas un heading
                target_para = para
                if idx + 1 < len(all_paras):
                    next_para = all_paras[idx + 1]
                    if not is_heading(next_para):
                        # Le prochain paragraphe est probablement la phrase automatique
                        target_para = next_para

                insert_image_after(target_para, img, w or width_inches, text_before, text_after)


def apply_images_after_paragraphs(doc: Document, images_after: Dict[str, str], width_inches: float = 1.5,
                                  per_image_widths: Optional[Dict[str, float]] = None,
                                  image_texts: Optional[Dict[str, Dict[str, str]]] = None):
    """Insere des images apres les paragraphes dont le texte matche exactement la cle du mapping."""
    if not images_after:
        return

    # Utiliser un set pour éviter les insertions multiples du même paragraphe
    processed = set()
    for p in iter_all_paragraphs(doc):
        # Créer un identifiant unique basé sur le contenu et la position
        para_id = id(p._element)
        if para_id in processed:
            continue

        key = p.text.strip()
        img = images_after.get(key)
        if img:
            w = per_image_widths.get(key) if per_image_widths else None
            text_data = image_texts.get(key, {}) if image_texts else {}
            text_before = text_data.get("before", "") if text_data.get("position") == "before" else ""
            text_after = text_data.get("after", "") if text_data.get("position") == "after" else ""
            insert_image_after(p, img, w or width_inches, text_before, text_after)
            processed.add(para_id) 


def _insert_image_to_paragraph(p, img_path: Path, width_inches: float):
    """Helper pour insérer une image dans un paragraphe avec compression."""
    try:
        # Utiliser la compression pour réduire la taille du fichier Word
        compressed = compress_image_for_docx(str(img_path))
        if compressed:
            p.add_run().add_picture(compressed, width=Inches(width_inches))
            return True
        # Fallback: essayer sans compression
        p.add_run().add_picture(str(img_path), width=Inches(width_inches))
        return True
    except Exception as e:
        print(f"[FALLBACK] Erreur insertion ({type(e).__name__}), reconversion...")
        try:
            img = Image.open(img_path)
            if img.mode in ('RGBA', 'LA', 'P'):
                rgb_image = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                if img.mode == 'RGBA':
                    rgb_image.paste(img, mask=img.split()[-1])
                else:
                    rgb_image.paste(img)
                img = rgb_image
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            # Compresser avec les mêmes paramètres
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                tmp_path = Path(tmp.name)
                # Redimensionner si trop grand
                if img.width > IMAGE_MAX_WIDTH:
                    ratio = IMAGE_MAX_WIDTH / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((IMAGE_MAX_WIDTH, new_height), Image.Resampling.LANCZOS)
                img.save(tmp_path, 'JPEG', quality=IMAGE_QUALITY, optimize=True)
            p.add_run().add_picture(str(tmp_path), width=Inches(width_inches))
            try:
                tmp_path.unlink()
            except:
                pass
            return True
        except Exception as e2:
            print(f"ERREUR FALLBACK: {type(e2).__name__}: {str(e2)}")
            return False


def apply_images_at_markers(doc: Document, images_at_markers: Dict[str, str], width_inches: float = 1.5,
                            per_image_widths: Optional[Dict[str, float]] = None,
                            image_texts: Optional[Dict[str, Dict[str, str]]] = None):
    """Remplace les marqueurs [[IMG:cle]] par l'image correspondante inseree a cet endroit."""
    if not images_at_markers:
        return

    for p in iter_all_paragraphs(doc):
        full_text = "".join(run.text for run in p.runs)
        if "[[IMG" not in full_text:
            continue

        # Cas spécial: supp1 et supp2 sur la même ligne -> tableau côte à côte
        has_supp1 = "[[IMG:supp1]]" in full_text
        has_supp2 = "[[IMG:supp2]]" in full_text

        if has_supp1 and has_supp2:
            # Créer un tableau 1x2 pour les images côte à côte
            supp1_path = images_at_markers.get("supp1")
            supp2_path = images_at_markers.get("supp2")

            # Effacer le contenu du paragraphe
            for run in list(p.runs):
                run.text = ""

            # Si au moins une image existe, créer le tableau
            if supp1_path or supp2_path:
                # Créer un tableau 1x2 et l'insérer après le paragraphe
                tbl = doc.add_table(rows=1, cols=2)
                tbl.autofit = False

                # Déplacer le tableau à la position du paragraphe
                p._p.addnext(tbl._tbl)

                # Cellule gauche (supp1)
                left_cell = tbl.cell(0, 0)
                left_para = left_cell.paragraphs[0]
                left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if supp1_path and Path(supp1_path).exists():
                    w1 = per_image_widths.get("supp1") if per_image_widths else None
                    w1 = w1 or per_image_widths.get("marker:supp1") if per_image_widths else None
                    _insert_image_to_paragraph(left_para, Path(supp1_path), w1 or width_inches)

                # Cellule droite (supp2)
                right_cell = tbl.cell(0, 1)
                right_para = right_cell.paragraphs[0]
                right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if supp2_path and Path(supp2_path).exists():
                    w2 = per_image_widths.get("supp2") if per_image_widths else None
                    w2 = w2 or per_image_widths.get("marker:supp2") if per_image_widths else None
                    _insert_image_to_paragraph(right_para, Path(supp2_path), w2 or width_inches)

                # Supprimer les bordures du tableau (tableau invisible)
                tbl_element = tbl._tbl
                tbl_pr = tbl_element.tblPr
                if tbl_pr is None:
                    tbl_pr = OxmlElement('w:tblPr')
                    tbl_element.insert(0, tbl_pr)
                tbl_borders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')
                    tbl_borders.append(border)
                tbl_pr.append(tbl_borders)

                # Définir les largeurs des cellules
                try:
                    left_cell.width = Cm(8)
                    right_cell.width = Cm(8)
                except:
                    pass

            continue  # Passer au paragraphe suivant

        # Traitement standard pour les autres marqueurs
        parts = IMAGE_MARKER.split(full_text)  # [texte, cle, texte, cle, ...]
        for run in list(p.runs):
            run.text = ""
        for idx, chunk in enumerate(parts):
            if idx % 2 == 0:
                if chunk:
                    p.add_run(chunk)
            else:
                key = chunk.strip()
                img_path = images_at_markers.get(key)
                if img_path:
                    img_file = Path(img_path)
                    if not img_file.exists():
                        print(f"ATTENTION: Image introuvable pour le marqueur '{key}': {img_path}")
                        p.add_run(f"[[IMG:{key} - INTROUVABLE]]")
                        continue

                    text_data = image_texts.get(key, {}) if image_texts else {}
                    if text_data.get("position") == "before" and text_data.get("before"):
                        p.add_run(text_data["before"] + " ")

                    w = per_image_widths.get(key) if per_image_widths else None
                    if not _insert_image_to_paragraph(p, img_file, w or width_inches):
                        p.add_run(f"[[IMG:{key} - ERREUR]]")
                        continue

                    if text_data.get("position") == "after" and text_data.get("after"):
                        p.add_run(" " + text_data["after"])
                else:
                    p.add_run(f"[[IMG:{key}]]")


def apply_heading_content_blocks(doc: Document, heading_content: Dict[str, List[Dict]], default_width_inches: float = 1.5, mapping: Optional[Dict[str, str]] = None, account_images: Optional[Dict[str, str]] = None):
    """Insere les blocs de contenu (texte/images) apres chaque heading specifie."""
    # Déclaration explicite pour utiliser les imports globaux
    global OxmlElement, Paragraph

    if not heading_content:
        return

    if account_images is None:
        account_images = {}

    def replace_placeholders(text: str) -> str:
        """Remplace les placeholders dans le texte avec les valeurs du mapping."""
        if not mapping or not text:
            return text
        result = text
        for placeholder, value in mapping.items():
            result = result.replace(placeholder, value or "")
        return result

    def get_next_sibling_paragraph(para):
        """Trouve le prochain paragraphe frère dans la structure XML du document."""
        current_element = para._p
        next_element = current_element.getnext()
        while next_element is not None:
            if next_element.tag.endswith('}p'):
                # C'est un paragraphe
                return Paragraph(next_element, para._parent)
            elif next_element.tag.endswith('}tbl'):
                # C'est un tableau, pas de paragraphe direct après
                return None
            next_element = next_element.getnext()
        return None

    # Debug: afficher les clés attendues vs les headings du document
    print(f"[DEBUG apply_heading_content_blocks] Clés heading_content: {list(heading_content.keys())}")

    # Utiliser iter_block_items pour maintenir l'ordre correct du document
    block_items = list(iter_block_items(doc))
    headings_in_doc = [b.text.strip() for b in block_items if isinstance(b, Paragraph) and is_heading(b)]
    print(f"[DEBUG apply_heading_content_blocks] Headings dans le doc: {headings_in_doc}")

    # Créer un mapping normalisé pour matching flexible
    def normalize_title(title):
        """Normalise un titre pour le matching (trim, lowercase, espaces multiples)"""
        import re
        return re.sub(r'\s+', ' ', title.strip().lower())

    scoped_prefix = "__PLATFORM__|||"
    scoped_heading_content: Dict[tuple, tuple] = {}
    flat_heading_content: Dict[str, List[Dict]] = {}
    for k, v in heading_content.items():
        if isinstance(k, str) and k.startswith(scoped_prefix):
            parts = k.split("|||", 2)
            if len(parts) == 3:
                parent_title = parts[1]
                child_title = parts[2]
                scoped_heading_content[(normalize_title(parent_title), normalize_title(child_title))] = (k, v)
                continue
        flat_heading_content[k] = v

    heading_content_normalized = {normalize_title(k): (k, v) for k, v in flat_heading_content.items()}
    print(f"[DEBUG MATCHING] flat_heading_content keys: {list(flat_heading_content.keys())}")
    print(f"[DEBUG MATCHING] heading_content_normalized keys: {list(heading_content_normalized.keys())}")
    heading_context_by_level: Dict[int, str] = {}

    for idx, block in enumerate(block_items):
        if not isinstance(block, Paragraph) or not is_heading(block):
            continue

        para = block
        heading_text = para.text.strip()
        heading_level = get_heading_level(para)
        heading_context_by_level[heading_level] = heading_text
        for lvl in list(heading_context_by_level.keys()):
            if lvl > heading_level:
                del heading_context_by_level[lvl]
        parent_heading_text = heading_context_by_level.get(heading_level - 1, "")

        # Essayer le matching exact d'abord
        blocks = flat_heading_content.get(heading_text, [])

        # Si pas de matching exact, essayer le matching normalisé
        if not blocks:
            normalized = normalize_title(heading_text)
            if normalized in heading_content_normalized:
                original_key, blocks = heading_content_normalized[normalized]
                print(f"[DEBUG] Matching normalisé: '{heading_text}' -> '{original_key}'")

        # Si toujours pas de match, essayer le matching partiel (titre sans numérotation)
        # Ex: "2.6.1 Images" dans le doc -> "Images" dans heading_content
        if not blocks:
            # Enlever la numérotation du début (ex: "2.6.1 " -> "")
            import re
            heading_without_number = re.sub(r'^[\d\.]+\s*', '', heading_text).strip()
            if heading_without_number:
                normalized_without_number = normalize_title(heading_without_number)
                if normalized_without_number in heading_content_normalized:
                    original_key, blocks = heading_content_normalized[normalized_without_number]
                    print(f"[DEBUG] Matching partiel (sans numéro): '{heading_text}' -> '{original_key}'")

        # Matching contextuel parent->enfant pour éviter les collisions entre sous-sous-titres homonymes
        if not blocks and parent_heading_text:
            scoped_key = (normalize_title(parent_heading_text), normalize_title(heading_text))
            if scoped_key in scoped_heading_content:
                original_key, blocks = scoped_heading_content[scoped_key]
                print(f"[DEBUG] Matching contextuel: parent='{parent_heading_text}', child='{heading_text}' -> '{original_key}'")

            # Matching alternatif pour "Compte utilisateur" -> "Accounts" (transformation Comptes associés)
            if not blocks and normalize_title(heading_text) == "compte utilisateur":
                alt_scoped_key = (normalize_title(parent_heading_text), "accounts")
                if alt_scoped_key in scoped_heading_content:
                    original_key, blocks = scoped_heading_content[alt_scoped_key]
                    print(f"[DEBUG] Matching alternatif Compte utilisateur->Accounts: '{original_key}'")

        if blocks:
            print(f"[DEBUG] Heading '{heading_text}' trouvé, {len(blocks)} blocs à insérer")
        if not blocks:
            continue

        # Trouver le paragraphe directement après le heading (en utilisant la structure XML)
        target_para = para
        next_sibling = get_next_sibling_paragraph(para)
        if next_sibling is not None and not is_heading(next_sibling):
            target_para = next_sibling
            print(f"[DEBUG] Insertion après le paragraphe suivant: '{target_para.text[:50] if target_para.text else '(vide)'}...'")
        else:
            print(f"[DEBUG] Insertion directement après le heading: '{heading_text}'")

        # Inserer chaque bloc dans l'ordre (TOUJOURS, pas seulement dans le else)
        current_para = target_para

        # Ajouter un espacement après le titre (avant le premier bloc)
        try:
            current_para.paragraph_format.space_after = Pt(14)
        except:
            pass

        for block_idx, block in enumerate(blocks):
            print(f"[DEBUG] Bloc {block_idx}: type={block.get('type')}, content={str(block.get('content', ''))[:50] if block.get('content') else 'N/A'}, src={block.get('src', 'N/A')}")
            if block.get("type") == "text":
                content = block.get("content", "")
                if content:
                    # Remplacer les placeholders dans le contenu
                    content = replace_placeholders(content)
                    print(f"[DEBUG] Insertion texte: '{content[:80]}...'")

                    # Vérifier s'il y a une image de compte associée
                    account_image_src = block.get("accountImage")
                    account_image_width = block.get("accountImageWidth", 2.0)  # Défaut 2 pouces (~5cm)

                    # Résoudre le chemin de l'image de compte (si présente et trouvée)
                    acc_img_path = None
                    if account_image_src:
                        _acc_src = account_image_src.replace("/uploads/", "uploads/") if account_image_src.startswith("/uploads/") else account_image_src
                        _acc_p = Path(_acc_src)
                        if _acc_p.exists():
                            acc_img_path = _acc_p
                        else:
                            print(f"[DEBUG] Image compte non trouvée: {_acc_p}")

                    # Si image présente : tableau 2 colonnes (texte | image) pour les afficher CÔTE À CÔTE (comme le frontend)
                    # En cas d'échec, on retombe sur une insertion de l'image APRÈS le texte (repli garanti).
                    acc_table = None
                    acc_left_cell = None
                    if acc_img_path is not None:
                        try:
                            from docx.oxml.ns import qn as _qn
                            from docx.shared import Inches as _Inches
                            _acc_table = doc.add_table(rows=1, cols=2)
                            _acc_table.autofit = False
                            current_para._p.addnext(_acc_table._tbl)
                            _left = _acc_table.cell(0, 0)
                            _right = _acc_table.cell(0, 1)
                            try:
                                _left.width = Cm(11); _right.width = Cm(5.5)
                            except Exception:
                                pass
                            # Retirer toutes les bordures du tableau
                            _tblpr = _acc_table._tbl.tblPr
                            if _tblpr is None:
                                _tblpr = OxmlElement('w:tblPr'); _acc_table._tbl.insert(0, _tblpr)
                            _borders = OxmlElement('w:tblBorders')
                            for _b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                _el = OxmlElement(f'w:{_b}'); _el.set(_qn('w:val'), 'none'); _borders.append(_el)
                            _tblpr.append(_borders)
                            # Image dans la cellule de droite (centrée) — si ça échoue, on annule le tableau
                            _img_para = _right.paragraphs[0]
                            _img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            _img_para.add_run().add_picture(str(acc_img_path), width=_Inches(account_image_width))
                            # Tout est OK : on garde le tableau et on rend le texte dans la cellule gauche
                            acc_table = _acc_table
                            acc_left_cell = _left
                            current_para = acc_left_cell.paragraphs[0]
                            print(f"[DEBUG] Image compte insérée À CÔTÉ: {acc_img_path}, largeur={account_image_width} pouces")
                        except Exception as _e:
                            print(f"[DEBUG] Tableau compte échoué -> repli image après texte: {_e}")
                            acc_table = None
                            acc_left_cell = None

                    # Gérer les sauts de ligne - chaque ligne devient un paragraphe séparé
                    lines = content.split("\n")
                    for line in lines:
                        # Détecter le type de ligne pour appliquer le bon style
                        import re
                        is_compte_line = re.match(r'^Compte\s+.+\s*:', line)
                        is_bullet_detail = re.match(r'^[•\-\*]\s*(Pseudo|Nom d.utilisateur|Identifiant|Email|T.l.phone|Date et heure|Date de naissance)', line)
                        is_detail_line = re.match(r'^(Pseudo|Nom d.utilisateur|Identifiant|Email|T.l.phone|Date et heure|Date de naissance)\s*:', line)

                        if is_compte_line:
                            new_para = insert_after_in_cell(current_para, "", "Normal")
                            run = new_para.add_run(line)
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(30, 58, 95)
                            new_para.paragraph_format.space_before = Pt(14)
                            new_para.paragraph_format.space_after = Pt(6)
                        elif is_bullet_detail:
                            try:
                                new_para = insert_after_in_cell(current_para, "", "List Paragraph")
                            except:
                                new_para = insert_after_in_cell(current_para, "", "Normal")
                            clean_line = re.sub(r'^[•\-\*]\s*', '', line)
                            if " : " in clean_line:
                                label, value = clean_line.split(" : ", 1)
                                run_label = new_para.add_run(label + " : ")
                                run_label.bold = True
                                run_label.font.size = Pt(10)
                                run_label.font.color.rgb = RGBColor(80, 80, 80)
                                run_value = new_para.add_run(value)
                                run_value.font.size = Pt(10)
                                if value.strip() == "N/C":
                                    run_value.font.color.rgb = RGBColor(160, 160, 160)
                                    run_value.italic = True
                                else:
                                    run_value.font.color.rgb = RGBColor(30, 30, 30)
                            else:
                                run = new_para.add_run(clean_line)
                                run.bold = True
                                run.font.size = Pt(10)
                            new_para.paragraph_format.space_after = Pt(2)
                            new_para.paragraph_format.space_before = Pt(1)
                        elif is_detail_line:
                            try:
                                new_para = insert_after_in_cell(current_para, "", "List Paragraph")
                            except:
                                new_para = insert_after_in_cell(current_para, "", "Normal")
                            if " : " in line:
                                label, value = line.split(" : ", 1)
                                run_label = new_para.add_run(label + " : ")
                                run_label.bold = True
                                run_label.font.size = Pt(10)
                                run_label.font.color.rgb = RGBColor(80, 80, 80)
                                run_value = new_para.add_run(value)
                                run_value.font.size = Pt(10)
                                if value.strip() == "N/C":
                                    run_value.font.color.rgb = RGBColor(160, 160, 160)
                                    run_value.italic = True
                                else:
                                    run_value.font.color.rgb = RGBColor(30, 30, 30)
                            else:
                                run = new_para.add_run(line)
                                run.bold = True
                                run.font.size = Pt(10)
                            new_para.paragraph_format.space_after = Pt(2)
                            new_para.paragraph_format.space_before = Pt(1)
                        else:
                            new_para = insert_after_in_cell(current_para, line)

                        current_para = new_para
                    # Espacement après le bloc texte
                    try:
                        current_para.paragraph_format.space_after = Pt(12)
                    except:
                        pass

                    # Si on a utilisé le tableau côte à côte : retirer le paragraphe vide de tête
                    # de la cellule gauche et repositionner le curseur APRÈS le tableau
                    if acc_table is not None:
                        try:
                            _first_p = acc_left_cell.paragraphs[0]
                            if (not _first_p.text or not _first_p.text.strip()) and len(acc_left_cell.paragraphs) > 1:
                                _first_p._p.getparent().remove(_first_p._p)
                        except Exception:
                            pass
                        _after_acc = OxmlElement("w:p")
                        acc_table._tbl.addnext(_after_acc)
                        current_para = Paragraph(_after_acc, doc)
                        current_para.paragraph_format.space_before = Pt(8)
                    elif acc_img_path is not None:
                        # Repli garanti : insérer l'image APRÈS le texte (alignée à droite)
                        try:
                            from docx.shared import Inches as _Inches2
                            _np = insert_after_in_cell(current_para, "")
                            _np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            _np.add_run().add_picture(str(acc_img_path), width=_Inches2(account_image_width))
                            _np.paragraph_format.space_before = Pt(8)
                            _np.paragraph_format.space_after = Pt(12)
                            current_para = _np
                            print(f"[DEBUG] Image compte insérée APRÈS le texte (repli): {acc_img_path}")
                        except Exception as _e:
                            print(f"[DEBUG] Erreur repli image compte: {_e}")
                else:
                    print(f"[DEBUG] Bloc texte ignoré: content vide")
            elif block.get("type") == "image":
                src = block.get("src", "")
                if src:
                    width_raw = block.get("width")
                    print(f"[DEBUG BACKEND IMAGE] width_raw from frontend: {width_raw}, default_width_inches: {default_width_inches}")
                    try:
                        width_cm = float(width_raw) if width_raw is not None else None
                    except (TypeError, ValueError):
                        width_cm = None

                    # Le frontend TOUJOURS envoie les largeurs en cm (slider 3-20 cm)
                    # Il faut TOUJOURS convertir en pouces pour python-docx
                    if width_cm is not None:
                        width = width_cm / 2.54  # Conversion cm vers pouces
                        print(f"[DEBUG BACKEND IMAGE] width_cm={width_cm} cm -> width={width:.2f} inches")
                    else:
                        width = float(default_width_inches)
                        print(f"[DEBUG BACKEND IMAGE] Using default width: {width} inches")

                    caption = block.get("caption", "")
                    # Remplacer les placeholders dans la caption
                    caption = replace_placeholders(caption)
                    caption_pos = block.get("captionPosition", "bottom")

                    # Convertir les URLs /uploads/... en chemins de fichiers locaux
                    if src.startswith("/uploads/"):
                        src = src.replace("/uploads/", "uploads/")

                    img_path = Path(src)
                    is_chart_image = "chart_" in img_path.name.lower()

                    if not img_path.exists():
                        print(f"ATTENTION: Image introuvable: {src}")
                        continue

                    # Rendre les graphiques vraiment lisibles dans Word.
                    # Définir aussi la hauteur pour les images
                    if is_chart_image:
                        width = 17.0 / 2.54  # 17 cm fixe pour tous les graphiques
                        height = None  # Hauteur auto pour les graphiques (conserver ratio)
                    else:
                        # Adapter intelligemment selon les dimensions originales (ratio conservé)
                        # Limites en cm - MAX_HEIGHT est PRIORITAIRE (jamais dépassé)
                        MAX_WIDTH_CM = 4.5
                        MAX_HEIGHT_CM = 7.5  # Hauteur max stricte pour éviter images trop hautes

                        try:
                            from PIL import Image as PILImage
                            with PILImage.open(img_path) as pil_img:
                                orig_width_px, orig_height_px = pil_img.size

                            # Calculer le ratio original
                            ratio = orig_width_px / orig_height_px if orig_height_px > 0 else 1.0

                            # Commencer avec la largeur max
                            final_width_cm = MAX_WIDTH_CM
                            final_height_cm = final_width_cm / ratio

                            # Si la hauteur dépasse le max, réduire proportionnellement
                            # MAX_HEIGHT est PRIORITAIRE - ne jamais dépasser
                            if final_height_cm > MAX_HEIGHT_CM:
                                final_height_cm = MAX_HEIGHT_CM
                                final_width_cm = final_height_cm * ratio

                            width = final_width_cm / 2.54  # Conversion en pouces
                            height = final_height_cm / 2.54  # Conversion en pouces

                            print(f"[DEBUG IMAGE] Dimensions adaptées: {final_width_cm:.1f}cm x {final_height_cm:.1f}cm (ratio={ratio:.2f}, original={orig_width_px}x{orig_height_px}px)")
                        except Exception as dim_err:
                            # Fallback si impossible de lire les dimensions
                            print(f"[DEBUG IMAGE] Erreur lecture dimensions: {dim_err}, utilisation valeurs par défaut")
                            width = 6.0 / 2.54  # 6 cm par défaut
                            height = None  # Hauteur auto

                    print(f"[DEBUG IMAGE] Traitement image: {img_path}, suffix: {img_path.suffix.lower()}")

                    # Convertir les fichiers .decrypted, .file ou formats non standards en JPEG
                    img_suffix = img_path.suffix.lower()
                    needs_conversion = img_suffix in ('.decrypted', '.file', '.heic', '.heif', '.webp') or (img_suffix and img_suffix not in ('.jpg', '.jpeg', '.png', '.gif', '.bmp'))

                    if needs_conversion:
                        try:
                            print(f"[DEBUG] Conversion image format non standard: {img_path}")
                            from PIL import Image as PILImage
                            pil_img = PILImage.open(img_path)
                            if pil_img.mode in ('RGBA', 'LA', 'P'):
                                rgb_image = PILImage.new('RGB', pil_img.size, (255, 255, 255))
                                if pil_img.mode == 'P':
                                    pil_img = pil_img.convert('RGBA')
                                if pil_img.mode == 'RGBA':
                                    rgb_image.paste(pil_img, mask=pil_img.split()[-1])
                                else:
                                    rgb_image.paste(pil_img)
                                pil_img = rgb_image
                            elif pil_img.mode != 'RGB':
                                pil_img = pil_img.convert('RGB')
                            # Sauvegarder en JPEG temporaire
                            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                                tmp_converted_path = Path(tmp.name)
                                pil_img.save(tmp_converted_path, 'JPEG', quality=90, optimize=True)
                            img_path = tmp_converted_path
                            print(f"[DEBUG] Image convertie: {tmp_converted_path}")
                        except Exception as conv_err:
                            print(f"[DEBUG] Erreur conversion image: {conv_err}")
                            # Continuer avec le fichier original si la conversion échoue

                    # Fonction helper pour insérer l'image
                    # Import local pour éviter les problèmes de closure
                    from docx.shared import Inches as InchesLocal

                    def insert_image_to_run(run, img_path, width, height=None):
                        try:
                            if height is not None:
                                run.add_picture(str(img_path), width=InchesLocal(width), height=InchesLocal(height))
                            else:
                                run.add_picture(str(img_path), width=InchesLocal(width))
                            return True
                        except Exception as e:
                            print(f"[FALLBACK] Erreur insertion {img_path} ({type(e).__name__}), reconversion...")
                            try:
                                img = Image.open(img_path)
                                if img.mode in ('RGBA', 'LA', 'P'):
                                    rgb_image = Image.new('RGB', img.size, (255, 255, 255))
                                    if img.mode == 'P':
                                        img = img.convert('RGBA')
                                    if img.mode == 'RGBA':
                                        rgb_image.paste(img, mask=img.split()[-1])
                                    else:
                                        rgb_image.paste(img)
                                    img = rgb_image
                                elif img.mode != 'RGB':
                                    img = img.convert('RGB')
                                with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                                    tmp_path_fallback = Path(tmp.name)
                                    img.save(tmp_path_fallback, 'JPEG', quality=90, optimize=True, subsampling=0)
                                if height is not None:
                                    run.add_picture(str(tmp_path_fallback), width=InchesLocal(width), height=InchesLocal(height))
                                else:
                                    run.add_picture(str(tmp_path_fallback), width=InchesLocal(width))
                                try:
                                    tmp_path_fallback.unlink()
                                except:
                                    pass
                                return True
                            except Exception as e2:
                                print(f"ERREUR FALLBACK: {type(e2).__name__}: {str(e2)}")
                                return False

                    # Insertion selon la position du caption
                    if caption_pos in ("left", "right") and caption:
                        # Ajouter espacement avant l'image
                        spacer_para = insert_after_in_cell(current_para, "")
                        spacer_para.paragraph_format.space_after = Pt(12)
                        current_para = spacer_para

                        # Créer un tableau 1 ligne x 2 colonnes pour positionner texte et image côte à côte
                        # Calculer les largeurs: image en cm, texte prend le reste
                        img_width_cm = width * 2.54  # convertir inches en cm
                        total_width_cm = 16  # largeur totale approximative de la page
                        text_width_cm = max(4, total_width_cm - img_width_cm - 1)  # minimum 4cm pour le texte

                        # Créer le tableau dans le document (sera ajouté à la fin)
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.autofit = False

                        # Déplacer le tableau après le paragraphe courant
                        current_para._p.addnext(tbl._tbl)

                        # Configurer les cellules selon la position
                        if caption_pos == "right":
                            # Image à gauche, texte à droite
                            img_cell = tbl.cell(0, 0)
                            text_cell = tbl.cell(0, 1)
                            try:
                                img_cell.width = Cm(img_width_cm + 0.5)
                                text_cell.width = Cm(text_width_cm)
                            except:
                                pass
                        else:
                            # Texte à gauche, image à droite (caption_pos == "left")
                            text_cell = tbl.cell(0, 0)
                            img_cell = tbl.cell(0, 1)
                            try:
                                text_cell.width = Cm(text_width_cm)
                                img_cell.width = Cm(img_width_cm + 0.5)
                            except:
                                pass

                        # Ajouter le texte dans la cellule texte avec formatage
                        text_para = text_cell.paragraphs[0]
                        text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        # Parser le caption pour les métadonnées (format: **label:** valeur)
                        # Gérer les sauts de ligne
                        caption_lines = caption.split('\n')
                        first_line = True
                        for line in caption_lines:
                            if not first_line:
                                # Ajouter un nouveau paragraphe pour chaque ligne
                                text_para = text_cell.add_paragraph()
                                text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            # Parser les **bold** dans la ligne
                            import re
                            bold_pattern = re.compile(r'\*\*([^*]+)\*\*')
                            parts = bold_pattern.split(line)

                            for i, part in enumerate(parts):
                                if not part:
                                    continue
                                run = text_para.add_run(part)
                                run.font.size = Pt(10)
                                run.font.name = "Calibri"
                                # Les parties impaires sont celles qui étaient entre **
                                if i % 2 == 1:
                                    run.bold = True

                            first_line = False

                        # Ajouter l'image dans la cellule image
                        img_para = img_cell.paragraphs[0]
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        insert_image_to_run(run, img_path, width, height)

                        # Centrage vertical du texte dans la cellule
                        try:
                            from docx.oxml.ns import qn
                            tc_pr = text_cell._tc.get_or_add_tcPr()
                            v_align = OxmlElement('w:vAlign')
                            v_align.set(qn('w:val'), 'center')
                            tc_pr.append(v_align)
                        except:
                            pass

                        # Supprimer les bordures du tableau (tableau invisible)
                        tbl_element = tbl._tbl
                        tbl_pr = tbl_element.tblPr
                        if tbl_pr is None:
                            tbl_pr = OxmlElement('w:tblPr')
                            tbl_element.insert(0, tbl_pr)
                        tbl_borders = OxmlElement('w:tblBorders')
                        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')
                            tbl_borders.append(border)
                        tbl_pr.append(tbl_borders)

                        # Créer un paragraphe après le tableau pour continuer avec espacement
                        after_tbl_para = OxmlElement("w:p")
                        tbl._tbl.addnext(after_tbl_para)
                        current_para = Paragraph(after_tbl_para, doc)
                        try:
                            current_para.paragraph_format.space_before = Pt(12)
                            current_para.paragraph_format.space_after = Pt(18)
                        except:
                            pass

                    elif caption_pos in ("left", "right") and not caption:
                        # Pas de caption, juste l'image
                        # Ajouter un paragraphe vide avant pour l'espacement
                        spacer_para = insert_after_in_cell(current_para, "")
                        spacer_para.paragraph_format.space_after = Pt(12)
                        current_para = spacer_para

                        img_para = insert_after_in_cell(current_para, "")
                        run = img_para.add_run()
                        insert_image_to_run(run, img_path, width, height)
                        current_para = img_para
                    else:
                        # Position top ou bottom (par défaut)
                        # Ajouter un espacement avant l'image
                        spacer_para = insert_after_in_cell(current_para, "")
                        spacer_para.paragraph_format.space_after = Pt(12)
                        current_para = spacer_para

                        if caption and caption_pos == "top":
                            caption_para = insert_after_in_cell(current_para, caption)
                            current_para = caption_para

                        img_para = insert_after_in_cell(current_para, "")
                        run = img_para.add_run()
                        insert_image_to_run(run, img_path, width, height)
                        current_para = img_para

                        if caption and caption_pos == "bottom":
                            caption_para = insert_after_in_cell(current_para, caption)
                            current_para = caption_para

                    # Ajouter espacement après l'image
                    try:
                        current_para.paragraph_format.space_before = Pt(12)
                        current_para.paragraph_format.space_after = Pt(18)
                    except:
                        pass

            elif block.get("type") == "video":
                # Insertion d'une vidéo avec miniatures et métadonnées
                video_name = block.get("name", "")
                thumbnails = block.get("thumbnails", [])
                metadata = block.get("metadata", {})

                file_path = metadata.get("filePath", "")
                created_date = metadata.get("createdDate", "")
                observation = metadata.get("observation", "")

                # Nettoyer le chemin: enlever le premier mot avant le premier / mais GARDER le /
                if file_path and "/" in file_path:
                    slash_idx = file_path.index("/")
                    if slash_idx < len(file_path) - 1:
                        file_path = file_path[slash_idx:]  # Garder le / au début

                # Nettoyer la date: enlever le numéro de série Excel et les formats datetime en double
                if created_date:
                    import re as re_local
                    # Supprimer le numéro de série Excel en début (ex: 45849.18085648148)
                    created_date = re_local.sub(r'^\d+(\.\d+)?\s+', '', created_date)
                    # Supprimer le format ISO datetime si présent en début (2024-01-15T10:30:00)
                    created_date = re_local.sub(r'^\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}(\.\d+)?\s*', '', created_date)
                    # Supprimer aussi le format datetime sans T
                    created_date = re_local.sub(r'^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}\s*', '', created_date)

                print(f"[DEBUG] Insertion vidéo '{video_name}' avec {len(thumbnails)} miniatures")

                if thumbnails:
                    # Ajouter espacement avant
                    spacer_para = insert_after_in_cell(current_para, "")
                    spacer_para.paragraph_format.space_after = Pt(8)
                    current_para = spacer_para

                    # Ajouter les métadonnées centrées AU-DESSUS des miniatures
                    meta_paragraph = insert_after_in_cell(current_para, "")
                    meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    meta_paragraph.paragraph_format.space_after = Pt(6)

                    # Nom de la vidéo en gras
                    name_run = meta_paragraph.add_run(f"Vidéo : {video_name}")
                    name_run.bold = True
                    name_run.font.size = Pt(11)
                    name_run.font.name = "Calibri"
                    meta_paragraph.add_run("\n")

                    # Chemin
                    path_label = meta_paragraph.add_run("Chemin : ")
                    path_label.bold = True
                    path_label.font.size = Pt(10)
                    path_value = meta_paragraph.add_run(file_path or "Non disponible")
                    path_value.font.size = Pt(10)
                    meta_paragraph.add_run("  |  ")

                    # Date de création
                    date_label = meta_paragraph.add_run("Date : ")
                    date_label.bold = True
                    date_label.font.size = Pt(10)
                    date_value = meta_paragraph.add_run(created_date or "Non disponible")
                    date_value.font.size = Pt(10)
                    meta_paragraph.add_run("\n")

                    # Observations
                    obs_label = meta_paragraph.add_run("Observations : ")
                    obs_label.bold = True
                    obs_label.font.size = Pt(10)
                    obs_value = meta_paragraph.add_run(observation or "(à compléter)")
                    obs_value.font.size = Pt(10)

                    current_para = meta_paragraph

                    # Créer un tableau 4x2 pour les miniatures EN-DESSOUS des métadonnées
                    thumb_tbl = doc.add_table(rows=2, cols=4)
                    thumb_tbl.autofit = False
                    thumb_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    current_para._p.addnext(thumb_tbl._tbl)

                    # Définir la largeur des colonnes (plus compact)
                    col_width = Cm(2.8)  # Largeur réduite pour 4 colonnes
                    for col_idx in range(4):
                        for row in thumb_tbl.rows:
                            row.cells[col_idx].width = col_width

                    # Centrer les miniatures et réduire les marges des cellules
                    for row in thumb_tbl.rows:
                        row.height = Cm(1.6)  # Hauteur réduite
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            try:
                                from docx.oxml.ns import qn
                                tc_pr_mini = cell._tc.get_or_add_tcPr()
                                v_align_mini = OxmlElement('w:vAlign')
                                v_align_mini.set(qn('w:val'), 'center')
                                tc_pr_mini.append(v_align_mini)
                                tc_mar = OxmlElement('w:tcMar')
                                for margin_name in ['top', 'left', 'bottom', 'right']:
                                    margin = OxmlElement(f'w:{margin_name}')
                                    margin.set(qn('w:w'), '20')
                                    margin.set(qn('w:type'), 'dxa')
                                    tc_mar.append(margin)
                                tc_pr_mini.append(tc_mar)
                            except:
                                pass

                    # Insérer les 8 miniatures
                    from docx.shared import Inches as InchesLocal
                    thumb_positions = [(0, 0), (0, 1), (0, 2), (0, 3), (1, 0), (1, 1), (1, 2), (1, 3)]

                    for i, thumb_url in enumerate(thumbnails[:8]):
                        if i >= len(thumb_positions):
                            break
                        row_idx, col_idx = thumb_positions[i]
                        cell = thumb_tbl.cell(row_idx, col_idx)

                        thumb_src = thumb_url
                        if thumb_src.startswith("/uploads/"):
                            thumb_src = thumb_src.replace("/uploads/", "uploads/")

                        thumb_path = Path(thumb_src)
                        print(f"[DEBUG VIDEO] Miniature {i}: URL={thumb_url}, Path={thumb_path}, Exists={thumb_path.exists()}")

                        if thumb_path.exists():
                            try:
                                para = cell.paragraphs[0]
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = para.add_run()
                                run.add_picture(str(thumb_path), width=InchesLocal(1.0))
                                print(f"[DEBUG VIDEO] Miniature {i} insérée avec succès")
                            except Exception as e:
                                print(f"[DEBUG VIDEO] Erreur miniature {i}: {e}")
                                para = cell.paragraphs[0]
                                para.add_run(f"[Mini {i+1}]")
                        else:
                            print(f"[DEBUG VIDEO] Miniature {i} non trouvée: {thumb_path}")
                            para = cell.paragraphs[0]
                            para.add_run(f"[Mini {i+1} - non trouvé]")

                    # Supprimer les bordures du tableau miniatures
                    try:
                        thumb_tbl_element = thumb_tbl._tbl
                        thumb_tbl_pr = thumb_tbl_element.tblPr
                        if thumb_tbl_pr is None:
                            thumb_tbl_pr = OxmlElement('w:tblPr')
                            thumb_tbl_element.insert(0, thumb_tbl_pr)
                        thumb_borders = OxmlElement('w:tblBorders')
                        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')
                            thumb_borders.append(border)
                        thumb_tbl_pr.append(thumb_borders)
                    except:
                        pass

                    # Paragraphe après le tableau des miniatures
                    after_tbl_para = OxmlElement("w:p")
                    thumb_tbl._tbl.addnext(after_tbl_para)
                    current_para = Paragraph(after_tbl_para, doc)
                    try:
                        current_para.paragraph_format.space_before = Pt(12)
                        current_para.paragraph_format.space_after = Pt(18)
                    except:
                        pass

            elif block.get("type") == "table":
                # Insertion d'un tableau Excel
                columns = block.get("columns", [])
                display_columns = block.get("displayColumns", columns)  # Noms affichés (éditables par l'utilisateur)
                data = block.get("data", [])
                sheet_name = block.get("sheet", "")

                if not columns or not data:
                    continue

                print(f"[DEBUG] Insertion tableau '{sheet_name}': {len(columns)} colonnes, {len(data)} lignes")
                print(f"[DEBUG] Colonnes originales: {columns}")
                print(f"[DEBUG] Colonnes affichées: {display_columns}")

                # Créer le tableau Word
                tbl = doc.add_table(rows=len(data) + 1, cols=len(columns))
                tbl.style = 'Table Grid'
                tbl.autofit = False  # Désactiver autofit pour contrôler les largeurs

                # Centrer le tableau si peu de colonnes (1-3 colonnes)
                if len(columns) <= 3:
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Déplacer le tableau après le paragraphe courant
                current_para._p.addnext(tbl._tbl)

                from docx.oxml.ns import nsdecls, qn
                from docx.oxml import parse_xml

                # Configurer le tableau pour utiliser la largeur de page
                try:
                    tbl_element = tbl._tbl
                    tbl_pr = tbl_element.tblPr
                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        tbl_element.insert(0, tbl_pr)
                    # Largeur du tableau = 100% de la page disponible
                    tbl_width = OxmlElement('w:tblW')
                    tbl_width.set(qn('w:w'), '5000')
                    tbl_width.set(qn('w:type'), 'pct')  # Pourcentage
                    tbl_pr.append(tbl_width)
                    # Layout auto pour permettre le retour à la ligne
                    tbl_layout = OxmlElement('w:tblLayout')
                    tbl_layout.set(qn('w:type'), 'autofit')
                    tbl_pr.append(tbl_layout)
                except:
                    pass

                # Calculer les largeurs optimales par colonne (basé sur le contenu)
                # Page A4 avec marges: environ 16cm de largeur disponible
                MAX_TABLE_WIDTH_CM = 16.0

                col_widths = []
                for col_idx, col_name in enumerate(columns):
                    display_name = display_columns[col_idx] if col_idx < len(display_columns) else col_name
                    max_len = len(str(display_name))
                    for row_data in data[:30]:  # Échantillon pour perf
                        val = str(row_data.get(col_name, "") or "")
                        if len(val) > max_len:
                            max_len = len(val)
                    # Largeur en cm: min 1.5cm, max 6cm, proportionnel au contenu
                    width_cm = min(6, max(1.5, max_len * 0.18))
                    col_widths.append(width_cm)

                # Si la largeur totale dépasse la page, réduire proportionnellement
                total_width = sum(col_widths)
                if total_width > MAX_TABLE_WIDTH_CM:
                    scale_factor = MAX_TABLE_WIDTH_CM / total_width
                    col_widths = [w * scale_factor for w in col_widths]
                    # S'assurer que chaque colonne a au moins 1cm
                    col_widths = [max(1.0, w) for w in col_widths]

                # En-tête du tableau - Design moderne bleu foncé
                header_row = tbl.rows[0]
                header_row.height = Pt(28)

                # Répéter l'en-tête sur chaque page (pour les tableaux longs)
                try:
                    tr = header_row._tr
                    trPr = tr.get_or_add_trPr()
                    tblHeader = OxmlElement('w:tblHeader')
                    trPr.append(tblHeader)
                except:
                    pass

                for col_idx, col_name in enumerate(columns):
                    cell = header_row.cells[col_idx]
                    display_name = display_columns[col_idx] if col_idx < len(display_columns) else col_name
                    cell.text = str(display_name)

                    # Largeur de colonne adaptative
                    try:
                        cell.width = Cm(col_widths[col_idx])
                    except:
                        pass

                    # Style en-tête: gras, blanc sur fond bleu foncé
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(255, 255, 255)

                    # Fond bleu foncé moderne pour l'en-tête
                    try:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="1E3A5F"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        pass
                    

                    
                    # Padding cellule
                    try:
                        tc_pr = cell._tc.get_or_add_tcPr()
                        tc_mar = OxmlElement('w:tcMar')
                        for side in ['top', 'bottom', 'left', 'right']:
                            margin = OxmlElement(f'w:{side}')
                            margin.set(qn('w:w'), '100')
                            margin.set(qn('w:type'), 'dxa')
                            tc_mar.append(margin)
                        tc_pr.append(tc_mar)
                        # S'assurer que le texte reste horizontal
                        for td in tc_pr.findall(qn('w:textDirection')):
                            tc_pr.remove(td)
                    except:
                        pass

                # Données du tableau avec alternance de couleurs
                # Vérifier si c'est la feuille Emails pour nettoyer le HTML de Body
                is_emails_sheet = "email" in sheet_name.lower() if sheet_name else False

                for row_idx, row_data in enumerate(data):
                    row = tbl.rows[row_idx + 1]
                    row.height = Pt(22)

                    # Couleur alternée: blanc / gris très clair
                    bg_color = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"

                    for col_idx, col_name in enumerate(columns):
                        cell = row.cells[col_idx]
                        value = row_data.get(col_name, "")
                        # Nettoyer le HTML de la colonne Body pour les Emails
                        if is_emails_sheet and col_name == "Body" and value:
                            value = strip_html_tags(str(value))
                        cell.text = str(value) if value not in [None, ""] else "—"

                        # Largeur de colonne
                        try:
                            cell.width = Cm(col_widths[col_idx])
                        except:
                            pass

                        # Style du texte
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                if cell.text == "—":
                                    run.font.color.rgb = RGBColor(160, 160, 160)
                                else:
                                    run.font.color.rgb = RGBColor(30, 30, 30)

                        # Fond alterné
                        try:
                            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                        except:
                            pass

                        # Padding cellule et retour à la ligne
                        try:
                            tc_pr = cell._tc.get_or_add_tcPr()
                            tc_mar = OxmlElement('w:tcMar')
                            for side in ['top', 'bottom']:
                                margin = OxmlElement(f'w:{side}')
                                margin.set(qn('w:w'), '60')
                                margin.set(qn('w:type'), 'dxa')
                                tc_mar.append(margin)
                            for side in ['left', 'right']:
                                margin = OxmlElement(f'w:{side}')
                                margin.set(qn('w:w'), '100')
                                margin.set(qn('w:type'), 'dxa')
                                tc_mar.append(margin)
                            tc_pr.append(tc_mar)
                            # S'assurer que le texte reste horizontal (pas de rotation)
                            # Supprimer tout textDirection existant
                            for td in tc_pr.findall(qn('w:textDirection')):
                                tc_pr.remove(td)
                        except:
                            pass

                # Bordures fines et modernes pour tout le tableau
                try:
                    tbl_element = tbl._tbl
                    tbl_pr = tbl_element.tblPr
                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        tbl_element.insert(0, tbl_pr)
                    tbl_borders = OxmlElement('w:tblBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')  # Bordure fine
                        border.set(qn('w:color'), 'E2E8F0')  # Gris clair
                        tbl_borders.append(border)
                    # Supprimer ancienne bordure si existe
                    for old_border in tbl_pr.findall(qn('w:tblBorders')):
                        tbl_pr.remove(old_border)
                    tbl_pr.append(tbl_borders)
                except Exception as e:
                    print(f"[DEBUG] Erreur bordures tableau: {e}")

                # Créer un paragraphe après le tableau pour continuer
                after_tbl_para = OxmlElement("w:p")
                tbl._tbl.addnext(after_tbl_para)
                current_para = Paragraph(after_tbl_para, doc)
                # Ajouter espacement après le tableau
                try:
                    current_para.paragraph_format.space_before = Pt(12)
                except:
                    pass

            elif block.get("type") == "conversation":
                # Insertion d'une conversation de chat comme image(s) PNG
                contact_name = block.get("contactName", "")
                source = block.get("source", "")
                image_paths = block.get("imagePaths", [])  # Liste d'images (paginées)
                message_images = block.get("messageImages", [])  # Liste d'images individuelles avec commentaires
                period_start = block.get("periodStart", "")
                period_end = block.get("periodEnd", "")
                contact_info = block.get("contactInfo", {}) or {}

                if message_images:
                    # Mode messages individuels avec commentaires
                    print(f"[DEBUG] Insertion {len(message_images)} message(s) individuel(s) avec commentaires: {contact_name}")

                    # Insérer le tableau récapitulatif du contact avant la conversation
                    source_lower = (contact_info.get("source", "") or "").lower()
                    is_signal = source_lower == "signal"

                    # Convertir en booléen (peut venir comme string "True"/"False" ou bool)
                    is_social_val = contact_info.get("is_social_source", None)
                    is_phone_val = contact_info.get("is_phone_source", None)

                    if is_social_val is not None:
                        is_social_source = is_social_val if isinstance(is_social_val, bool) else str(is_social_val).lower() == "true"
                    else:
                        is_social_source = False

                    if is_phone_val is not None:
                        is_phone_source = is_phone_val if isinstance(is_phone_val, bool) else str(is_phone_val).lower() == "true"
                    else:
                        is_phone_source = False

                    # Récupérer les infos du contact selon le type de source
                    # Pour sources sociales: ne PAS utiliser contact_name comme fallback pour pseudonyme
                    # Le pseudonyme doit venir uniquement du lookup Contacts (User ID-Username)
                    if is_social_source or is_signal:
                        # Sources sociales: pseudonyme "-" si pas trouvé dans Contacts
                        pseudonyme = contact_info.get("pseudonyme", "") or "-"
                        nom_utilisateur = contact_info.get("nom_utilisateur", "") or ""
                        identifiant_utilisateur = contact_info.get("identifiant_utilisateur", "") or ""
                        numero_telephone = ""
                    else:
                        # Sources téléphoniques: pseudonyme peut être le nom du contact
                        pseudonyme = contact_info.get("pseudonyme", "") or contact_name
                        nom_utilisateur = ""
                        identifiant_utilisateur = ""
                        numero_telephone = contact_info.get("numero_telephone", "") or contact_name

                    # Debug
                    print(f"[CONV TABLE] contact_info: {contact_info}")
                    print(f"[CONV TABLE] period_start: '{period_start}', period_end: '{period_end}'")
                    print(f"[CONV TABLE] is_social_source: {is_social_source}, is_phone_source: {is_phone_source}, is_signal: {is_signal}")

                    # Période formatée
                    if period_start and period_end:
                        period_text = f"{period_start}\nAu\n{period_end}"
                    elif period_start:
                        period_text = period_start
                    else:
                        period_text = "N/A"

                    # Override période si l'utilisateur a édité la cellule dans le frontend
                    _period_override = contact_info.get("period_text")
                    if _period_override is not None and str(_period_override).strip():
                        period_text = str(_period_override)

                    # Déterminer les colonnes selon le type de source
                    if is_signal:
                        # Signal: 3 colonnes (pas de Pseudonyme)
                        headers = ["Nom utilisateur", "Identifiant utilisateur", "Période de communication"]
                        data_values = [nom_utilisateur or contact_name, identifiant_utilisateur, period_text]
                        num_cols = 3
                    elif is_social_source:
                        # Sources sociales (Snapchat, Instagram, TikTok, Telegram, etc.): 4 colonnes
                        headers = ["Pseudonyme", "Nom utilisateur", "Identifiant utilisateur", "Période de communication"]
                        data_values = [pseudonyme, nom_utilisateur, identifiant_utilisateur, period_text]
                        num_cols = 4
                    else:
                        # Sources téléphoniques (Native messages, Natif, WhatsApp): 3 colonnes
                        headers = ["Pseudonyme", "Numéro de téléphone", "Période de communication"]
                        data_values = [pseudonyme, numero_telephone, period_text]
                        num_cols = 3

                    # Créer le tableau récapitulatif
                    summary_table = doc.add_table(rows=2, cols=num_cols)
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    current_para._p.addnext(summary_table._tbl)

                    # Définir les largeurs de colonnes (en cm)
                    if is_signal:
                        # 3 colonnes: Nom utilisateur (4.5cm), Identifiant (5.5cm), Période (5.5cm)
                        col_widths = [Cm(4.5), Cm(5.5), Cm(5.5)]
                    elif is_social_source:
                        # 4 colonnes: Pseudonyme (3cm), Nom utilisateur (3.5cm), Identifiant (4cm), Période (5cm)
                        col_widths = [Cm(3), Cm(3.5), Cm(4), Cm(5)]
                    else:
                        # 3 colonnes: Pseudonyme (4cm), Numéro (5cm), Période (5.5cm)
                        col_widths = [Cm(4), Cm(5), Cm(5.5)]

                    # Appliquer les largeurs
                    for i, width in enumerate(col_widths):
                        for row in summary_table.rows:
                            row.cells[i].width = width

                    # En-têtes du tableau
                    header_row = summary_table.rows[0]
                    for i, header_text in enumerate(headers):
                        cell = header_row.cells[i]
                        cell.text = header_text
                        # Style de l'en-tête
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        # Fond gris foncé pour l'en-tête
                        try:
                            from docx.oxml.ns import qn
                            tc_pr = cell._tc.get_or_add_tcPr()
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:val'), 'clear')
                            shading.set(qn('w:color'), 'auto')
                            shading.set(qn('w:fill'), '2d3748')
                            tc_pr.append(shading)
                        except:
                            pass

                    # Ligne de données
                    data_row = summary_table.rows[1]
                    for i, value in enumerate(data_values):
                        cell = data_row.cells[i]
                        cell.text = value
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(0, 0, 0)

                    # Style du tableau (bordures)
                    try:
                        from docx.oxml.ns import qn
                        tbl_element = summary_table._tbl
                        tbl_pr = tbl_element.tblPr
                        if tbl_pr is None:
                            tbl_pr = OxmlElement('w:tblPr')
                            tbl_element.insert(0, tbl_pr)
                        tbl_borders = OxmlElement('w:tblBorders')
                        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:color'), '888888')
                            tbl_borders.append(border)
                        tbl_pr.append(tbl_borders)
                    except Exception as e:
                        print(f"[DEBUG] Erreur style tableau récapitulatif: {e}")

                    # Paragraphe après le tableau récapitulatif
                    after_summary_para = OxmlElement("w:p")
                    summary_table._tbl.addnext(after_summary_para)
                    current_para = Paragraph(after_summary_para, doc)
                    current_para.paragraph_format.space_before = Pt(12)

                    print(f"[DEBUG] Tableau récapitulatif inséré: {pseudonyme}, période: {period_start} - {period_end}")

                    # Titre de la conversation (SOUS le tableau récapitulatif) — même style que le frontend
                    conv_title_para = insert_after_in_cell(current_para, "", "Normal")
                    conv_title_run = conv_title_para.add_run(f"Conversation avec {contact_name}")
                    conv_title_run.bold = True
                    conv_title_run.font.size = Pt(11)
                    conv_title_run.font.color.rgb = RGBColor(22, 32, 46)   # encre (#16202e)
                    if source:
                        conv_src_run = conv_title_para.add_run(f" ({source})")
                        conv_src_run.font.size = Pt(10)
                        conv_src_run.font.color.rgb = RGBColor(31, 95, 176)  # bleu institutionnel (#1f5fb0)
                    conv_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    conv_title_para.paragraph_format.space_before = Pt(10)
                    conv_title_para.paragraph_format.space_after = Pt(4)
                    current_para = conv_title_para

                    # Légende des messages / position des commentaires (gris #666)
                    # Cercles ● colorés en vrai vert/bleu (les emojis 🟢🔵 sortent en noir & blanc dans Word)
                    legend_para = insert_after_in_cell(current_para, "", "Normal")
                    grey = RGBColor(102, 102, 102)
                    _r = legend_para.add_run(f"\U0001F4AC {len(message_images)} message(s)  -  ")
                    _r.font.size = Pt(9); _r.font.color.rgb = grey
                    _rg = legend_para.add_run("● ")
                    _rg.font.size = Pt(9); _rg.font.color.rgb = RGBColor(31, 157, 87)   # vert (#1f9d57)
                    _r2 = legend_para.add_run("Vert : commentaire à gauche   |   ")
                    _r2.font.size = Pt(9); _r2.font.color.rgb = grey
                    _rb = legend_para.add_run("● ")
                    _rb.font.size = Pt(9); _rb.font.color.rgb = RGBColor(31, 95, 176)   # bleu (#1f5fb0)
                    _r3 = legend_para.add_run("Bleu : commentaire à droite")
                    _r3.font.size = Pt(9); _r3.font.color.rgb = grey
                    legend_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    legend_para.paragraph_format.space_after = Pt(8)
                    current_para = legend_para

                    from docx.shared import Inches
                    from docx.oxml.ns import nsdecls, qn
                    from docx.oxml import parse_xml

                    # Compteur de progression pour logs optimisés
                    total_messages = len(message_images)
                    log_interval = max(1, total_messages // 10)  # Log toutes les 10% environ
                    import time as time_module
                    batch_start_time = time_module.perf_counter()

                    for msg_idx, msg_data in enumerate(message_images):
                        # Pour le docx: preferer la bulle CLEAN (sans miniatures embarquees)
                        # car les miniatures video sont inserees separement APRES la bulle.
                        # Fallback sur image_path (full avec miniatures) si clean indisponible.
                        clean_path = msg_data.get("image_path_clean", "")
                        full_path = msg_data.get("image_path", "")
                        if clean_path and Path(clean_path).exists():
                            img_path = clean_path
                        else:
                            img_path = full_path
                        comment = msg_data.get("comment", "")
                        comment_position = msg_data.get("comment_position", "right")  # left ou right
                        is_owner = msg_data.get("is_owner", False)

                        if not img_path or not Path(img_path).exists():
                            print(f"[DEBUG] Image message non trouvée: {img_path}")
                            continue

                        if comment and comment.strip():
                            # Créer un tableau 1x2 pour image + commentaire
                            tbl = doc.add_table(rows=1, cols=2)
                            tbl.autofit = False
                            current_para._p.addnext(tbl._tbl)

                            # Configuration des colonnes selon la position du commentaire
                            if comment_position == "left":
                                # Commentaire à gauche, image à droite
                                comment_cell = tbl.cell(0, 0)
                                img_cell = tbl.cell(0, 1)
                            else:
                                # Image à gauche, commentaire à droite (par défaut)
                                img_cell = tbl.cell(0, 0)
                                comment_cell = tbl.cell(0, 1)

                            # Largeurs: image 8cm + commentaire 7cm = 15cm (sous la largeur utile A4 ~16cm).
                            # L'ancienne config (13+5=18cm) depassait la page Word, ce qui pouvait faire
                            # tronquer / disparaitre la cellule commentaire pour certains messages.
                            try:
                                img_cell.width = Cm(8)
                                comment_cell.width = Cm(7)
                            except Exception:
                                pass
                            # Forcer la table a NE PAS s'autofit (sinon Word recalcule selon le contenu image)
                            try:
                                from docx.oxml.ns import qn as _qn
                                tbl_pr = tbl._element.tblPr
                                tbl_layout = tbl_pr.find(_qn('w:tblLayout'))
                                if tbl_layout is None:
                                    tbl_layout = OxmlElement('w:tblLayout')
                                    tbl_pr.append(tbl_layout)
                                tbl_layout.set(_qn('w:type'), 'fixed')
                            except Exception:
                                pass
                            # Forcer la largeur preferee de chaque cellule via tcW (sinon Word ignore .width)
                            try:
                                from docx.oxml.ns import qn as _qn
                                for _c, _w_cm in ((img_cell, 8), (comment_cell, 7)):
                                    _tc_pr = _c._tc.get_or_add_tcPr()
                                    _tc_w = _tc_pr.find(_qn('w:tcW'))
                                    if _tc_w is None:
                                        _tc_w = OxmlElement('w:tcW')
                                        _tc_pr.append(_tc_w)
                                    # tcW en twentieths of a point. 1 cm = 567 twips.
                                    _tc_w.set(_qn('w:w'), str(int(_w_cm * 567)))
                                    _tc_w.set(_qn('w:type'), 'dxa')
                            except Exception as _wcell_err:
                                print(f"[DEBUG] Erreur tcW: {_wcell_err}")

                            # Retirer toutes les marges de cellules (defaut Word ~0.19cm)
                            # et le retrait de table, sinon la bulle se decale par rapport
                            # a une bulle sans commentaire (qui est en paragraphe direct,
                            # sans cellule).
                            try:
                                from docx.oxml.ns import qn as _qn
                                # Marges de chaque cellule a 0
                                for _c in (img_cell, comment_cell):
                                    _tc_pr = _c._tc.get_or_add_tcPr()
                                    _tc_mar = _tc_pr.find(_qn('w:tcMar'))
                                    if _tc_mar is None:
                                        _tc_mar = OxmlElement('w:tcMar')
                                        _tc_pr.append(_tc_mar)
                                    for _side in ('top', 'left', 'bottom', 'right'):
                                        _el = _tc_mar.find(_qn(f'w:{_side}'))
                                        if _el is None:
                                            _el = OxmlElement(f'w:{_side}')
                                            _tc_mar.append(_el)
                                        _el.set(_qn('w:w'), '0')
                                        _el.set(_qn('w:type'), 'dxa')
                                # Retrait de la table a 0
                                _tbl_ind = tbl_pr.find(_qn('w:tblInd'))
                                if _tbl_ind is None:
                                    _tbl_ind = OxmlElement('w:tblInd')
                                    tbl_pr.append(_tbl_ind)
                                _tbl_ind.set(_qn('w:w'), '0')
                                _tbl_ind.set(_qn('w:type'), 'dxa')
                                # Marges par defaut de la table elle-meme a 0
                                _tbl_cell_mar = tbl_pr.find(_qn('w:tblCellMar'))
                                if _tbl_cell_mar is None:
                                    _tbl_cell_mar = OxmlElement('w:tblCellMar')
                                    tbl_pr.append(_tbl_cell_mar)
                                for _side in ('top', 'left', 'bottom', 'right'):
                                    _el = _tbl_cell_mar.find(_qn(f'w:{_side}'))
                                    if _el is None:
                                        _el = OxmlElement(f'w:{_side}')
                                        _tbl_cell_mar.append(_el)
                                    _el.set(_qn('w:w'), '0')
                                    _el.set(_qn('w:type'), 'dxa')
                            except Exception as _mar_err:
                                print(f"[DEBUG] Erreur marges cellule/table: {_mar_err}")

                            # Aligner la TABLE de meme cote que la bulle aurait ete sans commentaire :
                            #  - owner (vert)   -> table a DROITE (comment a gauche, bulle a droite)
                            #  - non-owner      -> table a GAUCHE (bulle a gauche, comment a droite)
                            # Sans ca, la table s'aligne par defaut a gauche et la bulle owner
                            # "se rapproche" du commentaire au lieu de rester collee au bord droit.
                            try:
                                if comment_position == "left":
                                    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
                                else:
                                    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                            except Exception:
                                pass

                            # Insérer l'image dans la cellule image (avec compression).
                            # Aligner la bulle DANS sa cellule du meme cote que sans commentaire :
                            # owner -> bulle a droite de sa cellule, non-owner -> bulle a gauche.
                            img_para = img_cell.paragraphs[0]
                            if comment_position == "left":
                                img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            try:
                                run = img_para.add_run()
                                # Compresser l'image pour réduire la taille du fichier Word
                                compressed = compress_image_for_docx(img_path)
                                if compressed:
                                    run.add_picture(compressed, width=Cm(_compute_bubble_cm(img_path)))
                                else:
                                    run.add_picture(img_path, width=Cm(_compute_bubble_cm(img_path)))
                            except Exception as img_err:
                                print(f"[DEBUG] Erreur insertion image message: {img_err}")
                                error_run = img_para.add_run("[Image non disponible]")
                                error_run.font.color.rgb = RGBColor(150, 150, 150)
                                error_run.italic = True

                            # Insérer le commentaire dans la cellule commentaire
                            # On preserve les sauts de ligne (\n) via run.add_break() :
                            # add_run("a\nb") sur python-docx ne genere PAS de saut de ligne
                            # dans Word — il faut un <w:br/> entre les runs.
                            comment_para = comment_cell.paragraphs[0]
                            comment_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            comment_lines = comment.split("\n")
                            for line_idx, line in enumerate(comment_lines):
                                if line_idx > 0:
                                    br_run = comment_para.add_run()
                                    br_run.add_break()
                                line_run = comment_para.add_run(line)
                                line_run.font.name = "Calibri"
                                line_run.font.size = Pt(11)
                                line_run.font.color.rgb = RGBColor(0, 0, 0)

                            # Centrage vertical du commentaire dans la cellule
                            try:
                                tc_pr = comment_cell._tc.get_or_add_tcPr()
                                v_align = OxmlElement('w:vAlign')
                                v_align.set(qn('w:val'), 'center')
                                tc_pr.append(v_align)
                            except:
                                pass

                            # Supprimer les bordures du tableau (tableau invisible)
                            tbl_element = tbl._tbl
                            tbl_pr = tbl_element.tblPr
                            if tbl_pr is None:
                                tbl_pr = OxmlElement('w:tblPr')
                                tbl_element.insert(0, tbl_pr)
                            tbl_borders = OxmlElement('w:tblBorders')
                            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')
                                tbl_borders.append(border)
                            tbl_pr.append(tbl_borders)

                            # Créer un paragraphe après le tableau pour continuer avec espacement
                            after_tbl_para = OxmlElement("w:p")
                            tbl._tbl.addnext(after_tbl_para)
                            current_para = Paragraph(after_tbl_para, doc)
                            # Espacement après chaque message
                            current_para.paragraph_format.space_before = Pt(12)
                            current_para.paragraph_format.space_after = Pt(12)

                            # Log de progression optimisé (toutes les 10% ou dernier message)
                            if msg_idx == 0 or (msg_idx + 1) % log_interval == 0 or msg_idx == total_messages - 1:
                                elapsed = time_module.perf_counter() - batch_start_time
                                rate = (msg_idx + 1) / elapsed if elapsed > 0 else 0
                                print(f"[CONV] Progression: {msg_idx + 1}/{total_messages} ({100*(msg_idx+1)//total_messages}%) - {rate:.1f} msg/s")

                            # Insérer les images attachées en haute qualité après la bulle
                            # L'utilisateur peut les redimensionner dans Word
                            image_attachments = msg_data.get("image_attachments", [])
                            if image_attachments:
                                print(f"[WORD IMG] Message {msg_idx}: {len(image_attachments)} image(s) à insérer")
                                for img_att in image_attachments:
                                    att_path = img_att.get("path", "")
                                    att_name = img_att.get("name", "")
                                    path_exists = Path(att_path).exists() if att_path else False
                                    print(f"[WORD IMG] -> {att_name}: path={att_path}, exists={path_exists}")
                                    if att_path and path_exists:
                                        try:
                                            att_para = insert_after_in_cell(current_para, "", "Normal")
                                            att_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_owner else WD_ALIGN_PARAGRAPH.LEFT
                                            att_run = att_para.add_run()
                                            # Compresser et insérer l'image (largeur 6cm)
                                            compressed_att = compress_image_for_docx(att_path)
                                            if compressed_att:
                                                att_run.add_picture(compressed_att, width=Cm(6))
                                            else:
                                                att_run.add_picture(att_path, width=Cm(6))
                                            att_para.paragraph_format.space_before = Pt(4)
                                            att_para.paragraph_format.space_after = Pt(8)
                                            current_para = att_para
                                            print(f"[WORD IMG] ✓ Image insérée: {att_name}")
                                        except Exception as att_err:
                                            print(f"[WORD IMG] ✗ Erreur: {att_err}")

                            # Insérer les PDF attachés (convertis en images) après les images
                            # L'utilisateur peut les redimensionner dans Word
                            pdf_attachments = msg_data.get("pdf_attachments", [])
                            if pdf_attachments:
                                for pdf_att in pdf_attachments:
                                    pdf_img_path = pdf_att.get("path", "")
                                    pdf_name = pdf_att.get("name", "")
                                    if pdf_img_path and Path(pdf_img_path).exists():
                                        try:
                                            pdf_para = insert_after_in_cell(current_para, "", "Normal")
                                            pdf_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_owner else WD_ALIGN_PARAGRAPH.LEFT
                                            pdf_run = pdf_para.add_run()
                                            # Compresser et insérer l'image du PDF (largeur 12cm pour meilleure visibilité)
                                            compressed_pdf = compress_image_for_docx(pdf_img_path)
                                            if compressed_pdf:
                                                pdf_run.add_picture(compressed_pdf, width=Cm(12))
                                            else:
                                                pdf_run.add_picture(pdf_img_path, width=Cm(12))
                                            pdf_para.paragraph_format.space_before = Pt(4)
                                            pdf_para.paragraph_format.space_after = Pt(8)
                                            current_para = pdf_para
                                        except Exception as pdf_err:
                                            print(f"[DEBUG] Erreur insertion PDF: {pdf_err}")

                            # Insérer le composite des miniatures vidéo APRÈS la bulle dans le docx
                            # (la bulle utilisée ici est la version CLEAN sans miniatures).
                            video_attachments = msg_data.get("video_attachments", [])
                            if video_attachments:
                                # Dedupliquer si plusieurs videos pointent vers le meme composite
                                seen_composites = set()
                                for vid_att in video_attachments:
                                    vid_composite = vid_att.get("composite_path", "")
                                    if not vid_composite or vid_composite in seen_composites:
                                        continue
                                    if not Path(vid_composite).exists():
                                        continue
                                    seen_composites.add(vid_composite)
                                    try:
                                        vid_para = insert_after_in_cell(current_para, "", "Normal")
                                        vid_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_owner else WD_ALIGN_PARAGRAPH.LEFT
                                        vid_run = vid_para.add_run()
                                        compressed_vid = compress_image_for_docx(vid_composite)
                                        if compressed_vid:
                                            vid_run.add_picture(compressed_vid, width=Cm(_compute_bubble_cm(vid_composite)))
                                        else:
                                            vid_run.add_picture(vid_composite, width=Cm(_compute_bubble_cm(vid_composite)))
                                        vid_para.paragraph_format.space_before = Pt(4)
                                        vid_para.paragraph_format.space_after = Pt(8)
                                        current_para = vid_para
                                    except Exception as vid_err:
                                        print(f"[WORD VID] Erreur insertion miniatures: {vid_err}")

                        else:
                            # Pas de commentaire, juste l'image
                            # Positionner selon is_owner: vert (owner) à droite, bleu (autre) à gauche
                            new_para = insert_after_in_cell(current_para, "", "Normal")

                            # Alignement: owner (vert) à droite, autre (bleu) à gauche
                            if is_owner:
                                new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            try:
                                run = new_para.add_run()
                                # Compresser l'image pour réduire la taille du fichier Word
                                compressed = compress_image_for_docx(img_path)
                                if compressed:
                                    run.add_picture(compressed, width=Cm(_compute_bubble_cm(img_path)))
                                else:
                                    run.add_picture(img_path, width=Cm(_compute_bubble_cm(img_path)))
                                # Log de progression optimisé (toutes les 10% ou dernier message)
                                if msg_idx == 0 or (msg_idx + 1) % log_interval == 0 or msg_idx == total_messages - 1:
                                    elapsed = time_module.perf_counter() - batch_start_time
                                    rate = (msg_idx + 1) / elapsed if elapsed > 0 else 0
                                    print(f"[CONV] Progression: {msg_idx + 1}/{total_messages} ({100*(msg_idx+1)//total_messages}%) - {rate:.1f} msg/s")
                            except Exception as img_err:
                                print(f"[DEBUG] Erreur insertion image message: {img_err}")
                                error_run = new_para.add_run("[Image non disponible]")
                                error_run.font.color.rgb = RGBColor(150, 150, 150)
                                error_run.italic = True

                            # Espacement entre les messages
                            new_para.paragraph_format.space_before = Pt(10)
                            new_para.paragraph_format.space_after = Pt(10)
                            current_para = new_para

                            # Insérer les images attachées en haute qualité après la bulle
                            # L'utilisateur peut les redimensionner dans Word
                            image_attachments = msg_data.get("image_attachments", [])
                            if image_attachments:
                                print(f"[WORD IMG] Message {msg_idx}: {len(image_attachments)} image(s) à insérer")
                                for img_att in image_attachments:
                                    att_path = img_att.get("path", "")
                                    att_name = img_att.get("name", "")
                                    path_exists = Path(att_path).exists() if att_path else False
                                    print(f"[WORD IMG] -> {att_name}: path={att_path}, exists={path_exists}")
                                    if att_path and path_exists:
                                        try:
                                            att_para = insert_after_in_cell(current_para, "", "Normal")
                                            att_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_owner else WD_ALIGN_PARAGRAPH.LEFT
                                            att_run = att_para.add_run()
                                            # Compresser et insérer l'image (largeur 6cm)
                                            compressed_att = compress_image_for_docx(att_path)
                                            if compressed_att:
                                                att_run.add_picture(compressed_att, width=Cm(6))
                                            else:
                                                att_run.add_picture(att_path, width=Cm(6))
                                            att_para.paragraph_format.space_before = Pt(4)
                                            att_para.paragraph_format.space_after = Pt(8)
                                            current_para = att_para
                                            print(f"[WORD IMG] ✓ Image insérée: {att_name}")
                                        except Exception as att_err:
                                            print(f"[WORD IMG] ✗ Erreur: {att_err}")

                            # Insérer les PDF attachés (convertis en images) après les images
                            # L'utilisateur peut les redimensionner dans Word
                            pdf_attachments = msg_data.get("pdf_attachments", [])
                            if pdf_attachments:
                                for pdf_att in pdf_attachments:
                                    pdf_img_path = pdf_att.get("path", "")
                                    pdf_name = pdf_att.get("name", "")
                                    if pdf_img_path and Path(pdf_img_path).exists():
                                        try:
                                            pdf_para = insert_after_in_cell(current_para, "", "Normal")
                                            pdf_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_owner else WD_ALIGN_PARAGRAPH.LEFT
                                            pdf_run = pdf_para.add_run()
                                            # Compresser et insérer l'image du PDF (largeur 12cm pour meilleure visibilité)
                                            compressed_pdf = compress_image_for_docx(pdf_img_path)
                                            if compressed_pdf:
                                                pdf_run.add_picture(compressed_pdf, width=Cm(12))
                                            else:
                                                pdf_run.add_picture(pdf_img_path, width=Cm(12))
                                            pdf_para.paragraph_format.space_before = Pt(4)
                                            pdf_para.paragraph_format.space_after = Pt(8)
                                            current_para = pdf_para
                                        except Exception as pdf_err:
                                            print(f"[DEBUG] Erreur insertion PDF: {pdf_err}")

                            # Insérer le composite des miniatures vidéo APRÈS la bulle dans le docx
                            # (la bulle utilisée ici est la version CLEAN sans miniatures).
                            video_attachments = msg_data.get("video_attachments", [])
                            if video_attachments:
                                # Dedupliquer si plusieurs videos pointent vers le meme composite
                                seen_composites = set()
                                for vid_att in video_attachments:
                                    vid_composite = vid_att.get("composite_path", "")
                                    if not vid_composite or vid_composite in seen_composites:
                                        continue
                                    if not Path(vid_composite).exists():
                                        continue
                                    seen_composites.add(vid_composite)
                                    try:
                                        vid_para = insert_after_in_cell(current_para, "", "Normal")
                                        vid_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_owner else WD_ALIGN_PARAGRAPH.LEFT
                                        vid_run = vid_para.add_run()
                                        compressed_vid = compress_image_for_docx(vid_composite)
                                        if compressed_vid:
                                            vid_run.add_picture(compressed_vid, width=Cm(_compute_bubble_cm(vid_composite)))
                                        else:
                                            vid_run.add_picture(vid_composite, width=Cm(_compute_bubble_cm(vid_composite)))
                                        vid_para.paragraph_format.space_before = Pt(4)
                                        vid_para.paragraph_format.space_after = Pt(8)
                                        current_para = vid_para
                                    except Exception as vid_err:
                                        print(f"[WORD VID] Erreur insertion miniatures: {vid_err}")

                    # Résumé de fin après tous les messages
                    total_elapsed = time_module.perf_counter() - batch_start_time
                    avg_rate = total_messages / total_elapsed if total_elapsed > 0 else 0
                    print(f"[CONV] ✓ {total_messages} messages insérés en {total_elapsed:.1f}s ({avg_rate:.1f} msg/s)")

                elif image_paths:
                    # Insérer chaque image de la conversation (mode paginé)
                    print(f"[DEBUG] Insertion {len(image_paths)} image(s) conversation: {contact_name}")

                    from docx.shared import Inches

                    for img_idx, image_path in enumerate(image_paths):
                        if not Path(image_path).exists():
                            print(f"[DEBUG] Image non trouvée: {image_path}")
                            continue

                        new_para = insert_after_in_cell(current_para, "", "Normal")
                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        try:
                            # Largeur de l'image: 15 cm = ~5.9 pouces
                            run = new_para.add_run()
                            # Compresser l'image pour réduire la taille du fichier
                            compressed = compress_image_for_docx(image_path)
                            if compressed:
                                run.add_picture(compressed, width=Inches(5.9))
                            else:
                                run.add_picture(image_path, width=Inches(5.9))
                            print(f"[DEBUG] Image {img_idx + 1}/{len(image_paths)} insérée: {contact_name}")
                        except Exception as img_err:
                            print(f"[DEBUG] Erreur insertion image conversation: {img_err}")
                            error_run = new_para.add_run(f"[Image conversation non disponible: {contact_name}]")
                            error_run.font.color.rgb = RGBColor(150, 150, 150)
                            error_run.italic = True

                        new_para.paragraph_format.space_before = Pt(6)
                        new_para.paragraph_format.space_after = Pt(6)
                        current_para = new_para

                else:
                    # Fallback: affichage texte si pas d'image
                    messages = block.get("messages", [])
                    if not messages:
                        print(f"[DEBUG] Conversation ignorée: pas de messages ni d'image")
                        continue

                    print(f"[DEBUG] Insertion conversation texte avec {contact_name} ({source}): {len(messages)} messages")

                    # Titre de la conversation
                    title_para = insert_after_in_cell(current_para, "", "Normal")
                    title_run = title_para.add_run(f"Conversation avec {contact_name}")
                    title_run.bold = True
                    title_run.font.size = Pt(11)
                    title_run.font.color.rgb = RGBColor(30, 58, 95)
                    if source:
                        source_run = title_para.add_run(f" ({source})")
                        source_run.font.size = Pt(10)
                        source_run.font.color.rgb = RGBColor(100, 100, 100)
                    title_para.paragraph_format.space_before = Pt(12)
                    title_para.paragraph_format.space_after = Pt(8)
                    current_para = title_para

                    # Traiter chaque message
                    for msg in messages:
                        is_owner = msg.get("is_owner", False)
                        from_user = msg.get("from", "").replace("(owner)", "").strip()
                        body = msg.get("body", "")
                        attachments = msg.get("attachments", [])
                        timestamp = msg.get("timestamp_time", "") or msg.get("timestamp_date", "")

                        msg_para = insert_after_in_cell(current_para, "", "Normal")
                        sender_run = msg_para.add_run(from_user)
                        sender_run.bold = True
                        sender_run.font.size = Pt(9)
                        if is_owner:
                            sender_run.font.color.rgb = RGBColor(37, 211, 102)
                        else:
                            sender_run.font.color.rgb = RGBColor(0, 132, 255)

                        if body and body.strip() and body != "None":
                            msg_para.add_run("\n")
                            body_run = msg_para.add_run(body)
                            body_run.font.size = Pt(10)
                            body_run.font.color.rgb = RGBColor(30, 30, 30)

                        if attachments:
                            for att in attachments:
                                att_name = str(att).split("/")[-1].split("\\")[-1]
                                msg_para.add_run("\n")
                                att_run = msg_para.add_run(f"📎 {att_name}")
                                att_run.font.size = Pt(9)
                                att_run.font.color.rgb = RGBColor(100, 100, 100)
                                att_run.italic = True

                        if timestamp and timestamp != "nan" and timestamp != "None":
                            msg_para.add_run("\n")
                            ts_run = msg_para.add_run(timestamp)
                            ts_run.font.size = Pt(8)
                            ts_run.font.color.rgb = RGBColor(150, 150, 150)
                            ts_run.italic = True

                        if is_owner:
                            msg_para.paragraph_format.left_indent = Cm(2)
                        else:
                            msg_para.paragraph_format.left_indent = Cm(0)

                        msg_para.paragraph_format.space_before = Pt(6)
                        msg_para.paragraph_format.space_after = Pt(4)
                        current_para = msg_para

                    try:
                        current_para.paragraph_format.space_after = Pt(14)
                    except:
                        pass


def get_heading_number_prefix(para):
    """Extrait le préfixe numéroté d'un heading (ex: '2.3' de '2.3 Contacts')."""
    text = para.text.strip()
    match = re.match(r'^(\d+(?:\.\d+)*)\s+', text)
    if match:
        return match.group(1)
    return ""


def insert_platform_subheadings(doc, platforms_data: Dict[str, List[str]], edited_platform_titles: Optional[Dict[str, str]] = None):
    """Insère des sous-titres pour chaque plateforme sous les headings correspondants."""
    if not platforms_data:
        print("[DEBUG insert_platform_subheadings] Pas de platforms_data")
        return

    edited_titles = edited_platform_titles or {}
    print(f"[DEBUG insert_platform_subheadings] platforms_data: {platforms_data}")
    print(f"[DEBUG insert_platform_subheadings] edited_platform_titles: {edited_titles}")

    # Mapping des titres vers les clés de platforms_data
    # On cherche des correspondances partielles (le titre peut avoir un préfixe numéroté)
    heading_to_platforms_key = {
        "contacts": "contacts",
        "journal d'appel": "calls",
        "journal d'appels": "calls",
        "journal appel": "calls",
        "appels": "calls",
        "sms, mms, messageries et chats": "chats",
        "sms/mms/messageries/chats": "chats",
        "messageries et chats": "chats",
        "chats": "chats",
        "comptes associés": "accounts",
        "comptes associes": "accounts",
        "compte associé": "accounts",
        "user accounts": "accounts",
    }

    def normalize_text(text: str) -> str:
        t = (text or "").replace("’", "'").replace("`", "'").lower()
        t = re.sub(r"[^a-z0-9]+", " ", t)
        return re.sub(r"\s+", " ", t).strip()

    # Parcourir tous les blocs du document
    block_items = list(iter_block_items(doc))

    for idx, block in enumerate(block_items):
        if not isinstance(block, Paragraph) or not is_heading(block):
            continue

        para = block
        heading_text = para.text.strip()
        heading_level = get_heading_level(para)

        # Extraire le préfixe numéroté et le texte du titre
        prefix = get_heading_number_prefix(para)
        # Retirer le préfixe numéroté pour obtenir le texte pur du titre
        title_text_only = re.sub(r'^\d+(?:\.\d+)*\s+', '', heading_text).lower()
        title_text_norm = normalize_text(title_text_only)

        print(f"[DEBUG] Checking heading: '{heading_text}' (level={heading_level}, prefix='{prefix}', text='{title_text_only}')")

        # Chercher une correspondance avec les clés de platforms
        platforms_key = None
        for pattern, key in heading_to_platforms_key.items():
            pattern_norm = normalize_text(pattern)
            if pattern_norm in title_text_norm or title_text_norm in pattern_norm:
                platforms_key = key
                break

        if not platforms_key:
            continue

        platforms = platforms_data.get(platforms_key, [])
        if not platforms:
            print(f"[DEBUG] Heading '{heading_text}' correspond à '{platforms_key}' mais pas de plateformes")
            continue

        print(f"[DEBUG] Heading '{heading_text}' -> plateformes: {platforms}")

        # Créer les sous-titres pour chaque plateforme
        sub_level = heading_level + 1
        style_name = f"Heading {sub_level}"

        # Trouver le paragraphe après lequel insérer (après le heading et sa phrase descriptive)
        target_para = para

        # Si le prochain élément est un paragraphe (pas un heading), c'est probablement la phrase
        # descriptive, on insère après elle
        if idx + 1 < len(block_items):
            next_block = block_items[idx + 1]
            if isinstance(next_block, Paragraph) and not is_heading(next_block):
                target_para = next_block

        # Insérer les sous-titres en ordre inverse pour maintenir l'ordre final correct
        # Note: On n'ajoute PAS de numéro car le style Heading gère la numérotation automatique
        for sub_idx, platform in enumerate(reversed(platforms), 1):
            # Transformation spéciale pour Comptes associés
            display_platform = platform
            if platforms_key == "accounts":
                if platform.lower() == "accounts":
                    display_platform = "Compte utilisateur"

            # Vérifier si l'utilisateur a modifié le titre de cette plateforme
            # Les clés sont au format: "scoped:::ParentHeading|||PlatformTitle"
            for key, edited_title in edited_titles.items():
                if f"|||{platform}" in key or f"|||{display_platform}" in key:
                    print(f"[DEBUG] Titre modifié trouvé: '{key}' -> '{edited_title}'")
                    display_platform = edited_title
                    break

            # Insérer le sous-titre (juste le nom de la plateforme, Word gère la numérotation)
            # Ajouter 12pt d'espacement après pour aérer le contenu
            new_para = insert_after(target_para, display_platform, style_name, space_after_pt=12)
            print(f"[DEBUG] Inséré sous-titre '{display_platform}' (niveau {sub_level})")


def process_document(input_path, output_path, mapping_override: Optional[Dict[str, str]] = None,
                     decisions_override: Optional[List[str]] = None, interactive: bool = True,
                     heading_content: Optional[Dict[str, List[Dict]]] = None,
                     images_at_markers: Optional[Dict[str, str]] = None,
                     image_width_inches: float = 1.5,
                     images_at_markers_sizes: Optional[Dict[str, float]] = None,
                     headings_info: Optional[List[Dict]] = None,
                     platforms_data: Optional[Dict[str, List[str]]] = None,
                     edited_platform_titles: Optional[Dict[str, str]] = None,
                     extra_support_images: Optional[List[Dict]] = None,
                     account_images: Optional[Dict[str, str]] = None):
    input_path = Path(input_path)
    output_path = Path(output_path)

    doc = Document(str(input_path))

    placeholders = find_placeholders_in_order(doc)
    if mapping_override is not None:
        mapping = mapping_override
    elif interactive:
        mapping = prompt_placeholders(placeholders)
    else:
        mapping = {}

    # Remplacer dans les en-têtes
    for section in doc.sections:
        header = section.header
        # Paragraphes dans l'en-tête
        for p in header.paragraphs:
            replace_in_runs(p, mapping)
        # Tableaux dans l'en-tête
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_runs(p, mapping)

    # Supprimer les tableaux SIM vides AVANT le remplacement des placeholders
    remove_empty_sim_tables(doc, mapping)

    # Remplacer dans le corps du document
    for p in list(iter_all_paragraphs(doc)):
        removed = replace_in_runs(p, mapping)
        if removed:
            continue
    remove_empty_paragraphs(doc)

    # Note: L'insertion des titres personnalisés se fait APRÈS insert_platform_subheadings
    # pour qu'ils apparaissent après les sous-titres de plateforme

    headings = collect_headings_in_order(doc)

    # Créer un mapping titre->décision pour éviter les décalages d'index
    # (notamment quand des titres SIM sont supprimés par remove_empty_sim_tables)
    decisions_by_title = {}
    if decisions_override is not None:
        decisions = list(decisions_override)
        # S'assurer que decisions couvre tous les headings du document
        while len(decisions) < len(headings):
            decisions.append("__KEEP_TITLE_ONLY__")
        print(f"[DEBUG] decisions étendu de {len(decisions_override)} à {len(decisions)} pour couvrir {len(headings)} headings")

        # Créer le mapping titre->décision à partir de headings_info (titres originaux du frontend)
        if headings_info:
            for idx, info in enumerate(headings_info):
                title = info.get("title", "")
                if title and idx < len(decisions_override):
                    decisions_by_title[title] = decisions_override[idx]
                    print(f"[DEBUG] Mapping titre->décision: '{title}' -> '{decisions_override[idx][:30] if decisions_override[idx] else '(vide)'}...'")
        else:
            # Fallback: utiliser les titres actuels du document
            for idx, h in enumerate(headings):
                if idx < len(decisions):
                    decisions_by_title[h.text.strip()] = decisions[idx]
    elif interactive:
        decisions = collect_heading_decisions(headings, mapping)
    else:
        decisions = default_heading_decisions(headings, mapping)

    apply_heading_decisions(doc, decisions, decisions_by_title=decisions_by_title if decisions_by_title else None)
    remove_empty_paragraphs(doc)

    # Insertion automatique des sous-titres par plateforme (Contacts, Journal d'appel, Chats)
    if platforms_data:
        insert_platform_subheadings(doc, platforms_data, edited_platform_titles)

    # Insertion des titres personnalisés APRÈS les sous-titres de plateforme
    if headings_info:
        print(f"[DEBUG] headings_info reçu: {len(headings_info)} titres")
        for i, info in enumerate(headings_info):
            print(f"  [{i}] title='{info.get('title')}', level={info.get('level')}, isCustom={info.get('isCustom')}")

        # Collecter les titres custom à insérer
        custom_titles = []
        for info_idx, info in enumerate(headings_info):
            if info.get("isCustom", False):
                title = info.get("title", "Nouveau titre")
                level = info.get("level", 1)

                # Trouver le dernier heading non-custom avant celui-ci
                prev_idx = info_idx - 1
                while prev_idx >= 0 and headings_info[prev_idx].get("isCustom", False):
                    prev_idx -= 1

                prev_title = headings_info[prev_idx].get("title", "") if prev_idx >= 0 else None
                custom_titles.append({
                    "title": title,
                    "level": level,
                    "prev_title": prev_title,
                    "prev_level": headings_info[prev_idx].get("level", 1) if prev_idx >= 0 else 1
                })

        if custom_titles:
            print(f"[DEBUG] {len(custom_titles)} titres personnalisés à insérer")

            # Fonction pour normaliser le texte de comparaison
            def normalize_for_match(text):
                if not text:
                    return ""
                t = text.lower().replace("'", "'").replace("`", "'")
                t = re.sub(r'^\d+(?:\.\d+)*\s*', '', t)  # Retirer préfixe numéroté
                t = re.sub(r'[^a-z0-9àâäéèêëïîôùûüç]+', ' ', t)
                return ' '.join(t.split()).strip()

            # Parcourir les titres custom dans l'ordre
            for custom in custom_titles:
                title = custom["title"]
                level = custom["level"]
                prev_title = custom["prev_title"]
                prev_level = custom["prev_level"]
                style_name = f"Heading {level}"

                print(f"[DEBUG] Insertion de '{title}' (niveau {level}) après '{prev_title}' (niveau {prev_level})")

                # Rafraîchir la liste des headings
                all_blocks = list(iter_block_items(doc))
                all_heading_paras = [b for b in all_blocks if isinstance(b, Paragraph) and is_heading(b)]

                # Trouver le heading parent dans le document
                target_para = None
                parent_found_idx = -1
                prev_title_norm = normalize_for_match(prev_title)

                for idx, para in enumerate(all_heading_paras):
                    para_text_norm = normalize_for_match(para.text)
                    if prev_title_norm and (prev_title_norm in para_text_norm or para_text_norm in prev_title_norm):
                        parent_found_idx = idx
                        target_para = para
                        break

                if parent_found_idx >= 0:
                    # Si le parent a des sous-titres de plateforme (niveau parent+1),
                    # insérer après le dernier sous-titre de ce niveau
                    expected_sub_level = prev_level + 1
                    last_sub_para = target_para

                    # Parcourir les headings suivants pour trouver les sous-titres du parent
                    for check_idx in range(parent_found_idx + 1, len(all_heading_paras)):
                        check_para = all_heading_paras[check_idx]
                        check_level = get_heading_level(check_para)

                        if check_level <= prev_level:
                            # On est sorti des sous-titres du parent
                            break
                        elif check_level == expected_sub_level:
                            # C'est un sous-titre du parent
                            last_sub_para = check_para

                    # Insérer après le dernier sous-titre trouvé (ou après le parent)
                    new_para = insert_after(last_sub_para, title, style_name, space_after_pt=12)
                    print(f"[DEBUG] Inséré titre personnalisé '{title}' (niveau {level}) après '{last_sub_para.text[:50]}...'")
                else:
                    # Parent non trouvé, insérer à la fin des headings
                    if all_heading_paras:
                        last_para = all_heading_paras[-1]
                        new_para = insert_after(last_para, title, style_name, space_after_pt=12)
                        print(f"[DEBUG] Inséré titre personnalisé '{title}' à la fin (parent '{prev_title}' non trouvé)")
                    else:
                        print(f"[WARN] Impossible d'insérer '{title}': pas de headings dans le document")

    # Insertion des blocs de contenu (texte + images) après les headings APRÈS le remplacement des placeholders
    # Il faut remplacer les placeholders dans les clés du dictionnaire heading_content
    if heading_content:
        # Créer un nouveau dictionnaire avec les placeholders remplacés dans les clés
        heading_content_replaced = {}
        print(f"[DEBUG process_document] Remplacement des clés heading_content:")
        for heading_key, blocks in heading_content.items():
            # Remplacer les placeholders dans la clé du heading
            replaced_key = heading_key
            for placeholder, value in mapping.items():
                replaced_key = replaced_key.replace(placeholder, value or "")
            print(f"  '{heading_key}' -> '{replaced_key}' ({len(blocks)} blocs)")
            heading_content_replaced[replaced_key] = blocks

        apply_heading_content_blocks(doc, heading_content_replaced, default_width_inches=image_width_inches, mapping=mapping, account_images=account_images)

    # Insertion d'images sur les marqueurs
    if images_at_markers:
        apply_images_at_markers(doc, images_at_markers, width_inches=image_width_inches,
                                per_image_widths=images_at_markers_sizes)

    # Insertion des images supplémentaires pour "Photographies des supports numériques"
    if extra_support_images:
        insert_extra_support_images(doc, extra_support_images)

    # Désactiver les emails/URLs cliquables dans le PDF
    disable_clickable_emails(doc)

    # Forcer la mise à jour de la table des matières à l'ouverture
    update_table_of_contents(doc)

    # Corriger les tabulations des numéros de titres (évite les écarts avec 2+ chiffres)
    fix_heading_numbering_tabs(doc)

    # Détecter si c'est test2 (carte SIM) - les 3 premiers tableaux doivent rester sur page 1
    is_test2 = "test2" in str(input_path).lower()

    if is_test2:
        # Pour test2: réduire les espacements pour garder les 3 tableaux sur la page 1
        reduce_spacing_for_first_tables(doc, num_tables=3)
        print("[DEBUG] Mode test2: espacements réduits, pas de saut de page avant Photographies")
    else:
        # Pour test et test3: espacement normal et saut de page avant Photographies
        add_spacing_between_tables(doc, spacing_pt=12)
        add_page_break_before_heading(doc, "Photographies des supports numériques")

    # Sauvegarder avec gestion du fichier verrouillé
    try:
        doc.save(str(output_path))
        print(f"Document genere : {output_path}")
    except PermissionError:
        raise PermissionError(f"Le fichier '{output_path.name}' est ouvert dans Word. Fermez-le et réessayez.")


def main():
    parser = argparse.ArgumentParser(description="Automation interactive pour Word (placeholders + phrases sous titres)")
    parser.add_argument("--input", default="test.docx", help="Fichier Word source")
    parser.add_argument("--output", default="test_sortie.docx", help="Fichier Word de sortie")
    args = parser.parse_args()
    process_document(args.input, args.output)


if __name__ == "__main__":
    main()
