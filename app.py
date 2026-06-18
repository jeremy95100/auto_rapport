import asyncio
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import io
import json
import uuid
import time
import hashlib
from datetime import datetime
import zipfile
import re
import xml.etree.ElementTree as ET
from collections import Counter
import html  # Pour décoder les entités HTML (&amp; → &)
import threading  # Pour lock matplotlib (non thread-safe)

import polars as pl
import pandas as pd
import mammoth
import matplotlib
matplotlib.use('Agg')  # Backend non-interactif pour génération de graphiques
import matplotlib.pyplot as plt
from docx import Document
from fastapi import FastAPI, HTTPException, UploadFile, File, Header, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, StreamingResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from PIL import Image, ImageDraw, ImageFont, ImageFilter, ImageEnhance
import pillow_heif

from remplace_rapport import (
    collect_headings_in_order,
    collect_headings_with_levels,
    DEFAULT_ANALYSIS_TEMPLATE,
    default_heading_decisions,
    find_placeholders_in_order,
    find_image_markers_in_order,
    process_document,
)

# Module de traitement Excel streaming haute performance
from excel_processor import (
    process_excel_streaming,
    LazyAnalyzer,
    ImportResult,
    cleanup_import,
    cleanup_previous_import,
    list_imports,
    get_import_path,
    strip_html_tags,
    SKIP_SHEETS,
    EXCLUDED_CONTACT_SOURCES,
    EXCLUDED_CHAT_TOP15_SOURCES,
    MAX_UPLOAD_SIZE_MB,
)

TEMPLATES = {
    "test": Path("test.docx"),
    "test2": Path("test2.docx"),
    "test3": Path("test3.docx"),
}
OUTPUT_PATHS = {name: path.with_name(f"{path.stem}_sortie.docx") for name, path in TEMPLATES.items()}
FRONTEND_DIR = Path("frontend")
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# ÉTAT DE L'IMPORT COURANT (système streaming par UUID)

_current_import: Optional[ImportResult] = None
_current_analyzer: Optional[LazyAnalyzer] = None

# Rétrocompatibilité avec l'ancien système
_last_excel_content: bytes = b""
_current_excel_hash: str = ""

# Identifiants du propriétaire du téléphone (depuis User Accounts)
_owner_usernames: set = set()

# Lock pour la génération de graphiques matplotlib (non thread-safe)
_chart_lock = threading.Lock()

_cors_from_env = [o.strip() for o in os.getenv("CORS_ALLOWED_ORIGINS", "").split(",") if o.strip()]
ALLOWED_ORIGINS = _cors_from_env or [
    "http://127.0.0.1:8000",
    "http://localhost:8000",
]
API_KEY = os.getenv("RAPPORT_API_KEY", "").strip()

app = FastAPI(title="Rapport auto - API")  # HEIC support enabled
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["*"],
    allow_headers=["*"],
)

listeners: List[asyncio.Queue] = []
doc_lock = asyncio.Lock()

# ===== SYSTÈME DE PERSISTANCE DE SESSION =====
DATA_DIR = Path("data")
DATA_DIR.mkdir(exist_ok=True)
SESSION_DIR = DATA_DIR / "sessions"
SESSION_DIR.mkdir(exist_ok=True)
SESSION_FILE = SESSION_DIR / "current.json"

def save_session():
    """Sauvegarde l'état de la session courante dans un fichier JSON."""
    global _current_import, _tagged_data, _owner_usernames, _last_platforms_data

    try:
        session_data = {
            "timestamp": datetime.now().isoformat(),
            "current_import": None,
            "tagged_data": _tagged_data,
            "owner_usernames": list(_owner_usernames) if _owner_usernames else [],
            "platforms_data": _last_platforms_data,
        }

        # Sauvegarder les infos de l'import courant
        if _current_import:
            session_data["current_import"] = {
                "import_id": _current_import.import_id,
                "import_path": str(_current_import.import_path),
                "sheets": _current_import.sheets,
                "parquet_files": {k: str(v) for k, v in _current_import.parquet_files.items()},
                "device_info": _current_import.device_info,
                "row_counts": _current_import.row_counts,
            }

        # Convertir les Path en str pour JSON
        tagged_data_serializable = {}
        for key, value in _tagged_data.items():
            if key == "parquet_dir" and value:
                tagged_data_serializable[key] = str(value)
            elif key == "sheets_info":
                sheets_info_ser = {}
                for sheet_name, sheet_info in value.items():
                    sheet_info_copy = dict(sheet_info)
                    if "parquet_path" in sheet_info_copy:
                        sheet_info_copy["parquet_path"] = str(sheet_info_copy["parquet_path"])
                    sheets_info_ser[sheet_name] = sheet_info_copy
                tagged_data_serializable[key] = sheets_info_ser
            else:
                tagged_data_serializable[key] = value

        session_data["tagged_data"] = tagged_data_serializable

        with open(SESSION_FILE, "w", encoding="utf-8") as f:
            json.dump(session_data, f, ensure_ascii=False, indent=2, default=str)

        print(f"[SESSION] Sauvegardée: {SESSION_FILE}")
        return True
    except Exception as e:
        print(f"[SESSION] Erreur sauvegarde: {e}")
        import traceback
        traceback.print_exc()
        return False

def restore_session() -> bool:
    """Restaure l'état de la session depuis le fichier JSON."""
    global _current_import, _current_analyzer, _tagged_data, _owner_usernames, _last_platforms_data

    try:
        if not SESSION_FILE.exists():
            print("[SESSION] Aucune session sauvegardée")
            return False

        with open(SESSION_FILE, "r", encoding="utf-8") as f:
            session_data = json.load(f)

        print(f"[SESSION] Restauration depuis {session_data.get('timestamp', 'inconnu')}")

        # Restaurer l'import courant
        if session_data.get("current_import"):
            imp_data = session_data["current_import"]
            import_path = Path(imp_data["import_path"])

            # Vérifier que le dossier existe toujours
            if import_path.exists():
                _current_import = ImportResult(
                    import_id=imp_data["import_id"],
                    import_path=import_path,
                    sheets=imp_data["sheets"],
                    parquet_files={k: Path(v) for k, v in imp_data["parquet_files"].items()},
                    device_info=imp_data.get("device_info", {}),
                    row_counts=imp_data.get("row_counts", {}),
                )
                # Recréer l'analyzer
                _current_analyzer = LazyAnalyzer(import_path)
                print(f"[SESSION] Import restauré: {imp_data['import_id']}")
            else:
                print(f"[SESSION] Dossier import introuvable: {import_path}")

        # Restaurer tagged_data
        if session_data.get("tagged_data"):
            td = session_data["tagged_data"]
            _tagged_data["images"] = td.get("images", [])
            _tagged_data["videos"] = td.get("videos", [])
            _tagged_data["audios"] = td.get("audios", [])
            _tagged_data["chats"] = td.get("chats", [])
            _tagged_data["call_log"] = td.get("call_log", [])
            _tagged_data["conversations"] = td.get("conversations", {})
            if td.get("parquet_dir"):
                _tagged_data["parquet_dir"] = Path(td["parquet_dir"])
            if td.get("sheets_info"):
                sheets_info = {}
                for sheet_name, sheet_info in td["sheets_info"].items():
                    sheet_info_copy = dict(sheet_info)
                    if "parquet_path" in sheet_info_copy:
                        sheet_info_copy["parquet_path"] = Path(sheet_info_copy["parquet_path"])
                    sheets_info[sheet_name] = sheet_info_copy
                _tagged_data["sheets_info"] = sheets_info
            print(f"[SESSION] Tagged data restauré")

        # Restaurer owner_usernames
        if session_data.get("owner_usernames"):
            _owner_usernames = set(session_data["owner_usernames"])
            print(f"[SESSION] Owners restaurés: {_owner_usernames}")

        # Restaurer platforms_data
        if session_data.get("platforms_data"):
            _last_platforms_data.update(session_data["platforms_data"])
            print(f"[SESSION] Platforms restaurées")

        print("[SESSION] Restauration terminée avec succès")
        return True
    except Exception as e:
        print(f"[SESSION] Erreur restauration: {e}")
        import traceback
        traceback.print_exc()
        return False

# Restaurer la session au démarrage
restore_session()


class ImageTextData(BaseModel):
    before: str = ""
    after: str = ""
    position: str = "none"  # "before", "after", or "none"

class ContentBlock(BaseModel):
    type: str  # "text", "image", "table", "conversation", or "video"
    content: Optional[str] = None  # for text blocks
    accountImage: Optional[str] = None  # image associée au bloc de compte (pour les blocs texte)
    accountImageWidth: Optional[float] = None  # largeur de l'image de compte en pouces
    src: Optional[str] = None  # for image/video blocks
    width: Optional[float] = None  # for image blocks (in inches)
    caption: Optional[str] = None  # for image blocks - text caption
    captionPosition: Optional[str] = None  # for image blocks - "top", "bottom", "left", "right"
    # for table blocks
    sheet: Optional[str] = None  # Excel sheet name
    columns: Optional[List[str]] = None  # Selected columns (original names for data access)
    displayColumns: Optional[List[str]] = None  # Display names for column headers (user-editable)
    data: Optional[List[Dict[str, Any]]] = None  # Table data
    # for conversation blocks
    convId: Optional[str] = None  # Conversation ID
    contactName: Optional[str] = None  # Contact name
    source: Optional[str] = None  # Source (Snapchat, WhatsApp, etc.)
    periodStart: Optional[str] = None  # Début de la période de communication
    periodEnd: Optional[str] = None  # Fin de la période de communication
    contactInfo: Optional[Dict[str, Any]] = None  # Infos du contact (pseudonyme, nom_utilisateur, identifiant_utilisateur, is_social_source, is_phone_source)
    # for conversation with individual messages and comments
    messageImages: Optional[List[Dict[str, Any]]] = None  # Liste des messages avec images et commentaires
    # for video blocks
    name: Optional[str] = None  # Video name
    thumbnails: Optional[List[str]] = None  # List of thumbnail URLs/paths
    metadata: Optional[Dict[str, Any]] = None  # Video metadata (filePath, createdDate, observation)

class HeadingInfo(BaseModel):
    title: str
    level: int
    isCustom: bool = False  # True si ajouté par l'utilisateur

class PlatformsData(BaseModel):
    contacts: List[str] = Field(default_factory=list)
    calls: List[str] = Field(default_factory=list)
    chats: List[str] = Field(default_factory=list)
    accounts: List[str] = Field(default_factory=list)

class ExtraSupportImage(BaseModel):
    src: str
    width_inches: float = 3.0

class GeneratePayload(BaseModel):
    template: Optional[str] = None
    overwrite: bool = False
    mapping: Dict[str, str] = Field(default_factory=dict)
    decisions: Optional[List[str]] = None
    heading_content: Dict[str, List[ContentBlock]] = Field(default_factory=dict)  # nouvelle structure
    images_at_markers: Dict[str, str] = Field(default_factory=dict)
    image_width_inches: Optional[float] = None
    images_at_markers_sizes: Dict[str, float] = Field(default_factory=dict)
    # Nouvelle structure pour les titres modifiés/ajoutés
    headings_info: Optional[List[HeadingInfo]] = None  # Liste complète des titres avec niveaux
    # Auto-création des sous-titres par plateforme
    auto_platform_subheadings: bool = True  # Activer la création auto des sous-titres par plateforme
    platforms_data: Optional[PlatformsData] = None  # Données des plateformes (si non fourni, utilise les dernières données importées)
    edited_platform_titles: Dict[str, str] = Field(default_factory=dict)  # Titres de plateformes modifiés par l'utilisateur
    extra_support_images: List[ExtraSupportImage] = Field(default_factory=list)  # Images supplémentaires pour Photographies des supports
    account_images: Dict[str, str] = Field(default_factory=dict)  # Images associées aux blocs de compte: {headingKey_blockIdx: imageSrc}


def require_api_key(
    request: Request,
    x_api_key: Optional[str] = Header(default=None, alias="X-API-Key"),
):
    client_host = request.client.host if request.client else ""
    is_local_client = client_host in {"127.0.0.1", "::1", "localhost"}

    # Local development remains usable without key.
    if not API_KEY and is_local_client:
        return

    # Non-local usage requires a configured key.
    if not API_KEY:
        raise HTTPException(
            status_code=500,
            detail="Server misconfigured: RAPPORT_API_KEY is required for non-local access",
        )

    if x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")


def available_templates() -> List[str]:
    # On liste toutes les trames déclarées ; l'existence est vérifiée plus tard
    return list(TEMPLATES.keys())


def get_template_paths(selected: Optional[str] = None) -> Tuple[Path, Path, str]:
    if selected:
        if selected not in TEMPLATES:
            raise HTTPException(status_code=404, detail=f"Trame inconnue: {selected}")
        src = TEMPLATES[selected]
        if not src.exists():
            raise HTTPException(status_code=404, detail=f"Source file not found: {src}")
        return src, OUTPUT_PATHS[selected], selected
    for name, src in TEMPLATES.items():
        if src.exists():
            return src, OUTPUT_PATHS[name], name
    raise HTTPException(status_code=404, detail="Aucune trame disponible")


def convert_to_html(docx_path: Path) -> str:
    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="Document introuvable pour l'aperçu HTML")

    # Extraire les en-têtes avec python-docx
    doc = Document(str(docx_path))
    headers_html = ""

    # Récupérer les en-têtes de toutes les sections
    for section in doc.sections:
        header = section.header

        # Vérifier les paragraphes dans l'en-tête
        for para in header.paragraphs:
            if para.text.strip():
                style = 'style="text-align: center; font-weight: bold; margin-bottom: 10px;"'
                headers_html += f'<div {style}>{para.text}</div>\n'

        # Vérifier les tableaux dans l'en-tête (cas de test.docx et test2.docx)
        for table in header.tables:
            headers_html += '<table style="width: 100%; border-collapse: collapse; margin-bottom: 10px;">\n'
            for row in table.rows:
                headers_html += '  <tr>\n'
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Déterminer l'alignement et le style selon le contenu
                        if "Rapport" in cell_text or "Art." in cell_text:
                            style = 'style="text-align: center; font-weight: bold; padding: 5px; vertical-align: middle;"'
                        else:
                            style = 'style="text-align: right; padding: 5px; vertical-align: middle;"'
                        # Préserver les sauts de ligne
                        cell_html = cell_text.replace('\n', '<br>')
                        headers_html += f'    <td {style}>{cell_html}</td>\n'
                    else:
                        headers_html += '    <td style="padding: 5px;"></td>\n'
                headers_html += '  </tr>\n'
            headers_html += '</table>\n'

    # Si pas d'en-têtes trouvés dans les sections, chercher dans les premiers paragraphes
    if not headers_html:
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if text and "{" in text:
                style = 'style="text-align: center; font-weight: bold; margin-bottom: 10px;"'
                headers_html += f'<div {style}>{text}</div>\n'

    # Convertir le reste du document avec mammoth
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)

    # Combiner l'en-tête et le contenu
    if headers_html:
        full_html = f'<div style="border-bottom: 2px solid #333; padding-bottom: 20px; margin-bottom: 20px;">\n{headers_html}</div>\n{result.value}'
    else:
        full_html = result.value

    return full_html


async def broadcast(message: str) -> None:
    for q in list(listeners):
        await q.put(message)


async def event_stream():
    queue: asyncio.Queue = asyncio.Queue()
    listeners.append(queue)
    try:
        while True:
            msg = await queue.get()
            yield f"data: {msg}\n\n"
    finally:
        listeners.remove(queue)


@app.get("/")
def root():
    if FRONTEND_DIR.exists():
        return RedirectResponse(url="/ui/")
    return {"message": "API running", "routes": ["/templates", "/placeholders", "/generate", "/preview", "/events"]}


@app.get("/templates")
def list_templates():
    templates = available_templates()
    default = None
    for name in templates:
        if TEMPLATES[name].exists():
            default = name
            break
    if default is None:
        default = templates[0] if templates else None
    return {"templates": templates, "default": default}


@app.get("/placeholders")
def get_placeholders(template: Optional[str] = None):
    src_path, _, template_key = get_template_paths(template)
    doc = Document(str(src_path))
    placeholders = find_placeholders_in_order(doc)
    headings_with_levels = collect_headings_with_levels(doc)
    headings = [h["text"] for h in headings_with_levels]
    markers = find_image_markers_in_order(doc)
    return {
        "template": template_key,
        "templates": available_templates(),
        "placeholders": placeholders,
        "headings": headings,
        "headingsWithLevels": headings_with_levels,
        "markers": markers,
        "default_template": DEFAULT_ANALYSIS_TEMPLATE,
    }


@app.post("/generate", dependencies=[Depends(require_api_key)])
async def generate(payload: GeneratePayload):
    src_path, output_path, template_key = get_template_paths(payload.template)
    mapping = payload.mapping or {}
    if payload.overwrite:
        if output_path.exists():
            try:
                output_path.unlink()
            except PermissionError:
                print(f"[WARNING] Impossible de supprimer {output_path.name}, fichier probablement ouvert. Le document sera écrasé.")
    doc = Document(str(src_path))
    placeholders_in_doc = find_placeholders_in_order(doc)
    for missing in placeholders_in_doc:
        mapping.setdefault(missing, "")
    headings = collect_headings_in_order(doc)
    decisions = payload.decisions
    if decisions is None:
        decisions = default_heading_decisions(headings, mapping)
    else:
        # "__DEFAULT__" => phrase auto, "__KEEP_TITLE_ONLY__" => garder titre sans phrase, "" => supprimer, autre => texte personnalisé
        resolved = []
        defaults = default_heading_decisions(headings, mapping)
        for idx, h in enumerate(headings):
            choice = payload.decisions[idx] if idx < len(payload.decisions) else "__DEFAULT__"
            if choice == "__DEFAULT__" or choice is None:
                resolved.append(defaults[idx])
            elif choice == "__KEEP_TITLE_ONLY__":
                # Garder le titre mais ne rien ajouter en dessous
                resolved.append("__KEEP_TITLE_ONLY__")
            else:
                resolved.append(choice)
        decisions = resolved

    def resolve_path(path_str: str) -> str:
        if path_str.startswith("/uploads/"):
            return str(UPLOAD_DIR / Path(path_str).name)
        return path_str

    # Convert heading_content blocks to resolved paths
    heading_content_resolved = {}
    for heading, blocks in (payload.heading_content or {}).items():
        heading_content_resolved[heading] = []
        for block in blocks:
            block_dict = block.model_dump()
            # Pour les blocs texte avec image de compte
            if block.type == "text" and block.accountImage:
                block_dict["accountImage"] = resolve_path(block.accountImage)
                print(f"[DEBUG] accountImage résolu: {block_dict['accountImage']}")
            if block.type == "image" and block.src:
                block_dict["src"] = resolve_path(block.src)
            # Pour les vidéos, résoudre les chemins des miniatures
            elif block.type == "video":
                if block.src:
                    block_dict["src"] = resolve_path(block.src)
                # Résoudre les chemins des miniatures
                if hasattr(block, 'thumbnails') and block.thumbnails:
                    resolved_thumbs = [resolve_path(t) for t in block.thumbnails]
                    block_dict["thumbnails"] = resolved_thumbs
                    print(f"[DEBUG VIDEO] Bloc vidéo: name={block.name}, thumbnails={len(resolved_thumbs)}")
                else:
                    print(f"[DEBUG VIDEO] Bloc vidéo sans miniatures: name={getattr(block, 'name', 'N/A')}")
            # Pour les conversations, générer une image PNG
            elif block.type == "conversation" and block.convId:
                conv_id = block.convId
                if conv_id in _tagged_data.get("conversations", {}):
                    conv_data = _tagged_data["conversations"][conv_id]
                    messages = conv_data.get("messages", [])
                    contact_name = conv_data.get("contact_name", block.contactName or "")
                    source = conv_data.get("source", block.source or "")

                    # Vérifier si le frontend a envoyé les messageImages (avec commentaires)
                    if block.messageImages:
                        # Résoudre les chemins des images
                        resolved_messages = []
                        for msg_img in block.messageImages:
                            resolved_msg = dict(msg_img)
                            if "image_path" in resolved_msg:
                                resolved_msg["image_path"] = resolve_path(resolved_msg["image_path"])
                            resolved_messages.append(resolved_msg)
                        block_dict["messageImages"] = resolved_messages
                        block_dict["contactName"] = contact_name
                        block_dict["source"] = source
                        # Ajouter les infos du tableau récapitulatif (envoyées par le frontend)
                        block_dict["periodStart"] = block.periodStart or ""
                        block_dict["periodEnd"] = block.periodEnd or ""
                        block_dict["contactInfo"] = block.contactInfo or {}
                        print(f"[DEBUG] Conversation {conv_id}: {len(resolved_messages)} messages avec commentaires")
                        print(f"[DEBUG] Conversation {conv_id}: periodStart={block.periodStart}, periodEnd={block.periodEnd}")
                        print(f"[DEBUG] Conversation {conv_id}: contactInfo={block.contactInfo}")
                    else:
                        # Ancien système: générer les images groupées
                        image_paths = generate_conversation_image(messages, contact_name, source)
                        if image_paths:
                            block_dict["imagePaths"] = image_paths
                            block_dict["contactName"] = contact_name
                            block_dict["source"] = source
                            block_dict["periodStart"] = block.periodStart or ""
                            block_dict["periodEnd"] = block.periodEnd or ""
                            block_dict["contactInfo"] = block.contactInfo or {}
                            print(f"[DEBUG] Conversation {conv_id}: {len(image_paths)} image(s) générée(s)")
                        else:
                            block_dict["messages"] = messages
                            block_dict["contactName"] = contact_name
                            block_dict["source"] = source
                            block_dict["periodStart"] = block.periodStart or ""
                            block_dict["periodEnd"] = block.periodEnd or ""
                            block_dict["contactInfo"] = block.contactInfo or {}
                            print(f"[DEBUG] Conversation {conv_id}: {len(messages)} messages (fallback texte)")
            heading_content_resolved[heading].append(block_dict)

    # Resolve marker images
    markers_resolved = {}
    for k, v in (payload.images_at_markers or {}).items():
        markers_resolved[k] = resolve_path(v)

    # Préparer headings_info pour les titres personnalisés
    headings_info = None
    if payload.headings_info:
        headings_info = [{"title": h.title, "level": h.level, "isCustom": h.isCustom} for h in payload.headings_info]

    # Préparer les données des plateformes pour création auto des sous-titres
    platforms_data = None
    if payload.auto_platform_subheadings:
        if payload.platforms_data:
            platforms_data = {
                "contacts": payload.platforms_data.contacts,
                "calls": payload.platforms_data.calls,
                "chats": payload.platforms_data.chats,
                "accounts": getattr(payload.platforms_data, 'accounts', [])
            }
        else:
            # Utiliser les dernières données importées
            platforms_data = _last_platforms_data
        print(f"[DEBUG] platforms_data pour sous-titres auto: {platforms_data}")

    # Debug logging
    print(f"[DEBUG] mapping: {mapping}")
    print(f"[DEBUG] decisions: {decisions}")
    print(f"[DEBUG] heading_content_resolved: {heading_content_resolved}")
    print(f"[DEBUG] markers_resolved: {markers_resolved}")
    print(f"[DEBUG] headings_info: {headings_info}")

    async with doc_lock:
        try:
            # Convertir les images supplémentaires
            extra_support_images_resolved = []
            print(f"[DEBUG] Payload extra_support_images: {payload.extra_support_images}")
            print(f"[DEBUG] UPLOAD_DIR absolu: {UPLOAD_DIR.absolute()}")
            for img in (payload.extra_support_images or []):
                if img.src:
                    # Le src peut être "/uploads/xxx.jpg" ou juste "xxx.jpg"
                    src_clean = img.src.replace("/uploads/", "").replace("uploads/", "")
                    img_path = UPLOAD_DIR / Path(src_clean).name
                    img_path_abs = img_path.absolute()
                    print(f"[DEBUG] Extra support image: src={img.src}, src_clean={src_clean}, path={img_path_abs}, exists={img_path_abs.exists()}")
                    if img_path_abs.exists():
                        extra_support_images_resolved.append({
                            "path": str(img_path_abs),
                            "width_inches": img.width_inches
                        })
                        print(f"[DEBUG]   -> Résolu: {img_path_abs}")
                    else:
                        print(f"[ATTENTION] Image supplémentaire non trouvée: {img_path_abs}")
            print(f"[DEBUG] {len(extra_support_images_resolved)} images supplémentaires résolues: {extra_support_images_resolved}")

            # Résoudre les chemins des images de compte
            account_images_resolved = {}
            for key, src in (payload.account_images or {}).items():
                if src:
                    if src.startswith("/uploads/"):
                        account_images_resolved[key] = src.replace("/uploads/", "uploads/")
                    else:
                        account_images_resolved[key] = src
            print(f"[DEBUG] {len(account_images_resolved)} images de compte résolues")

            process_document(
                src_path,
                output_path,
                mapping_override=mapping,
                decisions_override=decisions,
                interactive=False,
                heading_content=heading_content_resolved,
                images_at_markers=markers_resolved,
                image_width_inches=payload.image_width_inches or 1.5,
                images_at_markers_sizes=payload.images_at_markers_sizes or {},
                headings_info=headings_info,
                platforms_data=platforms_data,
                edited_platform_titles=payload.edited_platform_titles or {},
                extra_support_images=extra_support_images_resolved,
                account_images=account_images_resolved,
            )
        except PermissionError as e:
            raise HTTPException(status_code=423, detail=str(e))
    await broadcast("updated")
    return {"status": "ok", "template": template_key, "output": str(output_path), "pdf": None}


@app.get("/download")
async def download(template: Optional[str] = None):
    """Télécharge le fichier Word généré"""
    _, output_path, _ = get_template_paths(template)
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Fichier de sortie non trouvé. Générez d'abord le document.")
    return FileResponse(
        path=str(output_path),
        filename=output_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.get("/preview")
async def preview(template: Optional[str] = None):
    src_path, output_path, template_key = get_template_paths(template)
    target = output_path if output_path.exists() else src_path
    if not target.exists():
        raise HTTPException(status_code=404, detail="Aucun fichier de reference disponible")
    async with doc_lock:
        html = await asyncio.get_running_loop().run_in_executor(None, convert_to_html, target)
    return HTMLResponse(html)


@app.get("/preview/html")
async def preview_html(template: Optional[str] = None):
    src_path, output_path, _ = get_template_paths(template)
    target = output_path if output_path.exists() else src_path
    if not target.exists():
        raise HTTPException(status_code=404, detail="Aucun fichier de reference disponible")
    async with doc_lock:
        html = await asyncio.get_running_loop().run_in_executor(None, convert_to_html, target)
    return HTMLResponse(html)


@app.get("/events")
async def events():
    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/upload", dependencies=[Depends(require_api_key)])
async def upload_image(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nom de fichier manquant")

    contents = await file.read()
    original_filename = file.filename
    filename_lower = original_filename.lower()

    print(f"[UPLOAD DEBUG] Fichier reçu: {original_filename}, extension détectée: {filename_lower}")

    # Convertir HEIC en JPEG si nécessaire
    if filename_lower.endswith('.heic') or filename_lower.endswith('.heif'):
        try:
            # Enregistrer le support HEIF
            pillow_heif.register_heif_opener()

            # Ouvrir l'image HEIC
            heic_image = Image.open(io.BytesIO(contents))

            # Convertir en RGB si nécessaire (HEIC peut être en RGBA)
            if heic_image.mode in ('RGBA', 'LA', 'P'):
                rgb_image = Image.new('RGB', heic_image.size, (255, 255, 255))
                if heic_image.mode == 'P':
                    heic_image = heic_image.convert('RGBA')
                rgb_image.paste(heic_image, mask=heic_image.split()[-1] if heic_image.mode == 'RGBA' else None)
                heic_image = rgb_image
            elif heic_image.mode != 'RGB':
                heic_image = heic_image.convert('RGB')

            # Nouveau nom de fichier en .jpg (garder le nom original sans l'extension)
            new_filename = Path(original_filename).stem + '.jpg'
            
            dest = UPLOAD_DIR / new_filename

            # Sauvegarder en JPEG avec bonne qualité
            heic_image.save(dest, 'JPEG', quality=95, optimize=True)

            web_path = f"/uploads/{new_filename}"
            print(f"[UPLOAD] HEIC converti: {original_filename} -> {new_filename}")
            return JSONResponse({"path": web_path})
        except Exception as e:
            print(f"[UPLOAD ERROR] Erreur conversion HEIC {original_filename}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Erreur conversion HEIC: {str(e)}")
    else:
        # FORCER la re-conversion de TOUS les fichiers image pour garantir la compatibilité avec python-docx
        file_size_mb = len(contents) / (1024 * 1024)

        try:
            # Ouvrir l'image
            img = Image.open(io.BytesIO(contents))
            original_format = img.format

            print(f"[UPLOAD] Traitement: {original_filename} ({file_size_mb:.2f} MB, format={original_format}, mode={img.mode})")

            # TOUJOURS convertir en JPEG avec paramètres standard pour compatibilité python-docx
            # Cela résout les problèmes de JPG mal encodés

            # Redimensionner si l'image est trop grande (max 5000px pour gérer jusqu'Ã  15-20 MB)
            max_dimension = 5000
            needs_resize = img.width > max_dimension or img.height > max_dimension

            if needs_resize:
                ratio = min(max_dimension / img.width, max_dimension / img.height)
                new_size = (int(img.width * ratio), int(img.height * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
                print(f"[UPLOAD] Image redimensionnée: {img.size}")

            # Convertir en RGB si nécessaire (obligatoire pour JPEG)
            if img.mode in ('RGBA', 'LA', 'P'):
                rgb_image = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                if img.mode == 'RGBA':
                    rgb_image.paste(img, mask=img.split()[-1])
                else:
                    rgb_image.paste(img)
                img = rgb_image
                print(f"[UPLOAD] Converti de {img.mode} vers RGB")
            elif img.mode != 'RGB':
                img = img.convert('RGB')
                print(f"[UPLOAD] Converti de {img.mode} vers RGB")

            # Toujours sauvegarder en JPEG avec paramètres compatibles python-docx
            # subsampling=0 (4:4:4) évite les problèmes de compression
            new_filename = Path(original_filename).stem + '.jpg'
            dest = UPLOAD_DIR / new_filename

            # Ajuster la qualité selon la taille
            if file_size_mb > 10:
                quality = 80
            elif file_size_mb > 5:
                quality = 85
            else:
                quality = 90

            img.save(dest, 'JPEG', quality=quality, optimize=True, subsampling=0)

            new_size_mb = dest.stat().st_size / (1024 * 1024)

            # Vérifier si le fichier final est trop grand
            if new_size_mb > 10:
                print(f"[UPLOAD WARNING] Fichier encore trop grand ({new_size_mb:.2f} MB), nouvelle compression...")
                # Réduire encore la qualité
                img.save(dest, 'JPEG', quality=75, optimize=True, subsampling=2)
                new_size_mb = dest.stat().st_size / (1024 * 1024)

            print(f"[UPLOAD] Image traitee: {original_filename} -> {new_filename} ({file_size_mb:.2f} MB -> {new_size_mb:.2f} MB)")

            web_path = f"/uploads/{new_filename}"
            return JSONResponse({"path": web_path})

        except Exception as e:
            # En cas d'erreur, essayer une conversion basique
            print(f"[UPLOAD ERROR] Erreur traitement {original_filename}: {type(e).__name__}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Impossible de traiter l'image: {str(e)}")


# Feuilles Ã  ignorer lors de l'import Excel (inutiles pour l'analyse forensique)
SKIP_SHEETS = [
    'Summary',
    'Image Hashes',
    'Aggregated Application Usage',
    'Applications Usage Log',
    'Applications',
    'Exchange',
    'File Uploads',
    'Configurations',
    'Databases',
    'Videos',
    'Video',
    'Autofill',
    'Log Entries',
    'Social Media',
    'Text',
    'Locations View',
    'Device Notifications',
    'Cookies',
    'Shortcut',
    'Watch List Results',
]

# Couleurs assignées pour les plateformes connues
SOURCE_COLORS_MAP = {
    'Snapchat': '#FFFC00',
    'WhatsApp': '#25D366',
    'Instagram': '#C13584',
    'Facebook': '#1877F2',
    'Facebook Messenger': '#00BFFF',
    'Telegram': '#0088CC',
    'iMessage': '#34C759',
    'SMS': '#5856D6',
    'Signal': '#3A76F0',
    'Messenger': '#0084FF',
    'Natif': '#B0B0B0',
    'TikTok': '#69C9D0',
    'Twitter': '#1DA1F2',
    'LinkedIn': '#0A66C2',
    'Discord': '#5865F2',
    'Viber': '#665CAC',
    'Line': '#00B900',
    'WeChat': '#07C160',
}

# Palette de couleurs pour les sources non assignées (évite le gris)
FALLBACK_COLORS = [
    '#E74C3C',  # Rouge
    '#9B59B6',  # Violet
    '#3498DB',  # Bleu clair
    '#1ABC9C',  # Turquoise
    '#F39C12',  # Orange
    '#2ECC71',  # Vert
    '#E91E63',  # Rose
    '#00BCD4',  # Cyan
    '#FF5722',  # Orange foncé
    '#8BC34A',  # Vert clair
    '#673AB7',  # Violet foncé
    '#FF9800',  # Ambre
    '#795548',  # Marron
    '#607D8B',  # Bleu gris
    '#CDDC39',  # Lime
]


def get_chart_colors(sources: list, colors_map: dict = None) -> list:
    """
    Retourne une liste de couleurs pour les sources.
    Les sources connues utilisent leurs couleurs assignées.
    Les sources inconnues reçoivent des couleurs uniques de la palette fallback.
    """
    if colors_map is None:
        colors_map = SOURCE_COLORS_MAP

    colors = []
    fallback_index = 0
    used_fallback_colors = set()

    for src in sources:
        if src in colors_map:
            colors.append(colors_map[src])
        else:
            # Assigner une couleur de la palette fallback
            color = FALLBACK_COLORS[fallback_index % len(FALLBACK_COLORS)]
            colors.append(color)
            fallback_index += 1

    return colors


# Mapping Excel Name -> Word placeholder
EXCEL_TO_WORD_MAPPING = {
    # Device Info mappings (Name column -> placeholder)
    "IMEI": "{imei1}",
    "IMEI1": "{imei1}",
    "IMEI2": "{imei2}",
    "Vendor": "{marque}",
    "Detected Phone Model": "{modele}",
    "Model": "{modele}",  # fallback
    "Serial": "{numserie}",
    "OS Version": "{vers}",
    "OS": "{vers}",  # fallback
}

# Sheets to count for statistics
COUNT_SHEETS = {
    "Contacts": "{nbcontact}",
    "Call Log": "{nbappel}",
    "Chats": "{nbmessage}",
    "Document": "{nbdocument}",
    "Images": "{nbimage}",
    "Videos": "{nbvideo}",
    "Video": "{nbvideo}",
    "Audio": "{nbaudio}",
}

# Sources Ã  exclure pour les contacts (filtrées lors du comptage)
EXCLUDED_CONTACT_SOURCES = {'Recents', 'InteractionC', 'KnowledgeC', 'Native Messages', 'Threads', 'Biome', 'Unified Logs', 'SIM'}


# ==============================================================================
# SYSTÈME PARQUET + POLARS - Conversion et lecture ultra-rapide
# ==============================================================================

def get_excel_hash(content: bytes) -> str:
    """Génère un hash unique pour identifier le fichier Excel"""
    return hashlib.md5(content[:50000]).hexdigest()[:16]


def get_parquet_path(excel_hash: str, sheet_name: str) -> Path:
    """Chemin du fichier Parquet pour une feuille donnée"""
    safe_name = re.sub(r'[^\w\-]', '_', sheet_name)
    return PARQUET_CACHE_DIR / f"{excel_hash}_{safe_name}.parquet"


def clear_parquet_cache(excel_hash: str = None):
    """Nettoie le cache Parquet (tout ou pour un fichier spécifique)"""
    global _parquet_dataframes
    if excel_hash:
        # Supprimer uniquement les fichiers de cet Excel
        for f in PARQUET_CACHE_DIR.glob(f"{excel_hash}_*.parquet"):
            f.unlink()
        _parquet_dataframes = {k: v for k, v in _parquet_dataframes.items() if not k.startswith(excel_hash)}
    else:
        # Tout supprimer
        for f in PARQUET_CACHE_DIR.glob("*.parquet"):
            f.unlink()
        _parquet_dataframes.clear()


def parse_sheet_to_parquet(content: bytes, sheet_name: str, excel_hash: str) -> Optional[Path]:
    """
    Parse une feuille Excel via XML et l'exporte en Parquet.
    Retourne le chemin du fichier Parquet créé.
    """
    parquet_path = get_parquet_path(excel_hash, sheet_name)

    # Si déjÃ  en cache, ne pas recréer
    if parquet_path.exists():
        return parquet_path

    ns_main = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver l'ID de la feuille
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            wb_root = ET.fromstring(workbook_xml)

            sheet_id = None
            for sheet in wb_root.iter(f'{{{ns_main}}}sheet'):
                if sheet.get('name') == sheet_name:
                    sheet_id = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    break

            if not sheet_id:
                return None

            # 2. Trouver le fichier XML de la feuille
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            rels_root = ET.fromstring(rels_xml)

            sheet_file = None
            for rel in rels_root.iter():
                if rel.get('Id') == sheet_id:
                    target = rel.get('Target', '')
                    if target.startswith('/'):
                        sheet_file = target[1:]
                    elif target.startswith('xl/'):
                        sheet_file = target
                    else:
                        sheet_file = 'xl/' + target
                    break

            if not sheet_file:
                return None

            # 3. Charger sharedStrings
            shared_strings = []
            try:
                ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                ss_root = ET.fromstring(ss_xml)
                for si in ss_root.iter(f'{{{ns_main}}}si'):
                    text_parts = []
                    for t in si.iter(f'{{{ns_main}}}t'):
                        if t.text:
                            text_parts.append(t.text)
                    shared_strings.append(''.join(text_parts))
            except:
                pass

            # 4. Lire la feuille XML
            sheet_xml = zf.read(sheet_file).decode('utf-8')
            sheet_root = ET.fromstring(sheet_xml)

            # 5. Extraire les en-têtes (ligne 2)
            headers = {}  # {col_letter: col_name}
            for row in sheet_root.iter(f'{{{ns_main}}}row'):
                if row.get('r') == '2':
                    for cell in row.iter(f'{{{ns_main}}}c'):
                        cell_ref = cell.get('r', '')
                        col_letter = ''.join(c for c in cell_ref if c.isalpha())

                        v_elem = cell.find(f'{{{ns_main}}}v')
                        if v_elem is not None and v_elem.text:
                            cell_type = cell.get('t', '')
                            if cell_type == 's':
                                idx = int(v_elem.text)
                                if idx < len(shared_strings):
                                    headers[col_letter] = clean_excel_value_fast(shared_strings[idx])
                            else:
                                headers[col_letter] = clean_excel_value_fast(v_elem.text)
                    break

            if not headers:
                return None

            # 6. Extraire toutes les données (lignes 3+)
            all_data = {col_name: [] for col_name in headers.values()}
            letter_to_name = headers

            for row in sheet_root.iter(f'{{{ns_main}}}row'):
                row_num = int(row.get('r', 0))
                if row_num <= 2:
                    continue

                row_values = {col: None for col in headers.values()}

                for cell in row.iter(f'{{{ns_main}}}c'):
                    cell_ref = cell.get('r', '')
                    col_letter = ''.join(c for c in cell_ref if c.isalpha())

                    if col_letter not in letter_to_name:
                        continue

                    col_name = letter_to_name[col_letter]
                    v_elem = cell.find(f'{{{ns_main}}}v')

                    if v_elem is not None and v_elem.text:
                        cell_type = cell.get('t', '')
                        if cell_type == 's':
                            idx = int(v_elem.text)
                            if idx < len(shared_strings):
                                row_values[col_name] = clean_excel_value_fast(shared_strings[idx])
                        else:
                            row_values[col_name] = clean_excel_value_fast(v_elem.text)

                for col_name, value in row_values.items():
                    all_data[col_name].append(value)

            # 7. Créer DataFrame Polars et exporter en Parquet
            if any(len(v) > 0 for v in all_data.values()):
                df = pl.DataFrame(all_data)
                df.write_parquet(parquet_path, compression="zstd", compression_level=3)
                print(f"[PARQUET] Créé: {parquet_path.name} ({len(df)} lignes)")
                return parquet_path

    except Exception as e:
        print(f"[PARQUET ERROR] {sheet_name}: {e}")

    return None


def clean_excel_value_fast(value: str) -> str:
    """Nettoie les caractères spéciaux XML des valeurs Excel"""
    if not value:
        return ""
    # Supprimer _x000D_, _x000A_, etc.
    cleaned = re.sub(r'_x[0-9A-Fa-f]{4}_', '', value)
    cleaned = cleaned.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def get_sheet_dataframe(content: bytes, sheet_name: str, excel_hash: str = None) -> Optional[pl.DataFrame]:
    """
    Récupère un DataFrame Polars pour une feuille.
    Utilise le cache mémoire > cache Parquet > parsing XML.
    """
    global _parquet_dataframes

    if excel_hash is None:
        excel_hash = get_excel_hash(content)

    cache_key = f"{excel_hash}_{sheet_name}"

    # 1. Cache mémoire (instantané)
    if cache_key in _parquet_dataframes:
        return _parquet_dataframes[cache_key]

    # 2. Cache Parquet (très rapide)
    parquet_path = get_parquet_path(excel_hash, sheet_name)
    if parquet_path.exists():
        df = pl.read_parquet(parquet_path)
        _parquet_dataframes[cache_key] = df
        return df

    # 3. Parsing XML et création Parquet (première fois seulement)
    parquet_path = parse_sheet_to_parquet(content, sheet_name, excel_hash)
    if parquet_path and parquet_path.exists():
        df = pl.read_parquet(parquet_path)
        _parquet_dataframes[cache_key] = df
        return df

    return None


def convert_all_sheets_to_parquet(content: bytes) -> str:
    """
    Convertit toutes les feuilles utiles d'un Excel en Parquet.
    Retourne le hash du fichier Excel.
    """
    global _current_excel_hash, _parquet_dataframes

    excel_hash = get_excel_hash(content)

    # Si c'est un nouveau fichier, nettoyer l'ancien cache mémoire
    if excel_hash != _current_excel_hash:
        _parquet_dataframes.clear()
        _current_excel_hash = excel_hash

    # Récupérer la liste des feuilles
    sheets = get_all_sheets_names(content)

    # Convertir chaque feuille en Parquet (en parallèle si possible)
    for sheet_name in sheets:
        if sheet_name not in SKIP_SHEETS:
            parse_sheet_to_parquet(content, sheet_name, excel_hash)

    return excel_hash


def get_all_sheets_names(content: bytes) -> List[str]:
    """Liste toutes les feuilles d'un Excel (sans filtrer SKIP_SHEETS)"""
    sheets = []
    ns_main = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            wb_root = ET.fromstring(workbook_xml)

            for sheet in wb_root.iter(f'{{{ns_main}}}sheet'):
                sheet_name = sheet.get('name')
                if sheet_name:
                    sheets.append(sheet_name)
    except:
        pass

    return sheets


# ==============================================================================
# FONCTIONS POLARS - Statistiques et extraction ultra-rapides
# ==============================================================================

def count_sheet_rows_polars(content: bytes, sheet_name: str, excel_hash: str = None) -> int:
    """Compte les lignes d'une feuille via Polars (instantané après conversion)"""
    df = get_sheet_dataframe(content, sheet_name, excel_hash)
    return len(df) if df is not None else 0


def get_contacts_stats_polars(content: bytes, excel_hash: str = None) -> Tuple[int, Dict[str, int]]:
    """
    Récupère les stats des contacts via Polars:
    - Nombre total (excluant les sources filtrées)
    - Comptage par source
    """
    df = get_sheet_dataframe(content, "Contacts", excel_hash)

    if df is None or "Source" not in df.columns:
        return 0, {}

    # Remplacer les valeurs null par "Natif"
    df = df.with_columns(
        pl.col("Source").fill_null("Natif").alias("Source")
    )

    # Filtrer les sources exclues
    df_filtered = df.filter(~pl.col("Source").is_in(list(EXCLUDED_CONTACT_SOURCES)))

    # Comptage par source
    source_counts = (
        df_filtered
        .group_by("Source")
        .agg(pl.count().alias("count"))
        .sort("count", descending=True)
    )

    counts_dict = {
        str(row["Source"]): row["count"]
        for row in source_counts.to_dicts()
    }

    total = sum(counts_dict.values())
    return total, counts_dict


def get_calls_stats_polars(content: bytes, excel_hash: str = None) -> Tuple[int, Dict[str, int]]:
    """Récupère les stats des appels via Polars"""
    df = get_sheet_dataframe(content, "Call Log", excel_hash)

    if df is None:
        return 0, {}

    total = len(df)

    # Chercher une colonne Source/Application
    source_col = None
    for col in ["Source", "Application", "App"]:
        if col in df.columns:
            source_col = col
            break

    if source_col:
        # Remplacer les valeurs null par "Natif"
        df = df.with_columns(
            pl.col(source_col).fill_null("Natif").alias(source_col)
        )
        source_counts = (
            df
            .group_by(source_col)
            .agg(pl.count().alias("count"))
            .sort("count", descending=True)
        )
        counts_dict = {
            str(row[source_col]): row["count"]
            for row in source_counts.to_dicts()
        }
    else:
        counts_dict = {"Natif": total}

    return total, counts_dict


def get_chats_stats_polars(content: bytes, excel_hash: str = None) -> Tuple[int, int, Dict[str, int]]:
    """
    Récupère les stats des chats via Polars:
    - Nombre de messages
    - Nombre de conversations
    - Comptage par source/application
    """
    df = get_sheet_dataframe(content, "Chats", excel_hash)

    if df is None:
        return 0, 0, {}

    nb_messages = len(df)

    # Nombre de conversations (colonne Chat # ou similaire)
    nb_conversations = 0
    for col in ["Chat #", "Chat", "Conversation"]:
        if col in df.columns:
            # La dernière valeur non-nulle représente le nombre de conversations
            try:
                nb_conversations = df[col].drop_nulls().cast(pl.Int64, strict=False).max()
                if nb_conversations is None:
                    nb_conversations = df[col].n_unique()
            except:
                nb_conversations = df[col].n_unique()
            break

    # Comptage par source/application
    source_col = None
    for col in ["Source", "Application", "App", "Platform"]:
        if col in df.columns:
            source_col = col
            break

    if source_col:
        # Remplacer les valeurs null par "Natif"
        df = df.with_columns(
            pl.col(source_col).fill_null("Natif").alias(source_col)
        )
        source_counts = (
            df
            .group_by(source_col)
            .agg(pl.count().alias("count"))
            .sort("count", descending=True)
        )
        counts_dict = {
            str(row[source_col]): row["count"]
            for row in source_counts.to_dicts()
        }
    else:
        counts_dict = {"Messages": nb_messages}

    return nb_messages, int(nb_conversations) if nb_conversations else 0, counts_dict


def get_sheet_data_polars(
    content: bytes,
    sheet_name: str,
    selected_columns: List[str],
    max_rows: int = 100,
    source_filter: Optional[str] = None,
    excel_hash: str = None
) -> List[Dict[str, str]]:
    """
    Récupère les données d'une feuille via Polars - ULTRA RAPIDE.
    """
    df = get_sheet_dataframe(content, sheet_name, excel_hash)

    if df is None:
        return []

    # Sélectionner uniquement les colonnes demandées qui existent
    available_cols = [col for col in selected_columns if col in df.columns]
    if not available_cols:
        return []

    result = df.select(available_cols)

    # Filtrer par source si demandé
    if source_filter and "Source" in result.columns:
        result = result.filter(pl.col("Source") == source_filter)

    # Limiter le nombre de lignes
    result = result.head(max_rows)

    # Convertir en liste de dicts avec valeurs string
    # Nettoyer le HTML de la colonne Body pour les Emails
    is_emails = "email" in sheet_name.lower()
    data = []
    for row in result.to_dicts():
        row_dict = {}
        for k, v in row.items():
            val = str(v) if v is not None else ""
            # Nettoyer HTML pour colonne Body des Emails
            if is_emails and k == "Body" and val:
                val = strip_html_tags(val)
            row_dict[k] = val
        data.append(row_dict)

    return data


def extract_device_info_fast(content: bytes) -> Dict[str, str]:
    """
    Extraction ULTRA RAPIDE: lit directement le XML de Device Info
    sans parser tout le fichier Excel (bypass openpyxl)
    """
    import zipfile
    import xml.etree.ElementTree as ET

    device_data = {}

    with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
        # 1. Lire workbook.xml pour trouver l'ID de "Device Info"
        workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
        wb_root = ET.fromstring(workbook_xml)
        ns = {'main': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main',
              'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}

        sheet_id = None
        for sheet in wb_root.findall('.//main:sheet', ns):
            if sheet.get('name') == 'Device Info':
                sheet_id = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                break

        if not sheet_id:
            raise ValueError("Feuille 'Device Info' introuvable")

        # 2. Lire les relations pour trouver le fichier XML de la feuille
        rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
        rels_root = ET.fromstring(rels_xml)

        sheet_file = None
        for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            if rel.get('Id') == sheet_id:
                sheet_file = 'xl/' + rel.get('Target')
                break

        if not sheet_file:
            raise ValueError("Fichier de la feuille introuvable")

        # 3. Lire sharedStrings.xml (textes partagés)
        shared_strings = []
        try:
            ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
            ss_root = ET.fromstring(ss_xml)
            for si in ss_root.findall('.//main:si', ns):
                t_elem = si.find('.//main:t', ns)
                shared_strings.append(t_elem.text if t_elem is not None and t_elem.text else "")
        except:
            pass  # Pas de shared strings

        # 4. Lire la feuille Device Info
        sheet_xml = zf.read(sheet_file).decode('utf-8')
        sheet_root = ET.fromstring(sheet_xml)

        # Parser les lignes (on cherche colonnes C et D, lignes 3+)
        for row in sheet_root.findall('.//main:row', ns):
            row_num = int(row.get('r', 0))
            if row_num < 3:
                continue
            if row_num > 100:
                break

            cells = {}
            for cell in row.findall('main:c', ns):
                cell_ref = cell.get('r', '')
                col = ''.join(c for c in cell_ref if c.isalpha())

                if col not in ('C', 'D'):
                    continue

                cell_type = cell.get('t', '')
                v_elem = cell.find('main:v', ns)

                if v_elem is not None and v_elem.text:
                    if cell_type == 's':  # Shared string
                        idx = int(v_elem.text)
                        cells[col] = shared_strings[idx] if idx < len(shared_strings) else ""
                    else:
                        cells[col] = v_elem.text

            # Si on a C et D, ajouter au dictionnaire
            if 'C' in cells and 'D' in cells:
                name = cells['C'].strip()
                value = cells['D'].strip()
                if name and value:
                    device_data[name] = value

    return device_data


def get_chats_stats_fast(content: bytes) -> Tuple[int, int]:
    """
    Récupère RAPIDEMENT:
    - Le nombre de messages (lignes - 2)
    - Le nombre de conversations (dernière valeur de la colonne Chat # / colonne B)
    Retourne (nb_messages, nb_conversations)
    """
    import zipfile
    import re

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver le fichier de la feuille Chats via regex (plus rapide que ElementTree)
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            match = re.search(r'<sheet[^>]*name="Chats"[^>]*r:id="(rId\d+)"', workbook_xml)
            if not match:
                match = re.search(r'<sheet[^>]*r:id="(rId\d+)"[^>]*name="Chats"', workbook_xml)
            if not match:
                return (0, 0)

            sheet_rid = match.group(1)

            # 2. Trouver le fichier XML via relations
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            match = re.search(rf'<Relationship[^>]*Id="{sheet_rid}"[^>]*Target="([^"]+)"', rels_xml)
            if not match:
                return (0, 0)

            target = match.group(1)
            sheet_file = 'xl/' + target if not target.startswith('xl/') else target

            # 3. Charger sharedStrings avec regex optimisé
            shared_strings = []
            try:
                ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                # Extraire chaque <si>...</si> comme une entrée
                si_pattern = r'<si[^>]*>(.*?)</si>'
                for si_match in re.finditer(si_pattern, ss_xml, re.DOTALL):
                    si_content = si_match.group(1)
                    texts = re.findall(r'<t[^>]*>([^<]*)</t>', si_content)
                    shared_strings.append(''.join(texts))
            except:
                pass

            # 4. Lire le XML de la feuille
            sheet_xml = zf.read(sheet_file).decode('utf-8')

            # Compter les lignes (rapide)
            row_count = sheet_xml.count('<row ')
            nb_messages = max(0, row_count - 2)

            # 5. Trouver la dernière valeur de colonne B
            row_pattern = r'<row r="(\d+)"[^>]*>(.*?)</row>'
            rows = list(re.finditer(row_pattern, sheet_xml, re.DOTALL))

            last_chat_value = 0
            for row_match in reversed(rows[-50:]):  # 50 dernières lignes max
                row_num = int(row_match.group(1))
                if row_num <= 2:
                    continue

                row_content = row_match.group(2)

                # Chercher cellule B avec ses attributs et valeur
                cell_match = re.search(r'<c r="B\d+"([^>]*)><v>(\d+)</v></c>', row_content)
                if cell_match:
                    attrs = cell_match.group(1)
                    v_val = int(cell_match.group(2))

                    if 't="s"' in attrs:
                        # C'est un index dans sharedStrings
                        if v_val < len(shared_strings):
                            try:
                                last_chat_value = int(shared_strings[v_val])
                            except:
                                pass
                    else:
                        # Valeur numérique directe
                        last_chat_value = v_val
                    break

            print(f"[CHATS] Messages={nb_messages}, Conversations={last_chat_value}")
            return (nb_messages, last_chat_value)
    except Exception as e:
        print(f"[CHATS STATS] Erreur: {e}")
        return (0, 0)


def count_sheet_rows_fast(content: bytes, sheet_name: str) -> int:
    """
    Compte le nombre de lignes dans une feuille Excel (sans l'en-tête)
    """
    import zipfile
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
        # 1. Lire workbook.xml pour trouver l'ID de la feuille
        workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
        wb_root = ET.fromstring(workbook_xml)
        ns = {'main': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main',
              'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}

        sheet_id = None
        for sheet in wb_root.findall('.//main:sheet', ns):
            if sheet.get('name') == sheet_name:
                sheet_id = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                break

        if not sheet_id:
            return 0  # Feuille non trouvée

        # 2. Lire les relations pour trouver le fichier XML de la feuille
        rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
        rels_root = ET.fromstring(rels_xml)

        sheet_file = None
        for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            if rel.get('Id') == sheet_id:
                sheet_file = 'xl/' + rel.get('Target')
                break

        if not sheet_file:
            return 0

        # 3. Lire la feuille et compter les lignes
        sheet_xml = zf.read(sheet_file).decode('utf-8')
        sheet_root = ET.fromstring(sheet_xml)

        # Compter toutes les lignes (éléments row) - 2 pour les en-têtes
        rows = sheet_root.findall('.//main:row', ns)
        row_count = len(rows)

        # Soustraire 2 pour les en-têtes (lignes 1 et 2)
        return max(0, row_count - 2)


def get_contacts_by_source_fast(content: bytes) -> Dict[str, int]:
    """
    Compte les contacts par source via lecture XML directe (ultra rapide).
    Retourne un dictionnaire {source: count}.
    """
    import zipfile
    import re
    from collections import Counter

    # Sources Ã  exclure
    EXCLUDED_SOURCES = {'Recents', 'InteractionC', 'KnowledgeC', 'Native Messages', 'Threads', 'Biome', 'Unified Logs', 'SIM'}

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver la feuille Contacts
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            match = re.search(r'<sheet[^>]*name="Contacts"[^>]*r:id="(rId\d+)"', workbook_xml)
            if not match:
                match = re.search(r'<sheet[^>]*r:id="(rId\d+)"[^>]*name="Contacts"', workbook_xml)
            if not match:
                return {}

            sheet_rid = match.group(1)

            # 2. Trouver le fichier XML
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            match = re.search(rf'<Relationship[^>]*Id="{sheet_rid}"[^>]*Target="([^"]+)"', rels_xml)
            if not match:
                return {}

            target = match.group(1)
            sheet_file = 'xl/' + target if not target.startswith('xl/') else target

            # 3. Charger sharedStrings
            shared_strings = []
            try:
                ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                si_pattern = r'<si[^>]*>(.*?)</si>'
                for si_match in re.finditer(si_pattern, ss_xml, re.DOTALL):
                    si_content = si_match.group(1)
                    texts = re.findall(r'<t[^>]*>([^<]*)</t>', si_content)
                    shared_strings.append(''.join(texts))
            except:
                pass

            # 4. Trouver la colonne Source (chercher dans ligne 2)
            sheet_xml = zf.read(sheet_file).decode('utf-8')

            # Trouver l'index de "Source" dans sharedStrings
            source_str_idx = None
            for idx, s in enumerate(shared_strings):
                if s == 'Source':
                    source_str_idx = idx
                    break

            if source_str_idx is None:
                print("[CHART] 'Source' non trouvé dans sharedStrings")
                return {}

            # Trouver quelle colonne contient "Source" dans la ligne 2
            row2_match = re.search(r'<row r="2"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL)
            if not row2_match:
                return {}

            source_col = None
            row2_content = row2_match.group(1)
            cell_pattern = r'<c r="([A-Z]+)2"[^>]*t="s"[^>]*><v>(\d+)</v></c>'
            for cell_match in re.finditer(cell_pattern, row2_content):
                col = cell_match.group(1)
                str_idx = int(cell_match.group(2))
                if str_idx == source_str_idx:
                    source_col = col
                    break

            if not source_col:
                print("[CHART] Colonne Source non trouvée")
                return {}

            # 5. Extraire toutes les valeurs de la colonne Source
            source_counts = Counter()
            row_pattern = r'<row r="(\d+)"[^>]*>(.*?)</row>'

            for row_match in re.finditer(row_pattern, sheet_xml, re.DOTALL):
                row_num = int(row_match.group(1))
                if row_num <= 2:  # Skip header
                    continue

                row_content = row_match.group(2)

                # Chercher la cellule de la colonne Source (plusieurs formats possibles)
                # Format 1: <c r="X123" ...><v>val</v></c>
                # Format 2: <c r="X123" .../> (auto-fermante, pas de valeur)
                # Format 3: <c r="X123" ...></c> (vide)
                cell_pattern = rf'<c r="{source_col}{row_num}"([^>]*)(?:><v>([^<]*)</v></c>|></c>|/>)'
                cell_match = re.search(cell_pattern, row_content)

                source_value = "Natif"  # Valeur par défaut

                if cell_match:
                    attrs = cell_match.group(1) or ""
                    v_val = cell_match.group(2)  # Peut être None si cellule vide

                    if v_val is not None and v_val != "":
                        if 't="s"' in attrs:
                            str_idx = int(v_val)
                            if str_idx < len(shared_strings):
                                val = shared_strings[str_idx]
                                if val and val.strip():
                                    source_value = val.strip()
                        else:
                            if v_val.strip():
                                source_value = v_val.strip()

                # Filtrer les sources exclues
                if source_value not in EXCLUDED_SOURCES:
                    source_counts[source_value] += 1

            return dict(source_counts)

    except Exception as e:
        print(f"[CHART] Erreur lecture sources: {e}")
        return {}


def _generate_vertical_chart(
    labels: List[str],
    values: List[float],
    title: str,
    filename_prefix: str,
    colors: Optional[List[str]] = None,
    y_label: Optional[str] = None,
    value_labels: Optional[List[str]] = None,
    x_label: Optional[str] = None,
) -> Optional[str]:
    # Lock pour éviter les conflits matplotlib en parallèle
    with _chart_lock:
        try:
            if not labels or not values:
                return None

            plt.style.use('seaborn-v0_8-whitegrid')

            # Configurer une police qui supporte les émojis (Windows: Segoe UI Emoji)
            import platform
            if platform.system() == "Windows":
                plt.rcParams['font.family'] = ['Segoe UI Emoji', 'Segoe UI', 'DejaVu Sans', 'sans-serif']
            else:
                plt.rcParams['font.family'] = ['Noto Color Emoji', 'DejaVu Sans', 'sans-serif']

            n = len(labels)

            # Tailles de police ADAPTATIVES selon le nombre d'éléments
            # Plus il y a d'éléments (TOP 15), plus on augmente les tailles pour compenser le redimensionnement Word
            if n >= 12:  # TOP 15 ou similaire
                fig_width = min(72, max(56, n * 4.0))
                fig_height = 36
                fontsize_values = 64       # Valeurs au-dessus des barres
                fontsize_labels = 56       # Identifiants sur l'axe X - ZOOM
                fontsize_yticks = 48       # Graduations axe Y
                fontsize_ylabel = 50       # Label axe Y
                fontsize_title = 68        # Titre
                bottom_margin = 0.38       # Plus de marge pour les labels multi-ligne
                bar_width = 0.60           # Barres plus fines pour meilleur espacement
            elif n >= 8:  # Graphiques moyens
                fig_width = min(66, max(50, n * 3.5))
                fig_height = 32
                fontsize_values = 58
                fontsize_labels = 50       # ZOOM identifiants
                fontsize_yticks = 44
                fontsize_ylabel = 46
                fontsize_title = 62
                bottom_margin = 0.35
                bar_width = 0.62
            else:  # Petits graphiques (< 8 éléments)
                fig_width = min(58, max(42, n * 3.2))
                fig_height = 30
                fontsize_values = 54
                fontsize_labels = 46       # ZOOM identifiants
                fontsize_yticks = 40
                fontsize_ylabel = 42
                fontsize_title = 58
                bottom_margin = 0.32
                bar_width = 0.65

            fig, ax = plt.subplots(figsize=(fig_width, fig_height))
            ax.set_facecolor('#FAFBFC')
            fig.patch.set_facecolor('#FFFFFF')

            palette = ['#4299E1', '#3BB0D1', '#957BBB', '#F8A031', '#45B63B',
                       '#F4C433', '#EE3B4B', '#42B7C1', '#9D46AF', '#3FA6BF']
            bar_colors = colors if colors and len(colors) == n else [palette[i % len(palette)] for i in range(n)]

            x_pos = range(n)
            bars = ax.bar(x_pos, values, width=bar_width, color=bar_colors, edgecolor='white', linewidth=2, zorder=2)

            max_val = max(values) if values else 1
            # Offset constant au-dessus de chaque barre
            label_offset = max_val * 0.04

            for idx, bar in enumerate(bars):
                h = bar.get_height()

                if value_labels and idx < len(value_labels):
                    # Graphique de DURÉE: utiliser la hauteur réelle + offset
                    txt = value_labels[idx]
                    y_pos = h + label_offset
                else:
                    # Graphique de COUNT: utiliser la valeur ARRONDIE pour l'alignement
                    # Ainsi tous les "3" seront à la même hauteur Y
                    rounded_h = int(round(h))
                    txt = str(rounded_h)
                    y_pos = rounded_h + label_offset

                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    y_pos,
                    txt,
                    ha='center',
                    va='bottom',
                    fontsize=fontsize_values,
                    fontweight='bold',
                    color='#000000',
                    zorder=3,
                )

            ax.set_xticks(list(x_pos))
            # Décomposer les identifiants longs (UUIDs, etc.) sur plusieurs lignes
            def split_long_label(label, max_chars=25):
                """Coupe les identifiants longs en morceaux pour meilleure lisibilité"""
                if len(label) <= max_chars:
                    return label
                # Couper en morceaux de max_chars caractères
                parts = []
                for i in range(0, len(label), max_chars):
                    parts.append(label[i:i+max_chars])
                return '\n'.join(parts)

            formatted_labels = [split_long_label(lbl) for lbl in labels]
            # IDENTIFIANTS ZOOMÉS et bien espacés pour lisibilité dans Word
            ax.set_xticklabels(formatted_labels, rotation=40, ha='right', fontsize=fontsize_labels, fontweight='bold', color='#000000')
            ax.tick_params(axis='y', labelsize=fontsize_yticks, colors='#000000', width=2)
            for label in ax.get_yticklabels():
                label.set_fontweight('bold')
            ax.set_ylim(0, max_val * 1.40)  # Plus d'espace pour les valeurs au-dessus des barres

            ax.yaxis.grid(True, linestyle='--', alpha=0.3, color='#CBD5E0')
            ax.xaxis.grid(False)
            for spine in ax.spines.values():
                spine.set_visible(False)

            if y_label:
                ax.set_ylabel(y_label, fontsize=fontsize_ylabel, fontweight='bold', color='#000000', labelpad=14)

            if x_label:
                ax.set_xlabel(x_label, fontsize=fontsize_ylabel, fontweight='bold', color='#000000', labelpad=14)

            ax.set_title(title, fontsize=fontsize_title, fontweight='bold', color='#000000', pad=35, loc='left')

            plt.tight_layout()
            # Espace généreux en bas pour les identifiants ZOOMÉS (adapté au nombre d'éléments)
            plt.subplots_adjust(left=0.06, right=0.97, top=0.92, bottom=bottom_margin)

            chart_filename = f"{filename_prefix}_{uuid.uuid4().hex[:8]}.png"
            chart_path = UPLOAD_DIR / chart_filename
            plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white', edgecolor='none', pad_inches=0.4)
            plt.close(fig)
            return f"/uploads/{chart_filename}"
        except Exception:
            plt.close('all')  # Fermer toutes les figures en cas d'erreur
            return None


def generate_contacts_chart(source_counts: Dict[str, int]) -> Optional[str]:
    try:
        if not source_counts:
            print("[CHART] Aucune donnée source")
            return None

        sorted_sources = sorted(source_counts.items(), key=lambda x: x[1], reverse=True)
        sources = [s[0] for s in sorted_sources]
        counts = [s[1] for s in sorted_sources]

        colors = get_chart_colors(sources, SOURCE_COLORS_MAP)
        return _generate_vertical_chart(
            labels=sources,
            values=counts,
            title='Volume de contacts par plateforme',
            filename_prefix='chart_contacts',
            colors=colors,
            y_label='Nombre de contacts',
        )
    except Exception as e:
        print(f"[CHART ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def get_calls_by_source_fast(content: bytes) -> Dict[str, int]:
    """
    Compte les appels par source via lecture XML directe (ultra rapide).
    Retourne un dictionnaire {source: count}.
    """
    import zipfile
    import re
    from collections import Counter

    # Sources Ã  exclure (mêmes que contacts)
    EXCLUDED_SOURCES = {'Recents', 'InteractionC', 'KnowledgeC', 'Native Messages', 'Threads', 'Biome', 'Unified Logs', 'SIM'}

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver la feuille Call Log
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            match = re.search(r'<sheet[^>]*name="Call Log"[^>]*r:id="(rId\d+)"', workbook_xml)
            if not match:
                match = re.search(r'<sheet[^>]*r:id="(rId\d+)"[^>]*name="Call Log"', workbook_xml)
            if not match:
                print("[CALLS CHART] Feuille 'Call Log' non trouvée")
                return {}

            sheet_rid = match.group(1)

            # 2. Trouver le fichier XML
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            match = re.search(rf'<Relationship[^>]*Id="{sheet_rid}"[^>]*Target="([^"]+)"', rels_xml)
            if not match:
                return {}

            target = match.group(1)
            sheet_file = 'xl/' + target if not target.startswith('xl/') else target

            # 3. Charger sharedStrings
            shared_strings = []
            try:
                ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                si_pattern = r'<si[^>]*>(.*?)</si>'
                for si_match in re.finditer(si_pattern, ss_xml, re.DOTALL):
                    si_content = si_match.group(1)
                    texts = re.findall(r'<t[^>]*>([^<]*)</t>', si_content)
                    shared_strings.append(''.join(texts))
            except:
                pass

            # 4. Trouver la colonne Source (chercher dans ligne 2)
            sheet_xml = zf.read(sheet_file).decode('utf-8')

            # Trouver l'index de "Source" dans sharedStrings
            source_str_idx = None
            for idx, s in enumerate(shared_strings):
                if s == 'Source':
                    source_str_idx = idx
                    break

            if source_str_idx is None:
                print("[CALLS CHART] 'Source' non trouvé dans sharedStrings")
                return {}

            # Trouver quelle colonne contient "Source" dans la ligne 2
            row2_match = re.search(r'<row r="2"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL)
            if not row2_match:
                return {}

            source_col = None
            row2_content = row2_match.group(1)
            cell_pattern = r'<c r="([A-Z]+)2"[^>]*t="s"[^>]*><v>(\d+)</v></c>'
            for cell_match in re.finditer(cell_pattern, row2_content):
                col = cell_match.group(1)
                str_idx = int(cell_match.group(2))
                if str_idx == source_str_idx:
                    source_col = col
                    break

            if not source_col:
                print("[CALLS CHART] Colonne Source non trouvée")
                return {}

            # 5. Extraire toutes les valeurs de la colonne Source
            source_counts = Counter()
            row_pattern = r'<row r="(\d+)"[^>]*>(.*?)</row>'

            for row_match in re.finditer(row_pattern, sheet_xml, re.DOTALL):
                row_num = int(row_match.group(1))
                if row_num <= 2:  # Skip header
                    continue

                row_content = row_match.group(2)

                cell_pattern = rf'<c r="{source_col}{row_num}"([^>]*)(?:><v>([^<]*)</v></c>|></c>|/>)'
                cell_match = re.search(cell_pattern, row_content)

                source_value = "Natif"  # Valeur par défaut

                if cell_match:
                    attrs = cell_match.group(1) or ""
                    v_val = cell_match.group(2)

                    if v_val is not None and v_val != "":
                        if 't="s"' in attrs:
                            str_idx = int(v_val)
                            if str_idx < len(shared_strings):
                                val = shared_strings[str_idx]
                                if val and val.strip():
                                    source_value = val.strip()
                        else:
                            if v_val.strip():
                                source_value = v_val.strip()

                # Filtrer les sources exclues
                if source_value not in EXCLUDED_SOURCES:
                    source_counts[source_value] += 1

            return dict(source_counts)

    except Exception as e:
        print(f"[CALLS CHART] Erreur lecture sources: {e}")
        return {}


def generate_calls_chart(source_counts: Dict[str, int]) -> Optional[str]:
    """
    Génère un graphique PNG du volume d'appels par plateforme.
    Design moderne et professionnel.
    Retourne le chemin du fichier PNG généré, ou None si échec.
    """
    try:
        if not source_counts:
            print("[CALLS CHART] Aucune donnée source")
            return None

        sorted_sources = sorted(source_counts.items(), key=lambda x: x[1], reverse=True)
        sources = [s[0] for s in sorted_sources]
        counts = [s[1] for s in sorted_sources]

        colors = get_chart_colors(sources, SOURCE_COLORS_MAP)
        return _generate_vertical_chart(
            labels=sources,
            values=counts,
            title="Volume d'appels par plateforme",
            filename_prefix='chart_calls',
            colors=colors,
            y_label="Nombre d'appels",
        )

    except Exception as e:
        print(f"[CALLS CHART ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def generate_calls_top15_count_chart(top15_data: List[Dict], source: str) -> Optional[str]:
    """
    Génère un graphique horizontal pour le top 15 des contacts par nombre d'appels.
    Priorité: Pseudonyme (Name) avec émojis, sinon Identifiant_utilisateur, sinon Identifier.
    """
    try:
        if not top15_data:
            print(f"[CALLS TOP15 COUNT] Aucune donnée pour {source}")
            return None

        source_lower = source.lower().strip()
        phone_sources = ["natif", "native messages", "whatsapp", "whatsapp business"]

        used_pseudonymes = 0
        used_phones = 0

        # Agréger par LABEL affiché (pseudonyme) : un même pseudonyme avec plusieurs
        # identifiants différents doit donner UNE seule barre dont le total = somme des appels
        # (cohérent avec le tableau résumé).
        agg_counts = {}   # label -> nombre d'appels cumulé
        label_order = []  # préserve l'ordre de première apparition
        for d in top15_data:
            # Priorité 1: Pseudonyme (Name) avec émojis
            name = (d.get("Name") or "").strip()
            if name and name.lower() != "inconnu":
                label = name[:50]
                used_pseudonymes += 1
            else:
                # Priorité 2: Identifiant_utilisateur ou Identifier selon la source
                ident = d.get("Identifiant_utilisateur") or d.get("Identifier") or "Inconnu"
                if ident and "@" in ident:
                    ident = ident.split("@")[0]
                    if ident and ident[0].isdigit():
                        ident = "+" + ident
                        used_phones += 1
                elif ident and ident[0].isdigit():
                    used_phones += 1
                label = (ident or "Inconnu")[:50]

            cnt = d.get("Nombre_appels", 0) or 0
            if label in agg_counts:
                agg_counts[label] += cnt
            else:
                agg_counts[label] = cnt
                label_order.append(label)

        # Re-trier par nombre d'appels décroissant après fusion des pseudonymes
        merged_labels = sorted(label_order, key=lambda l: agg_counts[l], reverse=True)
        identifiers = merged_labels
        counts = [agg_counts[l] for l in merged_labels]

        # Légende axe X adaptée selon ce qui est RÉELLEMENT affiché
        if used_pseudonymes >= used_phones and used_pseudonymes > 0:
            x_label = "Pseudonyme"
        elif used_phones > 0:
            x_label = "Numéro de téléphone"
        elif source_lower in phone_sources:
            x_label = "Numéro de téléphone"
        else:
            x_label = "Identifiant"

        return _generate_vertical_chart(
            labels=identifiers,
            values=counts,
            title=f"Top 15 des appels - {source}",
            filename_prefix=f"chart_calls_top15_count_{source.lower().replace(' ', '_')}",
            y_label="Nombre d'appels",
            x_label=x_label,
        )

    except Exception as e:
        print(f"[CALLS TOP15 COUNT ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def generate_calls_top15_duration_chart(top15_data: List[Dict], source: str) -> Optional[str]:
    """
    Génère un graphique horizontal pour le top 15 des contacts par durée d'appels.
    Priorité: Pseudonyme (Name) avec émojis, sinon Identifiant_utilisateur, sinon Identifier.
    """
    try:
        if not top15_data:
            print(f"[CALLS TOP15 DURATION] Aucune donnée pour {source}")
            return None

        source_lower = source.lower().strip()
        phone_sources = ["natif", "native messages", "whatsapp", "whatsapp business"]

        used_pseudonymes = 0
        used_phones = 0

        # Agréger par LABEL affiché (pseudonyme) : un même pseudonyme avec plusieurs
        # identifiants différents = UNE barre dont la durée = SOMME des durées (cohérent
        # avec le tableau résumé).
        agg_sec = {}      # label -> durée totale cumulée (secondes)
        label_order = []
        for d in top15_data:
            # Priorité 1: Pseudonyme (Name) avec émojis
            name = (d.get("Name") or "").strip()
            if name and name.lower() != "inconnu":
                label = name[:50]
                used_pseudonymes += 1
            else:
                # Priorité 2: Identifiant_utilisateur ou Identifier selon la source
                ident = d.get("Identifiant_utilisateur") or d.get("Identifier") or "Inconnu"
                if ident and "@" in ident:
                    ident = ident.split("@")[0]
                    if ident and ident[0].isdigit():
                        ident = "+" + ident
                        used_phones += 1
                elif ident and ident[0].isdigit():
                    used_phones += 1
                label = (ident or "Inconnu")[:50]

            # Durée de cette entrée en secondes
            dur_str = d.get("Duree_totale", "00:00:00")
            try:
                parts = str(dur_str).split(':')
                if len(parts) == 3:
                    secs = int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
                elif len(parts) == 2:
                    secs = int(parts[0]) * 60 + int(parts[1])
                else:
                    secs = 0
            except (ValueError, TypeError):
                secs = 0

            if label in agg_sec:
                agg_sec[label] += secs
            else:
                agg_sec[label] = secs
                label_order.append(label)

        # Re-trier par durée totale décroissante après fusion des pseudonymes
        merged_labels = sorted(label_order, key=lambda l: agg_sec[l], reverse=True)
        identifiers = merged_labels
        durations_sec: List[int] = [agg_sec[l] for l in merged_labels]
        duration_labels: List[str] = []
        for secs in durations_sec:
            h = secs // 3600
            m = (secs % 3600) // 60
            s = secs % 60
            duration_labels.append(f"{h}h{m:02d}m" if h > 0 else f"{m}m{s:02d}s")

        durations_min = [s / 60 for s in durations_sec]

        # Légende axe X adaptée selon ce qui est RÉELLEMENT affiché
        if used_pseudonymes >= used_phones and used_pseudonymes > 0:
            x_label = "Pseudonyme"
        elif used_phones > 0:
            x_label = "Numéro de téléphone"
        elif source_lower in phone_sources:
            x_label = "Numéro de téléphone"
        else:
            x_label = "Identifiant"

        return _generate_vertical_chart(
            labels=identifiers,
            values=durations_min,
            title=f"Top 15 duree des appels - {source}",
            filename_prefix=f"chart_calls_top15_duration_{source.lower().replace(' ', '_')}",
            y_label="Duree (minutes)",
            value_labels=duration_labels,
            x_label=x_label,
        )

    except Exception as e:
        print(f"[CALLS TOP15 DURATION ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def generate_chats_top15_count_chart(top15_data: List[Dict], source: str) -> Optional[str]:
    """
    Genere un graphique vertical pour le top 15 des contacts par nombre de messages.
    Priorité: Pseudonyme (Name) avec émojis, sinon Identifiant_utilisateur, sinon Identifier.
    """
    try:
        if not top15_data:
            print(f'[CHATS TOP15 COUNT] Aucune donnee pour {source}')
            return None

        source_lower = source.lower().strip()
        phone_sources = ["natif", "native messages", "whatsapp", "whatsapp business"]

        used_pseudonymes = 0
        used_phones = 0

        # Agréger par LABEL affiché (pseudonyme) : un même pseudonyme avec plusieurs
        # identifiants différents = UNE barre dont le total = somme des messages (cohérent
        # avec le tableau résumé).
        agg_counts = {}
        label_order = []
        for d in top15_data:
            # Priorité 1: Pseudonyme (Name) avec émojis
            name = (d.get("Name") or "").strip()
            if name and name.lower() != "inconnu":
                label = name[:50]
                used_pseudonymes += 1
            else:
                # Priorité 2: Identifiant_utilisateur ou Identifier selon la source
                ident = d.get('Identifiant_utilisateur') or d.get('Identifier') or 'Inconnu'
                # Nettoyer les numéros (590690196318@s.whatsapp.net → +590690196318)
                if ident and "@" in ident:
                    ident = ident.split("@")[0]
                    if ident and ident[0].isdigit():
                        ident = "+" + ident
                        used_phones += 1
                elif ident and ident[0].isdigit():
                    used_phones += 1
                label = (ident or 'Inconnu')[:50]

            cnt = d.get('Nombre_messages', 0) or 0
            if label in agg_counts:
                agg_counts[label] += cnt
            else:
                agg_counts[label] = cnt
                label_order.append(label)

        merged_labels = sorted(label_order, key=lambda l: agg_counts[l], reverse=True)
        identifiers = merged_labels
        counts = [agg_counts[l] for l in merged_labels]

        # Légende axe X adaptée selon ce qui est RÉELLEMENT affiché
        if used_pseudonymes >= used_phones and used_pseudonymes > 0:
            x_label = "Pseudonyme"
        elif used_phones > 0:
            x_label = "Numéro de téléphone"
        elif source_lower in phone_sources:
            x_label = "Numéro de téléphone"
        else:
            x_label = "Identifiant"

        return _generate_vertical_chart(
            labels=identifiers,
            values=counts,
            title=f'Top 15 des messages - {source}',
            filename_prefix=f'chart_chats_top15_count_{source.lower().replace(" ", "_")}',
            y_label='Nombre de messages',
            x_label=x_label,
        )

    except Exception as e:
        print(f'[CHATS TOP15 COUNT ERROR] {type(e).__name__}: {str(e)}')
        import traceback
        traceback.print_exc()
        return None


def get_chats_by_source_fast(content: bytes) -> Dict[str, int]:
    """
    Compte les messages par source via lecture XML directe (ultra rapide).
    Exclut les lignes où 'Instant Message #' est vide.
    Retourne un dictionnaire {source: count}.
    """
    import zipfile
    import re
    from collections import Counter

    # Sources Ã  exclure (mêmes que contacts/appels)
    EXCLUDED_SOURCES = {'Recents', 'InteractionC', 'KnowledgeC', 'Native Messages', 'Threads', 'Biome', 'Unified Logs', 'SIM'}

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver la feuille Chats
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            match = re.search(r'<sheet[^>]*name="Chats"[^>]*r:id="(rId\d+)"', workbook_xml)
            if not match:
                match = re.search(r'<sheet[^>]*r:id="(rId\d+)"[^>]*name="Chats"', workbook_xml)
            if not match:
                print("[CHATS CHART] Feuille 'Chats' non trouvée")
                return {}

            sheet_rid = match.group(1)

            # 2. Trouver le fichier XML
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            match = re.search(rf'<Relationship[^>]*Id="{sheet_rid}"[^>]*Target="([^"]+)"', rels_xml)
            if not match:
                return {}

            target = match.group(1)
            sheet_file = 'xl/' + target if not target.startswith('xl/') else target

            # 3. Charger sharedStrings
            shared_strings = []
            try:
                ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                si_pattern = r'<si[^>]*>(.*?)</si>'
                for si_match in re.finditer(si_pattern, ss_xml, re.DOTALL):
                    si_content = si_match.group(1)
                    texts = re.findall(r'<t[^>]*>([^<]*)</t>', si_content)
                    shared_strings.append(''.join(texts))
            except:
                pass

            # 4. Trouver les colonnes Source et Instant Message #
            sheet_xml = zf.read(sheet_file).decode('utf-8')

            # Index dans sharedStrings
            source_str_idx = None
            im_str_idx = None

            for idx, s in enumerate(shared_strings):
                if s == 'Source':
                    source_str_idx = idx
                elif s == 'Instant Message #':
                    im_str_idx = idx

            if source_str_idx is None:
                print("[CHATS CHART] 'Source' non trouvé dans sharedStrings")
                return {}

            # Trouver les colonnes dans la ligne 2
            row2_match = re.search(r'<row r="2"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL)
            if not row2_match:
                return {}

            source_col = None
            im_col = None

            row2_content = row2_match.group(1)
            for cell_match in re.finditer(r'<c r="([A-Z]+)2"[^>]*t="s"[^>]*><v>(\d+)</v></c>', row2_content):
                col = cell_match.group(1)
                str_idx = int(cell_match.group(2))

                if str_idx == source_str_idx:
                    source_col = col
                elif str_idx == im_str_idx:
                    im_col = col

            if not source_col:
                print("[CHATS CHART] Colonne Source non trouvée")
                return {}

            print(f"[CHATS CHART] Colonne 'Instant Message #': {im_col}")

            # Pattern pour extraire toutes les cellules d'une ligne en une passe
            cell_pattern = re.compile(r'<c r="([A-Z]+)\d+"([^>]*)(?:><v>([^<]*)</v></c>|></c>|/>)')

            # 5. Extraire les données
            source_counts = Counter()
            skipped_rows = 0

            for row_match in re.finditer(r'<row r="(\d+)"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL):
                row_num = int(row_match.group(1))
                if row_num <= 2:
                    continue

                row_content = row_match.group(2)

                # Extraire toutes les cellules en une seule passe
                source_value = None
                has_im = False

                for cell_match in cell_pattern.finditer(row_content):
                    col = cell_match.group(1)
                    attrs = cell_match.group(2) or ""
                    v_val = cell_match.group(3)

                    # Colonne Source
                    if col == source_col:
                        if v_val and 't="s"' in attrs:
                            idx = int(v_val)
                            if idx < len(shared_strings):
                                source_value = shared_strings[idx].strip()
                        elif v_val:
                            source_value = v_val.strip()

                    # Colonne Instant Message # - vérifier si elle a une valeur
                    elif col == im_col and not has_im:
                        if v_val:
                            has_im = True

                # Si Instant Message # vide, ignorer
                if im_col and not has_im:
                    skipped_rows += 1
                    continue

                # Valeur par défaut pour source
                if not source_value:
                    source_value = "Natif"

                # Filtrer les sources exclues
                if source_value not in EXCLUDED_SOURCES:
                    source_counts[source_value] += 1

            if skipped_rows > 0:
                print(f"[CHATS CHART] {skipped_rows} lignes ignorées (Instant Message # vide)")

            return dict(source_counts)

    except Exception as e:
        print(f"[CHATS CHART] Erreur lecture sources: {e}")
        return {}


def generate_chats_chart(source_counts: Dict[str, int]) -> Optional[str]:
    """
    Génère un graphique PNG du volume de messages par plateforme.
    Design moderne et professionnel.
    Retourne le chemin du fichier PNG généré, ou None si échec.
    """
    try:
        if not source_counts:
            print("[CHATS CHART] Aucune donnée source")
            return None

        sorted_sources = sorted(source_counts.items(), key=lambda x: x[1], reverse=True)
        sources = [s[0] for s in sorted_sources]
        counts = [s[1] for s in sorted_sources]

        colors = get_chart_colors(sources, SOURCE_COLORS_MAP)
        return _generate_vertical_chart(
            labels=sources,
            values=counts,
            title="Volume de messages par plateforme",
            filename_prefix='chart_chats',
            colors=colors,
            y_label="Nombre de messages",
        )

    except Exception as e:
        print(f"[CHATS CHART ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


# =========================================================================
# Rendu texte multi-langue (fallback de police par caractere)
# Probleme: ImageFont.truetype("seguiemj.ttf") n'a pas de glyphes pour l'Arabe,
# le Cyrillique etc. -> les caracteres apparaissent en "carres" dans les bulles.
# Solution: pour chaque caractere non supporte par la police principale, on
# bascule sur une chaine de fallback (Segoe UI, Arial, Tahoma, DejaVu...).
# =========================================================================
_CHAR_SUPPORT_CACHE: Dict[Tuple[int, str], bool] = {}
_FALLBACK_FONT_CACHE: Dict[int, List[ImageFont.FreeTypeFont]] = {}
_FONT_CMAP_CACHE: Dict[str, Optional[set]] = {}  # font path -> set de codepoints couverts


def _get_fallback_fonts(size: int) -> List[ImageFont.FreeTypeFont]:
    """Chaine de polices de fallback a la taille demandee (cache global)."""
    if size in _FALLBACK_FONT_CACHE:
        return _FALLBACK_FONT_CACHE[size]
    paths = [
        "C:/Windows/Fonts/seguiemj.ttf",    # Emojis (essentiel quand la principale est Calibri)
        "C:/Windows/Fonts/segoeui.ttf",     # Latin + Cyrillique + Grec
        "C:/Windows/Fonts/arial.ttf",       # Latin + Cyrillique + Arabe (basique)
        "C:/Windows/Fonts/tahoma.ttf",      # Large coverage (incl. Arabe, Thai)
        "C:/Windows/Fonts/arialuni.ttf",    # Arial Unicode MS - tres large (si MS Office)
        "C:/Windows/Fonts/seguisym.ttf",    # Symbols / fallback
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoColorEmoji.ttf",
    ]
    fonts: List[ImageFont.FreeTypeFont] = []
    for p in paths:
        try:
            fonts.append(ImageFont.truetype(p, size))
        except (OSError, IOError):
            continue
    _FALLBACK_FONT_CACHE[size] = fonts
    return fonts


def _get_font_cmap(font: ImageFont.FreeTypeFont) -> Optional[set]:
    """
    Renvoie l'ensemble des codepoints couverts par la police, en lisant
    directement sa table cmap via fontTools. Cache global par chemin.
    Plus fiable que rasteriser un masque (qui produit un "tofu" pour les
    glyphes absents et trompe getbbox()).
    """
    path = getattr(font, "path", None)
    if not path:
        return None
    if path in _FONT_CMAP_CACHE:
        return _FONT_CMAP_CACHE[path]
    try:
        from fontTools.ttLib import TTFont
        tt = TTFont(path, lazy=True)
        cmap = set(tt.getBestCmap().keys())
        tt.close()
    except Exception as e:
        print(f"[FONT CMAP] Erreur lecture {path}: {type(e).__name__}: {e}")
        cmap = None
    _FONT_CMAP_CACHE[path] = cmap
    return cmap


def _font_supports_char(font: ImageFont.FreeTypeFont, ch: str) -> bool:
    """
    True si la police a un glyphe pour ce caractere.
    On lit la table cmap via fontTools (fiable). Fallback rasterisation si
    cmap indisponible.
    """
    if not ch:
        return True
    key = (id(font), ch)
    cached = _CHAR_SUPPORT_CACHE.get(key)
    if cached is not None:
        return cached
    cmap = _get_font_cmap(font)
    if cmap is not None:
        try:
            supported = ord(ch) in cmap
        except TypeError:
            # ch est une sequence (combining marks, ZWJ) -> 1er codepoint
            supported = ord(ch[0]) in cmap if ch else False
    else:
        try:
            supported = font.getmask(ch).getbbox() is not None
        except Exception:
            supported = False
    _CHAR_SUPPORT_CACHE[key] = supported
    return supported


def _draw_text_safe(draw, xy, text, fill=None, font=None, **kwargs) -> None:
    """
    Drop-in remplacement de draw.text() avec fallback automatique de police
    pour les caracteres non supportes par la police principale (Arabe, Cyrillique,
    Hebreu, Thai, etc.). Les caracteres consecutifs supportes par la meme police
    sont regroupes en un seul appel pour la performance.
    """
    if text is None or text == "":
        return
    text = str(text)
    if font is None:
        draw.text(xy, text, fill=fill, **kwargs)
        return
    try:
        size = int(getattr(font, "size", 0) or 0)
    except Exception:
        size = 0
    fallbacks = _get_fallback_fonts(size) if size else []
    chain = [font] + [f for f in fallbacks if f is not font]

    def pick(ch):
        for f in chain:
            if _font_supports_char(f, ch):
                return f
        return font

    x, y = xy
    current_font = pick(text[0])
    run = text[0]
    for ch in text[1:]:
        f = pick(ch)
        if f is current_font:
            run += ch
        else:
            draw.text((x, y), run, font=current_font, fill=fill, **kwargs)
            try:
                x += int(draw.textlength(run, font=current_font))
            except Exception:
                try:
                    x += int(current_font.getlength(run))
                except Exception:
                    x += len(run) * max(1, size // 2)
            current_font = f
            run = ch
    if run:
        draw.text((x, y), run, font=current_font, fill=fill, **kwargs)


def generate_conversation_image(messages: List[Dict], contact_name: str, source: str) -> List[str]:
    """
    Génère une ou plusieurs images PNG de la conversation avec le style du frontend.
    Owner = bulle verte à droite, Other = bulle bleue à gauche.
    Retourne une liste de chemins d'images (paginées si conversation longue).
    """
    try:
        if not messages:
            print("[CONV IMAGE] Aucun message")
            return []

        # Configuration - Bulles plus compactes
        img_width = 700  # Réduit de 800 à 700
        max_page_height = 1000  # Hauteur max par page
        padding = 15  # Réduit de 20 à 15
        bubble_padding = 10  # Réduit de 12 à 10
        bubble_radius = 12  # Réduit de 15 à 12
        max_bubble_width = 400  # Réduit de 500 à 400
        line_height = 20  # Réduit de 22 à 20
        sender_height = 18  # Réduit de 20 à 18
        timestamp_height = 14  # Réduit de 16 à 14
        bubble_spacing = 10  # Réduit de 12 à 10
        title_height = 45  # Réduit de 50 à 45

        # Couleurs
        owner_color = (37, 211, 102)  # Vert WhatsApp #25D366
        other_color = (0, 132, 255)   # Bleu Messenger #0084ff
        text_color = (255, 255, 255)  # Blanc
        bg_color = (248, 249, 250)    # Gris clair
        title_color = (30, 58, 95)    # Bleu foncé

        # Polices: Calibri 11 (body), Calibri Bold 12 (identifiant + horodatage),
        # Calibri Bold 16 pour le titre. Emojis/Cyrillique/Arabe via fallback.
        try:
            font_sender = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 12)
            font_body = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 11)
            font_timestamp = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 12)
            font_title = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 16)
        except (OSError, IOError):
            try:
                # Fallback Linux/macOS
                font_sender = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 12)
                font_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 11)
                font_timestamp = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 12)
                font_title = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 16)
            except (OSError, IOError):
                font_sender = ImageFont.load_default()
                font_body = ImageFont.load_default()
                font_timestamp = ImageFont.load_default()
                font_title = ImageFont.load_default()

        # Étape 1: Précalculer toutes les bulles avec leurs infos
        bubbles = []
        for msg in messages:
            is_owner = msg.get("is_owner", False)
            from_user = msg.get("from", "").replace("(owner)", "").strip()
            body = msg.get("body", "") or ""
            attachments = msg.get("attachments", [])
            timestamp = msg.get("timestamp_time", "") or msg.get("timestamp_date", "")

            content_lines = [("sender", from_user)]

            if body and body.strip() and body != "None":
                words = body.split()
                line = ""
                for word in words:
                    test_line = line + " " + word if line else word
                    if len(test_line) * 7 > max_bubble_width - bubble_padding * 2:
                        content_lines.append(("body", line))
                        line = word
                    else:
                        line = test_line
                if line:
                    content_lines.append(("body", line))

            if attachments:
                for att in attachments:
                    att_name = str(att).split("/")[-1].split("\\")[-1]
                    content_lines.append(("attachment", f"📎 {att_name}"))

            if timestamp and timestamp != "nan" and timestamp != "None":
                content_lines.append(("timestamp", timestamp))

            # Calculer dimensions
            bubble_height = bubble_padding * 2
            bubble_width = 100
            for line_type, line_text in content_lines:
                if line_type == "sender":
                    bubble_height += sender_height
                    w = len(line_text) * 9
                elif line_type == "timestamp":
                    bubble_height += timestamp_height
                    w = len(line_text) * 6
                elif line_type == "attachment":
                    bubble_height += 18
                    w = len(line_text) * 7
                else:
                    bubble_height += line_height
                    w = len(line_text) * 7
                bubble_width = max(bubble_width, w + bubble_padding * 2)
            bubble_width = min(bubble_width, max_bubble_width)

            bubbles.append({
                "is_owner": is_owner,
                "content_lines": content_lines,
                "height": bubble_height,
                "width": bubble_width,
                "total_height": bubble_height + bubble_spacing
            })

        # Étape 2: Grouper les bulles par pages
        pages = []
        current_page = []
        current_height = title_height + padding

        for bubble in bubbles:
            if current_height + bubble["total_height"] > max_page_height and current_page:
                pages.append(current_page)
                current_page = []
                current_height = title_height + padding

            current_page.append(bubble)
            current_height += bubble["total_height"]

        if current_page:
            pages.append(current_page)

        # Étape 3: Générer une image pour chaque page
        image_paths = []
        base_id = uuid.uuid4().hex[:8]

        for page_idx, page_bubbles in enumerate(pages):
            # Calculer la hauteur de cette page
            page_height = title_height + padding
            for b in page_bubbles:
                page_height += b["total_height"]
            page_height += padding

            # Créer l'image
            img = Image.new('RGB', (img_width, page_height), bg_color)
            draw = ImageDraw.Draw(img)

            # Titre (avec numéro de page si plusieurs)
            if len(pages) > 1:
                title_text = f"Conversation avec {contact_name} ({source}) - Page {page_idx + 1}/{len(pages)}"
            else:
                title_text = f"Conversation avec {contact_name} ({source})"
            title_bbox = draw.textbbox((0, 0), title_text, font=font_title)
            title_x = (img_width - (title_bbox[2] - title_bbox[0])) // 2
            _draw_text_safe(draw,(title_x, padding), title_text, fill=title_color, font=font_title)

            y = title_height + padding

            # Dessiner les bulles de cette page
            for bubble in page_bubbles:
                bubble_color = owner_color if bubble["is_owner"] else other_color

                if bubble["is_owner"]:
                    bubble_x = img_width - padding - bubble["width"]
                else:
                    bubble_x = padding

                # Rectangle arrondi
                bubble_rect = [bubble_x, y, bubble_x + bubble["width"], y + bubble["height"]]
                draw.rounded_rectangle(bubble_rect, radius=bubble_radius, fill=bubble_color)

                # Contenu
                text_y = y + bubble_padding
                for line_type, line_text in bubble["content_lines"]:
                    text_x = bubble_x + bubble_padding
                    if line_type == "sender":
                        _draw_text_safe(draw,(text_x, text_y), line_text, fill=text_color, font=font_sender)
                        text_y += sender_height
                    elif line_type == "timestamp":
                        _draw_text_safe(draw,(text_x, text_y), line_text, fill=(220, 220, 220), font=font_timestamp)
                        text_y += timestamp_height
                    elif line_type == "attachment":
                        _draw_text_safe(draw,(text_x, text_y), line_text, fill=(200, 200, 200), font=font_timestamp)
                        text_y += 18
                    else:
                        _draw_text_safe(draw,(text_x, text_y), line_text, fill=text_color, font=font_body)
                        text_y += line_height

                y += bubble["total_height"]

            # Sauvegarder
            safe_name = re.sub(r'[^\w\-]', '_', contact_name)
            if len(pages) > 1:
                conv_filename = f"conv_{safe_name}_{base_id}_p{page_idx + 1}.png"
            else:
                conv_filename = f"conv_{safe_name}_{base_id}.png"
            conv_path = UPLOAD_DIR / conv_filename
            img.save(conv_path, "PNG", quality=95, dpi=(200, 200))
            image_paths.append(str(conv_path))

        print(f"[CONV IMAGE] {len(image_paths)} image(s) générée(s) pour {contact_name}")
        return image_paths

    except Exception as e:
        print(f"[CONV IMAGE ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return []


def generate_message_images(messages: List[Dict], contact_name: str, source: str, audio_transcriptions: Dict[str, str] = None, verbose: bool = False) -> List[Dict]:
    """
    Génère une image PNG pour CHAQUE message de la conversation.
    Retourne une liste de dictionnaires avec les infos de chaque message.
    FOND TRANSPARENT - seulement la bulle colorée visible.
    audio_transcriptions: dictionnaire {nom_audio: transcription}
    verbose: si True, affiche les logs détaillés (ralentit l'exécution)
    """
    if audio_transcriptions is None:
        audio_transcriptions = {}
    try:
        if not messages:
            return []

        # Compteur pour log de progression
        import time as time_mod
        gen_start = time_mod.perf_counter()

        # Configuration - HAUTE RÉSOLUTION avec largeur FIXE
        # Facteur de résolution pour améliorer la qualité
        RESOLUTION_SCALE = 2  # 2x pour haute résolution

        bubble_padding = 8 * RESOLUTION_SCALE
        bubble_radius = 10 * RESOLUTION_SCALE
        # Bulles FULL (frontend preview) alignees sur les CLEAN (Calibri 16).
        # Largeur fixe equivalente a 7 cm en Word pour les bulles texte ;
        # pour les bulles avec miniatures video, la largeur s'etend pour caser les thumbs.
        MIN_BUBBLE_WIDTH = 420 * RESOLUTION_SCALE        # 840 px = identique a CLEAN_BUBBLE_WIDTH
        MAX_BUBBLE_WIDTH = 420 * RESOLUTION_SCALE        # FIXE 840 px pour bulles texte
        MAX_BUBBLE_WIDTH_VIDEO = 560 * RESOLUTION_SCALE  # 1120 px pour caser les miniatures video
        line_height = 24 * RESOLUTION_SCALE
        sender_height = 26 * RESOLUTION_SCALE
        timestamp_height = 24 * RESOLUTION_SCALE
        attachment_height = 28 * RESOLUTION_SCALE
        transcription_height = 30 * RESOLUTION_SCALE

        # Couleurs
        owner_color = (37, 211, 102)  # Vert
        other_color = (0, 132, 255)   # Bleu
        text_color = (255, 255, 255)  # Blanc

        # Polices pour les bulles FULL (frontend preview): Calibri 20 (memes que CLEAN).
        try:
            font_sender = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 20 * RESOLUTION_SCALE)
            font_body = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 20 * RESOLUTION_SCALE)
            font_timestamp = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 20 * RESOLUTION_SCALE)
        except (OSError, IOError):
            try:
                font_sender = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 20 * RESOLUTION_SCALE)
                font_body = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20 * RESOLUTION_SCALE)
                font_timestamp = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 20 * RESOLUTION_SCALE)
            except (OSError, IOError):
                font_sender = ImageFont.load_default()
                font_body = ImageFont.load_default()
                font_timestamp = ImageFont.load_default()

        # Polices pour les bulles CLEAN (backend docx): Calibri 16 partout (plus gros).
        # La bulle est PLUS ETROITE (9 cm Word) -> la police parait plus grande
        # proportionnellement. Wrap a 6 mots par ligne pour rester confortable.
        # Densite PNG: 1080 px @ 9 cm = ~305 DPI -> tres net.
        try:
            font_sender_clean = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 16 * RESOLUTION_SCALE)
            font_body_clean = ImageFont.truetype("C:/Windows/Fonts/calibri.ttf", 16 * RESOLUTION_SCALE)
            font_timestamp_clean = ImageFont.truetype("C:/Windows/Fonts/calibrib.ttf", 16 * RESOLUTION_SCALE)
        except (OSError, IOError):
            try:
                font_sender_clean = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 16 * RESOLUTION_SCALE)
                font_body_clean = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16 * RESOLUTION_SCALE)
                font_timestamp_clean = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 16 * RESOLUTION_SCALE)
            except (OSError, IOError):
                font_sender_clean = font_sender
                font_body_clean = font_body
                font_timestamp_clean = font_timestamp

        # Parametres clean: largeur 7 cm en Word, wrap 6 mots/ligne, police 16pt (inchangee).
        # PNG reduit proportionnellement (1080 -> 840 px) pour conserver le meme DPI (305) et
        # la meme taille visuelle de police que la version 9 cm.
        CLEAN_BUBBLE_WIDTH = 420 * RESOLUTION_SCALE   # 840 px -> 7 cm Word -> ~305 DPI
        CLEAN_WORDS_PER_LINE = 6
        CLEAN_LINE_HEIGHT = 24 * RESOLUTION_SCALE
        CLEAN_SENDER_HEIGHT = 26 * RESOLUTION_SCALE
        CLEAN_TIMESTAMP_HEIGHT = 24 * RESOLUTION_SCALE
        CLEAN_ATTACHMENT_HEIGHT = 28 * RESOLUTION_SCALE
        CLEAN_TRANSCRIPTION_HEIGHT = 30 * RESOLUTION_SCALE

        def _wrap_by_words(text, n_per_line):
            """Wrap par mots: chaque ligne contient au plus n_per_line mots."""
            if not text:
                return []
            words = str(text).split()
            if not words:
                return []
            return [" ".join(words[i:i + n_per_line]) for i in range(0, len(words), n_per_line)]

        def _measure_text_width(text, font):
            try:
                return int(font.getlength(text))
            except Exception:
                try:
                    bbox = font.getbbox(text)
                    return bbox[2] - bbox[0]
                except Exception:
                    return len(text) * 8 * RESOLUTION_SCALE

        def _wrap_by_pixels(text, font, max_width_px):
            """
            Wrap par largeur de pixels: chaque ligne tient dans max_width_px.
            Si un mot seul depasse, on le split caractere par caractere pour eviter
            qu'il soit tronque par le masque arrondi de la bulle (cas long URL,
            email, identifiant @s.whatsapp.net, etc.).
            """
            if not text:
                return []
            words = str(text).split()
            if not words:
                return []
            lines = []
            current = ""
            for word in words:
                # Mot seul plus large que la limite : split char par char.
                if _measure_text_width(word, font) > max_width_px:
                    if current:
                        lines.append(current)
                        current = ""
                    chunk = ""
                    for ch in word:
                        if chunk and _measure_text_width(chunk + ch, font) > max_width_px:
                            lines.append(chunk)
                            chunk = ch
                        else:
                            chunk += ch
                    if chunk:
                        current = chunk
                    continue
                candidate = f"{current} {word}".strip() if current else word
                if _measure_text_width(candidate, font) <= max_width_px:
                    current = candidate
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)
            return lines

        # Largeur de texte disponible (sans padding gauche/droit) pour le wrap.
        BUBBLE_TEXT_MAX_WIDTH = MAX_BUBBLE_WIDTH - 2 * bubble_padding
        CLEAN_BUBBLE_TEXT_MAX_WIDTH = CLEAN_BUBBLE_WIDTH - 2 * bubble_padding

        result = []
        base_id = uuid.uuid4().hex[:8]

        for msg_idx, msg in enumerate(messages):
            is_owner = msg.get("is_owner", False)
            from_user = msg.get("from", "").replace("(owner)", "").strip()
            body = msg.get("body", "") or ""
            attachments = msg.get("attachments", [])
            attachment_urls = msg.get("attachment_urls", [])
            attachment_transcripts = msg.get("attachment_transcripts", {})  # Transcriptions depuis l'Excel
            timestamp = msg.get("timestamp_time", "") or msg.get("timestamp_date", "")

            # Log debug pour les attachments
            if attachments or attachment_urls:
                print(f"[MSG {msg_idx}] attachments: {attachments}, attachment_urls: {len(attachment_urls) if attachment_urls else 0}")

            # Construire le contenu - wrap par MOTS (6 par ligne) aligne sur la version CLEAN.
            # Garantit un design identique entre frontend et backend pour les bulles texte.
            content_lines = []

            if from_user:
                for line in _wrap_by_pixels(from_user, font_sender, BUBBLE_TEXT_MAX_WIDTH):
                    content_lines.append(("sender", line))

            if body and body.strip() and body != "None":
                for line in _wrap_by_pixels(body, font_body, BUBBLE_TEXT_MAX_WIDTH):
                    content_lines.append(("body", line))

            # Collecter les images et vidéos d'attachments pour les afficher
            embedded_images = []  # Images dans la bulle (vidéos miniatures UNIQUEMENT)
            video_attachments = []  # Pour stocker les infos des vidéos
            image_attachments = []  # Pour stocker les chemins des images originales (insertion séparée dans Word)
            pdf_attachments = []  # Pour stocker les chemins des PDF (insertion séparée dans Word)
            seen_base_names = set()  # Pour détecter les doublons (nom et nom (2))

            def get_base_name(filename):
                """Retire les suffixes (2), (3), etc. pour détecter les doublons"""
                import re
                # Séparer nom et extension
                name_part = filename.rsplit('.', 1)[0] if '.' in filename else filename
                ext_part = '.' + filename.rsplit('.', 1)[1] if '.' in filename else ''
                # Retirer les suffixes comme (2), (3), _1, _2, etc.
                base = re.sub(r'\s*[\(\[]\d+[\)\]]$', '', name_part)  # (2), [2]
                base = re.sub(r'_\d+$', '', base)  # _1, _2
                return (base + ext_part).lower()

            def get_clean_name(filename):
                """Retire les suffixes (2), (3), etc. pour l'affichage"""
                import re
                name_part = filename.rsplit('.', 1)[0] if '.' in filename else filename
                ext_part = '.' + filename.rsplit('.', 1)[1] if '.' in filename else ''
                base = re.sub(r'\s*[\(\[]\d+[\)\]]$', '', name_part)
                base = re.sub(r'_\d+$', '', base)
                return base + ext_part

            print(f"[MSG DEBUG] attachment_urls count: {len(attachment_urls) if attachment_urls else 0}")
            if attachment_urls:
                for att_info in attachment_urls:
                    att_name = att_info.get("name", "")
                    att_url = att_info.get("url")
                    print(f"[MSG DEBUG] Attachment: name='{att_name}', url='{att_url}'")

                    # Vérifier si c'est un doublon
                    base_name = get_base_name(att_name)
                    if base_name in seen_base_names:
                        if verbose: print(f"[MSG IMAGES] Doublon ignoré: {att_name}")
                        continue
                    seen_base_names.add(base_name)

                    if att_url:
                        att_lower = att_name.lower()
                        # Vérifier si c'est une image
                        if any(att_lower.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp']):
                            img_path = UPLOAD_DIR / att_url.replace("/uploads/", "")
                            print(f"[IMG ATTACH] Image détectée: {att_name}, URL: {att_url}, Path: {img_path}, Exists: {img_path.exists()}")
                            if img_path.exists():
                                # Stocker l'image originale pour Word (haute qualité)
                                image_attachments.append({
                                    "name": att_name,
                                    "path": str(img_path),
                                    "url": att_url
                                })
                                print(f"[IMG ATTACH] ✓ Image ajoutée aux attachments: {att_name}")
                                # Afficher le nom de l'image dans la bulle (sans suffixe doublon)
                                clean_name = get_clean_name(att_name)
                                content_lines.append(("attachment", f"🖼️ {clean_name}"))
                                # NOTE: Les images sont maintenant insérées séparément dans Word
                                # pour permettre à l'utilisateur de les redimensionner
                                # (Stockées dans image_attachments ci-dessus)
                            else:
                                print(f"[IMG ATTACH] ✗ Image non trouvée: {img_path}")
                                content_lines.append(("attachment", f"📎 {get_clean_name(att_name)}"))
                        # Vérifier si c'est un fichier audio
                        elif any(att_lower.endswith(ext) for ext in ['.opus', '.ogg', '.aac', '.m4a', '.wav', '.mp3', '.amr', '.wma']):
                            clean_audio_name = get_clean_name(att_name)
                            content_lines.append(("attachment", f"🎵 {clean_audio_name}"))
                            # Chercher la transcription d'abord dans attachment_transcripts (Excel Chats),
                            # puis fallback sur audio_transcriptions (feuille Audios)
                            audio_base = att_name.rsplit('.', 1)[0] if '.' in att_name else att_name
                            transcription = None

                            # 1. Chercher dans attachment_transcripts du message (colonnes Attachment #X - Transcript)
                            if attachment_transcripts:
                                for trans_key in attachment_transcripts:
                                    trans_key_lower = trans_key.lower()
                                    att_name_lower = att_name.lower()
                                    audio_base_lower = audio_base.lower()
                                    if (trans_key_lower == att_name_lower or
                                        trans_key_lower == audio_base_lower or
                                        att_name_lower in trans_key_lower or
                                        trans_key_lower in att_name_lower):
                                        transcription = attachment_transcripts[trans_key]
                                        break

                            # 2. Fallback: Chercher dans audio_transcriptions (feuille Audios)
                            if not transcription:
                                for audio_key in audio_transcriptions:
                                    # Correspondance exacte ou partielle du nom
                                    if audio_key and (audio_key == att_name or audio_key == audio_base or
                                        att_name in audio_key or audio_key in att_name or
                                        audio_base in audio_key or audio_key in audio_base):
                                        transcription = audio_transcriptions[audio_key]
                                        break

                            if transcription and str(transcription).strip() and str(transcription).lower() != 'nan':
                                # Découper la transcription en lignes si trop longue
                                trans_text = str(transcription).strip()
                                max_chars_per_line = 40  # Caractères max par ligne
                                trans_words = trans_text.split()
                                trans_lines = []
                                trans_line = ""
                                for word in trans_words:
                                    test_line = trans_line + " " + word if trans_line else word
                                    if len(test_line) > max_chars_per_line:
                                        trans_lines.append(trans_line)
                                        trans_line = word
                                    else:
                                        trans_line = test_line
                                if trans_line:
                                    trans_lines.append(trans_line)
                                # Première ligne avec icône 📝, les autres sans
                                for i, line in enumerate(trans_lines):
                                    if i == 0:
                                        content_lines.append(("transcription", f"📝 {line}"))
                                    else:
                                        content_lines.append(("transcription", f"    {line}"))  # Indentation sans icône
                        # Vérifier si c'est un PDF
                        elif att_lower.endswith('.pdf'):
                            pdf_path = UPLOAD_DIR / att_url.replace("/uploads/", "")
                            if verbose: print(f"[MSG IMAGES] PDF détecté: {att_name}, Path: {pdf_path}, Exists: {pdf_path.exists()}")
                            content_lines.append(("attachment", f"📄 {get_clean_name(att_name)}"))
                            if pdf_path.exists():
                                try:
                                    import fitz  # PyMuPDF
                                    # Ouvrir le PDF et extraire la première page
                                    pdf_doc = fitz.open(str(pdf_path))
                                    if len(pdf_doc) > 0:
                                        first_page = pdf_doc[0]
                                        # Convertir en image avec TRÈS HAUTE résolution (zoom 5x)
                                        zoom = 5.0  # Facteur de zoom très élevé pour qualité maximale
                                        mat = fitz.Matrix(zoom, zoom)
                                        pix = first_page.get_pixmap(matrix=mat, alpha=False)
                                        # Convertir en PIL Image
                                        from PIL import ImageFilter, ImageEnhance
                                        pdf_img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                                        if verbose: print(f"[MSG IMAGES] PDF résolution brute: {pdf_img.width}x{pdf_img.height}")
                                        # Améliorer la netteté avec UnsharpMask (optimal pour texte PDF)
                                        pdf_img = pdf_img.filter(ImageFilter.UnsharpMask(radius=1.2, percent=100, threshold=1))
                                        # Améliorer le contraste pour le texte
                                        contrast_enhancer = ImageEnhance.Contrast(pdf_img)
                                        pdf_img = contrast_enhancer.enhance(1.12)
                                        # Sauvegarder comme image PNG pour insertion séparée dans Word
                                        pdf_img_filename = f"pdf_{uuid.uuid4().hex[:8]}_{att_name.replace('.pdf', '.png')}"
                                        pdf_img_path = UPLOAD_DIR / pdf_img_filename
                                        pdf_img.save(pdf_img_path, "PNG", dpi=(300, 300))
                                        pdf_attachments.append({
                                            "name": att_name,
                                            "path": str(pdf_img_path),
                                            "url": f"/uploads/{pdf_img_filename}"
                                        })
                                        if verbose: print(f"[MSG IMAGES] PDF image créée: {att_name} ({pdf_img.width}x{pdf_img.height})")
                                    pdf_doc.close()
                                except Exception as e:
                                    if verbose: print(f"[MSG IMAGES] Erreur extraction PDF {att_name}: {e}")
                                    import traceback
                                    traceback.print_exc()
                        # Vérifier si c'est une vidéo
                        elif any(att_lower.endswith(ext) for ext in ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.m4v', '.3gp']):
                            # Construire le chemin vidéo correctement
                            if att_url:
                                video_filename = att_url.replace("/uploads/", "").replace("\\uploads\\", "").lstrip("/\\")
                            else:
                                # Si pas d'URL, essayer avec le nom directement
                                video_filename = att_name
                            video_path = UPLOAD_DIR / video_filename
                            if verbose: print(f"[MSG IMAGES] Vidéo détectée: {att_name}, URL: {att_url}, Path: {video_path}, Exists: {video_path.exists()}")

                            # Si le fichier n'existe pas, chercher par nom dans uploads
                            if not video_path.exists():
                                # Recherche dans le dossier uploads
                                for f in UPLOAD_DIR.iterdir():
                                    if f.is_file() and f.name.lower() == att_name.lower():
                                        video_path = f
                                        if verbose: print(f"[MSG IMAGES] Vidéo trouvée par recherche: {video_path}")
                                        break

                            if video_path.exists():
                                try:
                                    # Générer 8 miniatures pour la bulle (alignées en 4 colonnes)
                                    thumbnails = extract_video_thumbnails(video_path, num_thumbnails=8)
                                    if thumbnails:
                                        # Afficher le nom de la vidéo dans la bulle
                                        content_lines.append(("attachment", f"🎬 {get_clean_name(att_name)}"))
                                        # Charger les miniatures pour les intégrer dans la bulle
                                        video_thumbs = []
                                        for thumb_url in thumbnails[:8]:
                                            thumb_path = UPLOAD_DIR / thumb_url.replace("/uploads/", "")
                                            if thumb_path.exists():
                                                thumb_img = Image.open(thumb_path)
                                                thumb_img = thumb_img.convert("RGBA")
                                                # Taille reduite (etait 180) pour bulles video plus compactes
                                                max_thumb_width = 130 * RESOLUTION_SCALE
                                                ratio = max_thumb_width / thumb_img.width
                                                new_h = int(thumb_img.height * ratio)
                                                thumb_img = thumb_img.resize((max_thumb_width, new_h), Image.Resampling.LANCZOS)
                                                video_thumbs.append(thumb_img)
                                        # Deux composites distincts :
                                        # 1) EMBED (bulle) : RGBA + spacing (couleur bulle visible).
                                        # 2) DOCX (insertion separee) : RGBA + spacing=0 (thumbs colles,
                                        #    plus de blanc sur la page Word).
                                        if video_thumbs:
                                            thumb_w = video_thumbs[0].width
                                            thumb_h = video_thumbs[0].height
                                            cols = min(4, len(video_thumbs))
                                            rows = (len(video_thumbs) + cols - 1) // cols

                                            # --- Composite EMBED (bulle) ---
                                            spacing = 8 * RESOLUTION_SCALE
                                            embed_w = cols * thumb_w + (cols - 1) * spacing
                                            embed_h = rows * thumb_h + (rows - 1) * spacing
                                            composite = Image.new('RGBA', (embed_w, embed_h), (0, 0, 0, 0))
                                            for idx, thumb in enumerate(video_thumbs):
                                                row = idx // cols
                                                col = idx % cols
                                                x = col * (thumb_w + spacing)
                                                y = row * (thumb_h + spacing)
                                                composite.paste(thumb, (x, y))
                                            embedded_images.append(composite)

                                            # --- Composite DOCX (insertion separee, spacing reduit) ---
                                            # Petit espace visible entre miniatures, mais bien moins
                                            # que dans la bulle (8) pour eviter les grandes bandes blanches
                                            # sur la page Word.
                                            spacing_docx = 2 * RESOLUTION_SCALE
                                            docx_w = cols * thumb_w + (cols - 1) * spacing_docx
                                            docx_h = rows * thumb_h + (rows - 1) * spacing_docx
                                            composite_docx = Image.new('RGBA', (docx_w, docx_h), (0, 0, 0, 0))
                                            for idx, thumb in enumerate(video_thumbs):
                                                row = idx // cols
                                                col = idx % cols
                                                x = col * (thumb_w + spacing_docx)
                                                y = row * (thumb_h + spacing_docx)
                                                composite_docx.paste(thumb, (x, y))
                                            # Stocker en attribut transitoire ; sera sauvegarde sur disque
                                            # dans le bloc suivant et reference via composite_path.
                                            _composite_docx_for_save = composite_docx
                                        else:
                                            _composite_docx_for_save = None
                                        # Stocker les infos vidéo (thumbs embarquees dans bulle PNG,
                                        # composite docx sauvegarde plus bas et reference via composite_path)
                                        video_attachments.append({
                                            "name": att_name,
                                            "url": att_url,
                                            "path": str(video_path),
                                            "thumbnails": thumbnails,
                                            "_docx_composite_pending": _composite_docx_for_save,
                                        })
                                        if verbose: print(f"[MSG IMAGES] Vidéo {att_name}: {len(thumbnails)} miniatures (4 colonnes)")
                                    else:
                                        content_lines.append(("attachment", f"🎬 {get_clean_name(att_name)}"))
                                except Exception as e:
                                    if verbose: print(f"[MSG IMAGES] Erreur miniatures vidéo {att_name}: {e}")
                                    content_lines.append(("attachment", f"🎬 {get_clean_name(att_name)}"))
                            else:
                                if verbose: print(f"[MSG IMAGES] Vidéo NON TROUVÉE: {video_path}")
                                content_lines.append(("attachment", f"🎬 {get_clean_name(att_name)}"))
                        else:
                            # Pas une image ni vidéo, afficher le nom
                            content_lines.append(("attachment", f"📎 {get_clean_name(att_name)}"))
                    else:
                        content_lines.append(("attachment", f"📎 {get_clean_name(att_name)}"))
            elif attachments:
                for att in attachments:
                    att_name = str(att).split("/")[-1].split("\\")[-1]
                    content_lines.append(("attachment", f"📎 {get_clean_name(att_name)}"))

            if timestamp and timestamp != "nan" and timestamp != "None":
                content_lines.append(("timestamp", timestamp))

            # Polices unifiees pour TOUTES les bulles
            msg_font_body = font_body
            msg_font_sender = font_sender
            msg_font_timestamp = font_timestamp
            msg_line_height = line_height
            msg_sender_height = sender_height
            msg_timestamp_height = timestamp_height
            msg_attachment_height = attachment_height
            msg_transcription_height = transcription_height

            # Re-wrap par pixels les lignes attachment/transcription/timestamp ajoutees
            # plus haut sans controle de largeur (noms de fichier longs, contacts
            # partages, transcriptions audio...). Les lignes sender/body sont deja
            # wrappees a l'ajout, ce passage est idempotent pour elles.
            _font_map_full = {
                "sender": msg_font_sender,
                "body": msg_font_body,
                "attachment": msg_font_timestamp,
                "transcription": msg_font_body,
                "timestamp": msg_font_timestamp,
            }
            _rewrapped_full = []
            for _lt, _lx in content_lines:
                _f = _font_map_full.get(_lt, msg_font_body)
                for _sub in _wrap_by_pixels(_lx, _f, BUBBLE_TEXT_MAX_WIDTH):
                    _rewrapped_full.append((_lt, _sub))
            content_lines = _rewrapped_full

            # Largeur ADAPTATIVE: mesurer la ligne la plus large du contenu textuel
            # et ajouter le padding. La bulle s'etend entre MIN et MAX uniquement.
            content_max_px = 0
            for line_type, line_text in content_lines:
                if line_type == "sender":
                    f = msg_font_sender
                elif line_type == "timestamp":
                    f = msg_font_timestamp
                else:
                    f = msg_font_body
                try:
                    w = int(f.getlength(line_text))
                except Exception:
                    w = len(line_text) * 8 * RESOLUTION_SCALE
                if w > content_max_px:
                    content_max_px = w

            if embedded_images:
                # Pour les bulles avec miniatures video, le plafond grimpe a MAX_BUBBLE_WIDTH_VIDEO
                max_embedded_width = max(emb.width for emb in embedded_images)
                content_max_px = max(content_max_px, max_embedded_width)
                plafond = MAX_BUBBLE_WIDTH_VIDEO
            else:
                plafond = MAX_BUBBLE_WIDTH
            bubble_width = max(MIN_BUBBLE_WIDTH, min(plafond, content_max_px + bubble_padding * 2))
            bubble_height = bubble_padding * 2
            for line_type, line_text in content_lines:
                if line_type == "sender":
                    bubble_height += msg_sender_height
                elif line_type == "timestamp":
                    bubble_height += msg_timestamp_height
                elif line_type == "attachment":
                    bubble_height += msg_attachment_height
                elif line_type == "transcription":
                    bubble_height += msg_transcription_height
                else:
                    bubble_height += msg_line_height

            # Ajouter la hauteur des images embedded (dans la bulle - vidéos miniatures UNIQUEMENT)
            for emb_img in embedded_images:
                bubble_height += emb_img.height + 8 * RESOLUTION_SCALE

            # NOTE: Les images et PDF sont maintenant insérés séparément dans Word
            # pour permettre à l'utilisateur de les redimensionner

            # Créer l'image avec FOND TRANSPARENT (RGBA) - SANS MARGE
            img_width = bubble_width
            img_height = bubble_height

            bubble_color = owner_color if is_owner else other_color

            # Position de la bulle (collée aux bords, pas de marge)
            bubble_x = 0
            bubble_y = 0

            # Méthode avec masque pour coins parfaitement transparents
            # 1. Créer un masque avec la forme arrondie (blanc = opaque, noir = transparent)
            mask = Image.new('L', (img_width, img_height), 0)
            mask_draw = ImageDraw.Draw(mask)
            mask_draw.rounded_rectangle([bubble_x, bubble_y, img_width, img_height], radius=bubble_radius, fill=255)

            # 2. Créer l'image de la bulle (couleur unie)
            bubble_layer = Image.new('RGB', (img_width, img_height), bubble_color)

            # 3. Créer l'image finale avec transparence en utilisant le masque
            img = Image.new('RGBA', (img_width, img_height), (0, 0, 0, 0))
            img.paste(bubble_layer, (0, 0), mask)

            draw = ImageDraw.Draw(img)

            # Dessiner le contenu (polices adaptées selon taille de bulle)
            text_y = bubble_y + bubble_padding
            for line_type, line_text in content_lines:
                text_x = bubble_x + bubble_padding
                if line_type == "sender":
                    _draw_text_safe(draw,(text_x, text_y), line_text, fill=text_color, font=msg_font_sender)
                    text_y += msg_sender_height
                elif line_type == "timestamp":
                    _draw_text_safe(draw,(text_x, text_y), line_text, fill=(230, 230, 230), font=msg_font_timestamp)
                    text_y += msg_timestamp_height
                elif line_type == "attachment":
                    _draw_text_safe(draw,(text_x, text_y), line_text, fill=(220, 220, 220), font=msg_font_timestamp)
                    text_y += msg_attachment_height
                elif line_type == "transcription":
                    # Transcription en blanc, même taille que le texte normal
                    _draw_text_safe(draw,(text_x, text_y), line_text, fill=(255, 255, 255), font=msg_font_body)
                    text_y += msg_transcription_height
                else:
                    _draw_text_safe(draw,(text_x, text_y), line_text, fill=text_color, font=msg_font_body)
                    text_y += msg_line_height

            # Dessiner les images embedded (dans la bulle - miniatures vidéo UNIQUEMENT)
            for emb_img in embedded_images:
                img_x = bubble_x + bubble_padding
                img.paste(emb_img, (img_x, text_y), emb_img if emb_img.mode == 'RGBA' else None)
                text_y += emb_img.height + 8 * RESOLUTION_SCALE

            # NOTE: Les images et PDF sont insérés séparément dans Word (voir image_attachments et pdf_attachments)

            # Sauvegarder en PNG avec transparence et qualité maximale + DPI élevé pour Word
            safe_name = re.sub(r'[^\w\-]', '_', contact_name)
            msg_filename = f"msg_{safe_name}_{base_id}_{msg_idx + 1}.png"
            msg_path = UPLOAD_DIR / msg_filename
            # Sauvegarder avec DPI=300 pour haute qualité dans Word
            from PIL import PngImagePlugin
            png_info = PngImagePlugin.PngInfo()
            img.save(msg_path, "PNG", compress_level=1, dpi=(300, 300))

            # === BULLE CLEAN pour le BACKEND DOCX (toujours generee) ===
            # Calibri 16 partout, largeur FIXE 7 cm, polices preservees.
            # Les miniatures video sont desormais INTEGREES dans la bulle CLEAN
            # (et non plus inserees separement apres la bulle dans le docx).
            clean_path_str = None

            # 1) Recuperer les composites video (spacing serre 4 px) prepares pour
            # le docx et les redimensionner pour tenir dans la largeur utile de la
            # bulle clean (sans toucher aux polices texte).
            clean_embedded_videos = []
            if embedded_images:
                for _va in video_attachments:
                    cand = _va.pop("_docx_composite_pending", None)
                    if cand is None:
                        continue
                    if cand.width > CLEAN_BUBBLE_TEXT_MAX_WIDTH:
                        new_w = CLEAN_BUBBLE_TEXT_MAX_WIDTH
                        new_h = int(cand.height * (new_w / cand.width))
                        cand = cand.resize((new_w, new_h), Image.LANCZOS)
                    clean_embedded_videos.append(cand)

            # 2) Construire content_lines_clean avec wrap par largeur de pixels.
            # Garantit qu'aucun mot/URL ne sera tronque par le masque arrondi de la bulle.
            content_lines_clean = []
            if from_user:
                for _line in _wrap_by_pixels(from_user, font_sender_clean, CLEAN_BUBBLE_TEXT_MAX_WIDTH):
                    content_lines_clean.append(("sender", _line))
            if body and body.strip() and body != "None":
                for _line in _wrap_by_pixels(body, font_body_clean, CLEAN_BUBBLE_TEXT_MAX_WIDTH):
                    content_lines_clean.append(("body", _line))
            # Re-attacher les lignes attachment/transcription/timestamp en les wrappant
            # par largeur de pixels (noms de fichier longs, contacts partages, transcriptions).
            _font_map_clean = {
                "attachment": font_timestamp_clean,
                "transcription": font_body_clean,
                "timestamp": font_timestamp_clean,
            }
            for _lt, _lx in content_lines:
                if _lt in ("attachment", "transcription", "timestamp"):
                    _f = _font_map_clean[_lt]
                    for _sub in _wrap_by_pixels(_lx, _f, CLEAN_BUBBLE_TEXT_MAX_WIDTH):
                        content_lines_clean.append((_lt, _sub))

            # 3) Calculer la hauteur (largeur FIXE = CLEAN_BUBBLE_WIDTH)
            bubble_height_clean = bubble_padding * 2
            for _lt, _lx in content_lines_clean:
                if _lt == "sender": bubble_height_clean += CLEAN_SENDER_HEIGHT
                elif _lt == "timestamp": bubble_height_clean += CLEAN_TIMESTAMP_HEIGHT
                elif _lt == "attachment": bubble_height_clean += CLEAN_ATTACHMENT_HEIGHT
                elif _lt == "transcription": bubble_height_clean += CLEAN_TRANSCRIPTION_HEIGHT
                else: bubble_height_clean += CLEAN_LINE_HEIGHT
            # Ajouter la hauteur des miniatures video dans la bulle CLEAN
            for _emb in clean_embedded_videos:
                bubble_height_clean += _emb.height + 8 * RESOLUTION_SCALE
            bubble_width_clean = CLEAN_BUBBLE_WIDTH

            # 4) Rendre la bulle clean
            mask_c = Image.new('L', (bubble_width_clean, bubble_height_clean), 0)
            ImageDraw.Draw(mask_c).rounded_rectangle([0, 0, bubble_width_clean, bubble_height_clean], radius=bubble_radius, fill=255)
            layer_c = Image.new('RGB', (bubble_width_clean, bubble_height_clean), bubble_color)
            img_c = Image.new('RGBA', (bubble_width_clean, bubble_height_clean), (0, 0, 0, 0))
            img_c.paste(layer_c, (0, 0), mask_c)
            draw_c = ImageDraw.Draw(img_c)

            ty_c = bubble_padding
            for _lt, _lx in content_lines_clean:
                tx_c = bubble_padding
                if _lt == "sender":
                    _draw_text_safe(draw_c, (tx_c, ty_c), _lx, fill=text_color, font=font_sender_clean)
                    ty_c += CLEAN_SENDER_HEIGHT
                elif _lt == "timestamp":
                    _draw_text_safe(draw_c, (tx_c, ty_c), _lx, fill=(230, 230, 230), font=font_timestamp_clean)
                    ty_c += CLEAN_TIMESTAMP_HEIGHT
                elif _lt == "attachment":
                    _draw_text_safe(draw_c, (tx_c, ty_c), _lx, fill=(220, 220, 220), font=font_timestamp_clean)
                    ty_c += CLEAN_ATTACHMENT_HEIGHT
                elif _lt == "transcription":
                    _draw_text_safe(draw_c, (tx_c, ty_c), _lx, fill=(255, 255, 255), font=font_body_clean)
                    ty_c += CLEAN_TRANSCRIPTION_HEIGHT
                else:
                    _draw_text_safe(draw_c, (tx_c, ty_c), _lx, fill=text_color, font=font_body_clean)
                    ty_c += CLEAN_LINE_HEIGHT
            # Coller les miniatures video dans la bulle CLEAN (apres les lignes de texte)
            for _emb in clean_embedded_videos:
                img_c.paste(_emb, (bubble_padding, ty_c), _emb if _emb.mode == 'RGBA' else None)
                ty_c += _emb.height + 8 * RESOLUTION_SCALE
            try:
                clean_filename = f"msg_{safe_name}_{base_id}_{msg_idx + 1}_clean.png"
                clean_disk_path = UPLOAD_DIR / clean_filename
                img_c.save(clean_disk_path, "PNG", compress_level=1, dpi=(300, 300))
                clean_path_str = str(clean_disk_path)
            except Exception as _e:
                print(f"[CLEAN BUBBLE] Erreur sauvegarde: {_e}")

            # Les miniatures video sont maintenant DANS la bulle CLEAN, donc on
            # ne propage plus composite_path pour eviter une double insertion
            # cote remplace_rapport.py.

            # Ajouter aux résultats
            result.append({
                "image_path": str(msg_path),  # bulle FULL (frontend preview)
                "image_path_clean": clean_path_str,  # bulle CLEAN sans miniatures (backend docx)
                "image_url": f"/uploads/{msg_filename}",
                "is_owner": is_owner,
                "from_user": from_user,
                "body": body if body != "None" else "",
                "timestamp": timestamp if timestamp != "nan" else "",
                "attachment_urls": attachment_urls,
                "video_attachments": video_attachments,  # composite_path inclus si applicable
                "image_attachments": image_attachments,  # Images originales (insertion séparée dans Word)
                "pdf_attachments": pdf_attachments,  # PDF en images (insertion séparée dans Word)
                "comment": "",
                "comment_position": "left" if is_owner else "right"  # owner=vert à droite, commentaire à gauche
            })

        gen_elapsed = time_mod.perf_counter() - gen_start
        rate = len(result) / gen_elapsed if gen_elapsed > 0 else 0
        print(f"[MSG IMAGES] ✓ {len(result)} images générées pour {contact_name} en {gen_elapsed:.1f}s ({rate:.1f} img/s)")
        return result

    except Exception as e:
        print(f"[MSG IMAGES ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return []


@app.post("/api/import-excel", dependencies=[Depends(require_api_key)])
async def import_excel(file: UploadFile = File(...)):
    """
    Import Excel - SYSTÈME STREAMING HAUTE PERFORMANCE
    ==================================================
    Pipeline:
    1. Upload streaming (pas de chargement complet en mémoire)
    2. Parsing XML avec lxml iterparse
    3. Conversion multi-parquet par UUID
    4. Analyses via Polars lazy (scan_parquet)
    5. Génération graphiques
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Aucun fichier fourni")

    ext = Path(file.filename).suffix.lower()
    if ext != '.xlsx':
        raise HTTPException(status_code=400, detail="Format non supporté. Utilisez uniquement .xlsx")

    try:
        global _current_import, _current_analyzer, _last_platforms_data, _owner_usernames
        start_time = time.perf_counter()

        # Nettoyage des fichiers du dump précédent
        cleanup_result = cleanup_previous_import()
        print(f"[IMPORT] Nettoyage: {cleanup_result['data_cleaned']} dossiers data, {cleanup_result['uploads_cleaned']} graphiques")

        # Lecture du fichier (streaming pour gros fichiers)
        content = await file.read()
        file_size_mb = len(content) / (1024 * 1024)

        if file_size_mb > MAX_UPLOAD_SIZE_MB:
            raise HTTPException(
                status_code=413,
                detail=f"Fichier trop volumineux ({file_size_mb:.1f} MB). Maximum: {MAX_UPLOAD_SIZE_MB} MB"
            )

        print(f"[STREAMING] Fichier reçu: {file_size_mb:.1f} MB")

        # PHASE 1: Traitement streaming complet (XML -> Parquet)
        parse_start = time.perf_counter()
        result = process_excel_streaming(content)
        _current_import = result
        _current_analyzer = LazyAnalyzer(result.import_path)
        parse_time = time.perf_counter() - parse_start

        print(f"[STREAMING] Import {result.import_id}: {len(result.sheets)} feuilles en {parse_time:.2f}s")
        print(f"[STREAMING] Stockage: {result.import_path}")

        # PHASE 2: Mapping Device Info -> Placeholders Word
        device_data = result.device_info
        mapping: Dict[str, str] = {}

        # IMEI
        imei1 = device_data.get("IMEI") or device_data.get("IMEI1") or "N/C"
        imei2 = device_data.get("IMEI2") or "N/C"
        mapping["{imei}"] = imei1
        mapping["{imei1}"] = imei1
        mapping["{imei2}"] = imei2

        # Autres infos device
        modele_val = device_data.get("Detected Phone Model") or device_data.get("Model") or "N/C"

        # PRIORITÉ MARQUE: Vendor en premier, sinon déduction depuis le modèle
        vendor_val = device_data.get("Vendor") or ""

        # Si Vendor est vide ou générique, essayer de déduire depuis le modèle
        if not vendor_val or vendor_val.strip() == "" or vendor_val.lower() in ["android", "ios", ""]:
            model_lower = modele_val.lower() if modele_val else ""

            # Liste de correspondances modèle -> marque
            brand_patterns = [
                (["iphone", "ipad", "ipod", "macbook", "apple watch", "airpods"], "Apple"),
                (["samsung", "galaxy"], "Samsung"),
                (["xiaomi", "redmi", "poco", " mi "], "Xiaomi"),
                (["huawei", "mate 2", "mate 3", "mate 4", "nova"], "Huawei"),
                (["honor"], "Honor"),
                (["oneplus", "one plus"], "OnePlus"),
                (["google", "pixel", "nexus"], "Google"),
                (["oppo", "find x", "reno"], "Oppo"),
                (["vivo"], "Vivo"),
                (["realme"], "Realme"),
                (["motorola", "moto g", "moto e", "moto x"], "Motorola"),
                (["nokia", "lumia"], "Nokia"),
                ([" lg ", "lg-", "lg v", "lg g"], "LG"),
                (["sony", "xperia"], "Sony"),
                (["asus", "zenfone", "rog phone"], "Asus"),
                (["nothing phone"], "Nothing"),
                (["fairphone"], "Fairphone"),
                (["tcl"], "TCL"),
                (["zte", "blade"], "ZTE"),
                (["alcatel"], "Alcatel"),
                (["htc"], "HTC"),
                (["blackberry"], "BlackBerry"),
                (["lenovo"], "Lenovo"),
                (["wiko"], "Wiko"),
            ]
            for patterns, brand in brand_patterns:
                if any(p in model_lower for p in patterns):
                    vendor_val = brand
                    break

        mapping["{marque}"] = vendor_val if vendor_val else "N/C"
        mapping["{modele}"] = modele_val
        mapping["{vers}"] = device_data.get("OS Version") or device_data.get("OS") or "N/C"
        mapping["{stockage}"] = device_data.get("Storage") or device_data.get("Total Storage") or ""

        # Numéro de série: Serial en priorité, sinon Factory Number (insensible à la casse)
        serial_val = device_data.get("Serial") or ""
        if not serial_val:
            # Recherche insensible à la casse pour "Factory Number"
            for key in device_data.keys():
                if key.lower() == "factory number":
                    serial_val = device_data.get(key) or ""
                    break
        serial_val = serial_val or "N/C"
        mapping["{num}"] = serial_val
        mapping["{numseries}"] = serial_val  # Numéro de série appareil pour test3.docx
        mapping["{devicefactoryreset}"] = device_data.get("DeviceInfoFactoryReset") or "N/C"

        # SIM: IMSI et ICCID depuis Device Info
        mapping["{imsi}"] = device_data.get("IMSI") or "N/C"
        mapping["{iccid}"] = device_data.get("ICCID") or "N/C"

        # MSISDN depuis Sim Data
        msisdn_val = "N/C"
        if _current_analyzer:
            msisdn_val = _current_analyzer.get_sim_msisdn() or "N/C"
        mapping["{msisdn}"] = msisdn_val

                # PHASE 3: Statistiques via Polars LAZY (scan_parquet)
        stats_start = time.perf_counter()
        analyzer = _current_analyzer

        # Contacts
        nb_contacts, source_counts = analyzer.get_contacts_stats()
        mapping["{nbcontact}"] = str(nb_contacts)

        # Appels
        nb_appels, calls_source_counts = analyzer.get_calls_stats()
        mapping["{nbappel}"] = str(nb_appels)

        # Chats
        nb_messages, nb_conversations, chats_source_counts = analyzer.get_chats_stats()
        mapping["{nbmessage}"] = str(nb_messages)
        mapping["{nbconversation}"] = str(nb_conversations)

        # User Accounts
        nb_accounts, accounts_source_counts = analyzer.get_accounts_stats()
        mapping["{nbcompte}"] = str(nb_accounts)

        # Comptages simples
        mapping["{nbdocument}"] = str(analyzer.count_rows("Document"))
        # Images, Videos: compter uniquement ceux avec DCIM dans le Path
        mapping["{nbimage}"] = str(analyzer.count_media_dcim("Images"))
        mapping["{nbvideo}"] = str(analyzer.count_media_dcim("Videos") or analyzer.count_media_dcim("Video"))
        # Audio: compter uniquement ceux dont le Path contient une application installée
        mapping["{nbaudio}"] = str(analyzer.count_audio_with_installed_apps())
        mapping["{nblien}"] = str(analyzer.count_rows("Web Bookmarks"))

        stats_time = time.perf_counter() - stats_start
        print(f"[LAZY] Statistiques calculées en {stats_time*1000:.1f}ms")

                # PHASE 4: Génération des graphiques
        chart_contacts_path = generate_contacts_chart(source_counts)
        chart_calls_path = generate_calls_chart(calls_source_counts)
        chart_chats_path = generate_chats_chart(chats_source_counts)

        # Plateformes pour sous-titres auto
        # Pour accounts: "Accounts" (Compte utilisateur) toujours en premier (2.2.1)
        accounts_list = [p[0] for p in sorted(accounts_source_counts.items(), key=lambda x: x[1], reverse=True)] if accounts_source_counts else []
        if "Accounts" in accounts_list:
            accounts_list.remove("Accounts")
            accounts_list.insert(0, "Accounts")

        # Filtrer les plateformes vides
        accounts_list = [a for a in accounts_list if a and a.strip()]

        # Collecter les usernames du propriétaire depuis User Accounts
        _owner_usernames = set()
        try:
            user_accounts_sheet = None
            for sheet in result.sheets:
                if "user" in sheet.lower() and "account" in sheet.lower():
                    user_accounts_sheet = sheet
                    break

            if user_accounts_sheet and analyzer:
                ua_cols = analyzer.get_columns(user_accounts_sheet)
                if "Username" in ua_cols:
                    ua_data = analyzer.get_rows(user_accounts_sheet, ["Username"], max_rows=500)
                    for row in ua_data:
                        username = str(row.get("Username", "") or "").strip()
                        if username:
                            # Nettoyer le username (enlever @s.whatsapp.net, etc.)
                            clean_username = re.sub(r'@s\.whatsapp\.net$', '', username, flags=re.IGNORECASE)
                            clean_username = re.sub(r'[\s\-\.\+]', '', clean_username).lower()
                            if clean_username:
                                _owner_usernames.add(clean_username)
                    print(f"[IMPORT] Owner usernames collectés: {len(_owner_usernames)}")
        except Exception as e:
            print(f"[IMPORT] Erreur collecte owner usernames: {e}")

        _last_platforms_data = {
            "contacts": [p[0] for p in sorted(source_counts.items(), key=lambda x: x[1], reverse=True) if p[0] and p[0].strip()] if source_counts else [],
            "calls": [p[0] for p in sorted(calls_source_counts.items(), key=lambda x: x[1], reverse=True) if p[0] and p[0].strip()] if calls_source_counts else [],
            "chats": [p[0] for p in sorted(chats_source_counts.items(), key=lambda x: x[1], reverse=True) if p[0] and p[0].strip() and p[0] not in EXCLUDED_CHAT_TOP15_SOURCES] if chats_source_counts else [],
            "accounts": accounts_list
        }

        elapsed_time = time.perf_counter() - start_time
        print(f"[STREAMING] Plateformes: {_last_platforms_data}")
        print(f"[STREAMING] Stats: Contacts={nb_contacts}, Appels={nb_appels}, Messages={nb_messages}")
        print(f"[STREAMING] TOTAL: {elapsed_time:.2f}s")

        # Sauvegarder la session pour persistance
        save_session()

        return JSONResponse({
            "success": True,
            "import_id": result.import_id,
            "filename": file.filename,
            "file_size_mb": round(file_size_mb, 2),
            "mapping": mapping,
            "charts": {
                "contacts": chart_contacts_path,
                "calls": chart_calls_path,
                "chats": chart_chats_path
            },
            "platforms": _last_platforms_data,
            "sheets": result.sheets,
            "row_counts": result.row_counts,
            "processing_time_s": round(elapsed_time, 2),
            "message": f"Import streaming réussi: {len(result.sheets)} feuilles, {len(mapping)} valeurs"
        })

    except HTTPException:
        raise
    except Exception as e:
        print(f"[STREAMING ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


# Stockage des plateformes détectées par catégorie (pour création auto des sous-titres)
_last_platforms_data: Dict[str, List[str]] = {
    "contacts": [],
    "calls": [],
    "chats": [],
    "accounts": []
}


# ===== ENDPOINTS DE PERSISTANCE DE SESSION =====

@app.get("/api/session/status", dependencies=[Depends(require_api_key)])
async def get_session_status():
    """Retourne le statut de la session courante."""
    global _current_import, _tagged_data, _last_platforms_data

    has_import = _current_import is not None
    has_tagged = bool(_tagged_data.get("images") or _tagged_data.get("videos") or
                      _tagged_data.get("audios") or _tagged_data.get("conversations"))

    session_file_exists = SESSION_FILE.exists()
    session_timestamp = None

    if session_file_exists:
        try:
            with open(SESSION_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                session_timestamp = data.get("timestamp")
        except:
            pass

    return {
        "has_active_import": has_import,
        "has_tagged_data": has_tagged,
        "import_id": _current_import.import_id if _current_import else None,
        "sheets": _current_import.sheets if _current_import else [],
        "platforms": _last_platforms_data,
        "session_file_exists": session_file_exists,
        "session_timestamp": session_timestamp,
        "tagged_stats": {
            "images": len(_tagged_data.get("images", [])),
            "videos": len(_tagged_data.get("videos", [])),
            "audios": len(_tagged_data.get("audios", [])),
            "conversations": len(_tagged_data.get("conversations", {})),
        }
    }


@app.post("/api/session/restore", dependencies=[Depends(require_api_key)])
async def api_restore_session():
    """Restaure la session depuis le fichier sauvegardé."""
    success = restore_session()

    if success:
        return {
            "success": True,
            "message": "Session restaurée avec succès",
            "import_id": _current_import.import_id if _current_import else None,
            "sheets": _current_import.sheets if _current_import else [],
            "platforms": _last_platforms_data,
        }
    else:
        return {
            "success": False,
            "message": "Aucune session à restaurer ou erreur lors de la restauration"
        }


@app.post("/api/session/save", dependencies=[Depends(require_api_key)])
async def api_save_session():
    """Sauvegarde la session courante."""
    success = save_session()
    return {
        "success": success,
        "message": "Session sauvegardée" if success else "Erreur lors de la sauvegarde"
    }


def get_all_sheets_fast(content: bytes) -> List[str]:
    """Liste toutes les feuilles d'un fichier Excel"""
    import zipfile
    import xml.etree.ElementTree as ET

    sheets = []
    with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
        workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
        wb_root = ET.fromstring(workbook_xml)
        ns = {'main': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}

        for sheet in wb_root.findall('.//main:sheet', ns):
            sheet_name = sheet.get('name')
            if sheet_name and sheet_name not in SKIP_SHEETS:
                sheets.append(sheet_name)

    return sheets


def clean_excel_value(value: str) -> str:
    """Nettoie les caractères spéciaux XML des valeurs Excel (_x000D_, etc.)"""
    if not value:
        return value
    import re
    # Supprimer _x000D_, _x000d_, _x000A_, etc. (caractères de contrôle encodés)
    cleaned = re.sub(r'_x[0-9A-Fa-f]{4}_', '', value)
    # Supprimer les retours chariot et normaliser les espaces
    cleaned = cleaned.replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')
    # Supprimer les espaces multiples
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def get_sheet_columns_fast(content: bytes, sheet_name: str) -> List[str]:
    """Récupère les en-têtes de colonnes (ligne 2) d'une feuille Excel - version robuste avec ElementTree"""
    import zipfile
    import xml.etree.ElementTree as ET

    columns = []
    ns = {'': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
    ns_main = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Lire workbook.xml pour trouver l'ID de la feuille
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            wb_root = ET.fromstring(workbook_xml)

            sheet_id = None
            for sheet in wb_root.iter(f'{{{ns_main}}}sheet'):
                if sheet.get('name') == sheet_name:
                    sheet_id = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    break

            if not sheet_id:
                print(f"[COLUMNS] Feuille '{sheet_name}' non trouvée")
                return columns

            # 2. Lire les relations pour trouver le fichier XML de la feuille
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            rels_root = ET.fromstring(rels_xml)

            sheet_file = None
            for rel in rels_root.iter():
                if rel.get('Id') == sheet_id:
                    target = rel.get('Target', '')
                    if target.startswith('/'):
                        sheet_file = target[1:]
                    elif target.startswith('xl/'):
                        sheet_file = target
                    else:
                        sheet_file = 'xl/' + target
                    break

            if not sheet_file:
                print(f"[COLUMNS] Fichier feuille non trouvé")
                return columns

            # 3. Lire sharedStrings.xml
            shared_strings = []
            try:
                shared_strings_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                shared_root = ET.fromstring(shared_strings_xml)
                for si in shared_root.iter(f'{{{ns_main}}}si'):
                    text_parts = []
                    for t in si.iter(f'{{{ns_main}}}t'):
                        if t.text:
                            text_parts.append(t.text)
                    shared_strings.append(''.join(text_parts))
            except:
                pass

            # 4. Lire la feuille avec ElementTree
            sheet_xml = zf.read(sheet_file).decode('utf-8')
            sheet_root = ET.fromstring(sheet_xml)

            # Trouver la ligne 2
            row2 = None
            for row in sheet_root.iter(f'{{{ns_main}}}row'):
                if row.get('r') == '2':
                    row2 = row
                    break

            if row2 is None:
                print(f"[COLUMNS] Ligne 2 non trouvée")
                return columns

            # Extraire les colonnes de la ligne 2
            col_data = []
            for cell in row2.iter(f'{{{ns_main}}}c'):
                cell_ref = cell.get('r', '')
                cell_type = cell.get('t', '')

                col_letter = ''.join(c for c in cell_ref if c.isalpha())

                v_elem = cell.find(f'{{{ns_main}}}v')
                if v_elem is not None and v_elem.text:
                    raw_value = v_elem.text

                    if cell_type == 's' and shared_strings:
                        try:
                            idx = int(raw_value)
                            if idx < len(shared_strings):
                                col_name = clean_excel_value(shared_strings[idx])
                                if col_name:
                                    col_data.append((col_letter, col_name))
                        except ValueError:
                            pass
                    else:
                        col_name = clean_excel_value(raw_value)
                        if col_name:
                            col_data.append((col_letter, col_name))

            col_data.sort(key=lambda x: (len(x[0]), x[0]))
            columns = [name for _, name in col_data]
            print(f"[COLUMNS] {len(columns)} colonnes: {columns[:5]}...")

    except Exception as e:
        print(f"[COLUMNS ERROR] {type(e).__name__}: {e}")

    return columns


def get_sheet_row_count(content: bytes, sheet_name: str) -> int:
    """Compte le nombre total de lignes de données (sans les en-têtes) d'une feuille Excel"""
    import zipfile
    import xml.etree.ElementTree as ET

    ns_main = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver l'ID de la feuille
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            wb_root = ET.fromstring(workbook_xml)

            sheet_id = None
            for sheet in wb_root.iter(f'{{{ns_main}}}sheet'):
                if sheet.get('name') == sheet_name:
                    sheet_id = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    break

            if not sheet_id:
                return 0

            # 2. Trouver le fichier XML de la feuille
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            rels_root = ET.fromstring(rels_xml)

            sheet_file = None
            for rel in rels_root.iter():
                if rel.get('Id') == sheet_id:
                    target = rel.get('Target', '')
                    if target.startswith('/'):
                        sheet_file = target[1:]
                    elif target.startswith('xl/'):
                        sheet_file = target
                    else:
                        sheet_file = 'xl/' + target
                    break

            if not sheet_file:
                return 0

            # 3. Compter les lignes dans le XML (méthode rapide avec count)
            sheet_xml = zf.read(sheet_file).decode('utf-8')
            # Compter les balises <row - 2 (lignes d'en-tête: ligne 1 vide + ligne 2 en-têtes)
            row_count = sheet_xml.count('<row ')
            return max(0, row_count - 2)

    except Exception as e:
        print(f"[ROW_COUNT ERROR] {type(e).__name__}: {str(e)}")
        return 0


def get_sheet_row_count_filtered(content: bytes, sheet_name: str, source_filter: str) -> int:
    """Compte le nombre de lignes de données filtrées par source"""
    import zipfile
    import xml.etree.ElementTree as ET

    ns_main = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
    source_filter_norm = _normalize_source_text(source_filter)

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver l'ID de la feuille
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            wb_root = ET.fromstring(workbook_xml)

            sheet_id = None
            for sheet in wb_root.iter(f'{{{ns_main}}}sheet'):
                if sheet.get('name') == sheet_name:
                    sheet_id = sheet.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    break

            if not sheet_id:
                return 0

            # 2. Trouver le fichier XML de la feuille
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            rels_root = ET.fromstring(rels_xml)

            sheet_file = None
            for rel in rels_root.iter():
                if rel.get('Id') == sheet_id:
                    target = rel.get('Target', '')
                    if target.startswith('/'):
                        sheet_file = target[1:]
                    elif target.startswith('xl/'):
                        sheet_file = target
                    else:
                        sheet_file = 'xl/' + target
                    break

            if not sheet_file:
                return 0

            # 3. Lire sharedStrings.xml
            shared_strings = []
            try:
                shared_strings_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                shared_root = ET.fromstring(shared_strings_xml)
                for si in shared_root.iter(f'{{{ns_main}}}si'):
                    text_parts = []
                    for t in si.iter(f'{{{ns_main}}}t'):
                        if t.text:
                            text_parts.append(t.text)
                    shared_strings.append(''.join(text_parts))
            except:
                pass

            # 4. Lire la feuille
            sheet_xml = zf.read(sheet_file).decode('utf-8')
            sheet_root = ET.fromstring(sheet_xml)

            # 5. Trouver les colonnes d'en-tête (ligne 2) et identifier la colonne source
            source_col_idx = None
            source_col_candidates = ["source", "application", "app", "plateforme", "platform", "service", "provider"]

            row2 = None
            for row in sheet_root.iter(f'{{{ns_main}}}row'):
                if row.get('r') == '2':
                    row2 = row
                    break

            if row2 is None:
                return 0

            # Trouver quelle colonne contient "source" ou similaire
            for cell in row2.iter(f'{{{ns_main}}}c'):
                cell_ref = cell.get('r', '')
                col_letter = ''.join(c for c in cell_ref if c.isalpha())

                # Récupérer la valeur de la cellule
                cell_type = cell.get('t', '')
                v_elem = cell.find(f'{{{ns_main}}}v')
                if v_elem is not None and v_elem.text:
                    if cell_type == 's':
                        idx = int(v_elem.text)
                        if idx < len(shared_strings):
                            value = shared_strings[idx]
                        else:
                            value = ''
                    else:
                        value = v_elem.text
                else:
                    value = ''

                if _normalize_source_text(value) in source_col_candidates:
                    source_col_idx = col_letter
                    break

            if not source_col_idx:
                # Pas de colonne source trouvée, retourner le total sans filtre
                row_count = sheet_xml.count('<row ')
                return max(0, row_count - 2)

            # 6. Compter les lignes où la colonne source correspond au filtre
            count = 0
            for row in sheet_root.iter(f'{{{ns_main}}}row'):
                row_num = row.get('r', '0')
                if int(row_num) <= 2:  # Ignorer les lignes d'en-tête
                    continue

                for cell in row.iter(f'{{{ns_main}}}c'):
                    cell_ref = cell.get('r', '')
                    col_letter = ''.join(c for c in cell_ref if c.isalpha())

                    if col_letter == source_col_idx:
                        cell_type = cell.get('t', '')
                        v_elem = cell.find(f'{{{ns_main}}}v')
                        if v_elem is not None and v_elem.text:
                            if cell_type == 's':
                                idx = int(v_elem.text)
                                if idx < len(shared_strings):
                                    value = shared_strings[idx]
                                else:
                                    value = ''
                            else:
                                value = v_elem.text
                        else:
                            value = ''

                        if _normalize_source_text(value) == source_filter_norm:
                            count += 1
                        break

            return count

    except Exception as e:
        print(f"[ROW_COUNT_FILTERED ERROR] {type(e).__name__}: {str(e)}")
        return 0


def _normalize_source_text(value: str) -> str:
    """Normalise une valeur pour matching souple des plateformes/sources."""
    import re
    if value is None:
        return ""
    s = str(value).lower().replace("'", "'").replace("`", "'")
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def get_sheet_data_fast(
    content: bytes,
    sheet_name: str,
    selected_columns: List[str],
    max_rows: int = 100,
    source_filter: Optional[str] = None
) -> List[Dict[str, str]]:
    """Récupère les données d'une feuille Excel avec les colonnes sélectionnées - version ULTRA RAPIDE avec regex"""
    import zipfile
    import re

    data = []

    try:
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            # 1. Trouver la feuille via regex (rapide)
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')

            # Chercher l'ID de la feuille avec les deux formats possibles
            pattern1 = rf'<sheet[^>]*name="{re.escape(sheet_name)}"[^>]*r:id="(rId\d+)"'
            pattern2 = rf'<sheet[^>]*r:id="(rId\d+)"[^>]*name="{re.escape(sheet_name)}"'
            match = re.search(pattern1, workbook_xml) or re.search(pattern2, workbook_xml)

            if not match:
                print(f"[DATA FAST] Feuille '{sheet_name}' non trouvée")
                return data

            sheet_rid = match.group(1)

            # 2. Trouver le fichier XML via relations
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            match = re.search(rf'<Relationship[^>]*Id="{sheet_rid}"[^>]*Target="([^"]+)"', rels_xml)
            if not match:
                return data

            target = match.group(1)
            sheet_file = 'xl/' + target if not target.startswith('xl/') else target

            # 3. Charger sharedStrings avec regex (BEAUCOUP plus rapide qu'ElementTree)
            shared_strings = []
            try:
                ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
                si_pattern = r'<si[^>]*>(.*?)</si>'
                for si_match in re.finditer(si_pattern, ss_xml, re.DOTALL):
                    si_content = si_match.group(1)
                    texts = re.findall(r'<t[^>]*>([^<]*)</t>', si_content)
                    shared_strings.append(''.join(texts))
            except:
                pass

            # 4. Lire la feuille XML
            sheet_xml = zf.read(sheet_file).decode('utf-8')

            # 5. Extraire les en-têtes de la ligne 2 avec regex
            row2_match = re.search(r'<row[^>]*r="2"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL)
            if not row2_match:
                print(f"[DATA FAST] Ligne 2 non trouvée")
                return data

            row2_content = row2_match.group(1)
            col_to_letter = {}  # {col_name: col_letter}
            letter_to_col = {}  # {col_letter: col_name}

            # Pattern pour les cellules ligne 2: <c r="A2" t="s"><v>123</v></c>
            cell_pattern_row2 = r'<c\s+r="([A-Z]+)2"([^>]*)>(?:<v>([^<]*)</v>)?</c>'
            # Pattern pour les cellules données: <c r="A123" t="s"><v>val</v></c>
            cell_pattern_data = r'<c\s+r="([A-Z]+)\d+"([^>]*)>(?:<v>([^<]*)</v>)?</c>'

            for cell_match in re.finditer(cell_pattern_row2, row2_content):
                col_letter = cell_match.group(1)
                attrs = cell_match.group(2) or ""
                v_val = cell_match.group(3)

                if v_val is None:
                    continue

                col_name = None
                if 't="s"' in attrs and shared_strings:
                    try:
                        idx = int(v_val)
                        if idx < len(shared_strings):
                            col_name = clean_excel_value(shared_strings[idx])
                    except ValueError:
                        pass
                else:
                    col_name = clean_excel_value(v_val)

                if col_name and col_name in selected_columns:
                    col_to_letter[col_name] = col_letter
                    letter_to_col[col_letter] = col_name

            print(f"[DATA FAST] Colonnes mappées: {col_to_letter}")

            if not col_to_letter:
                print(f"[DATA FAST] Aucune colonne sélectionnée trouvée parmi: {selected_columns}")
                return data

            # 6. Extraire les données des lignes 3+ avec regex (arrêt précoce)
            row_pattern = r'<row[^>]*r="(\d+)"[^>]*>(.*?)</row>'
            row_count = 0

            for row_match in re.finditer(row_pattern, sheet_xml, re.DOTALL):
                row_num = int(row_match.group(1))
                if row_num <= 2:
                    continue
                if row_count >= max_rows:
                    break

                row_content = row_match.group(2)
                row_data = {col: "" for col in selected_columns}

                # Extraire les cellules de cette ligne
                for cell_match in re.finditer(cell_pattern_data, row_content):
                    col_letter = cell_match.group(1)

                    # Vérifier si c'est une colonne qu'on veut
                    if col_letter not in letter_to_col:
                        continue

                    col_name = letter_to_col[col_letter]
                    attrs = cell_match.group(2) or ""
                    v_val = cell_match.group(3)

                    if v_val is None:
                        continue

                    cell_value = v_val
                    if 't="s"' in attrs and shared_strings:
                        try:
                            idx = int(v_val)
                            if idx < len(shared_strings):
                                cell_value = shared_strings[idx]
                        except ValueError:
                            pass

                    row_data[col_name] = clean_excel_value(cell_value)

                if any(row_data.values()):
                    data.append(row_data)
                    row_count += 1

            print(f"[DATA FAST] {len(data)} lignes extraites")

    except Exception as e:
        print(f"[DATA FAST ERROR] {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        raise

    return data


@app.get("/api/excel-sheets", dependencies=[Depends(require_api_key)])
async def get_excel_sheets():
    """Retourne la liste des feuilles disponibles dans le dernier Excel importé"""
    global _current_import

    if _current_import is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé. Importez d'abord un fichier Excel.")

    return JSONResponse({
        "success": True,
        "sheets": _current_import.sheets,
        "import_id": _current_import.import_id
    })


@app.get("/api/sheet-preview/{sheet_name}", dependencies=[Depends(require_api_key)])
async def get_sheet_preview(sheet_name: str, max_rows: int = 50):
    """Retourne un aperçu des données d'une feuille avec toutes ses colonnes"""
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=404, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        sheet_name = unquote(sheet_name)

        # Récupérer toutes les colonnes de la feuille
        columns = _current_analyzer.get_columns(sheet_name)
        if not columns:
            raise HTTPException(status_code=404, detail=f"Feuille '{sheet_name}' non trouvée")

        # Récupérer les données
        data = _current_analyzer.get_sheet_data(sheet_name, columns, max_rows)

        # Décoder les entités HTML
        for row in data:
            for key in row:
                if isinstance(row[key], str):
                    row[key] = html.unescape(row[key])

        return JSONResponse({
            "success": True,
            "columns": columns,
            "data": data,
            "total_rows": len(data)
        })

    except HTTPException:
        raise
    except Exception as e:
        print(f"[SHEET-PREVIEW] Erreur: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/excel-columns/{sheet_name}", dependencies=[Depends(require_api_key)])
async def get_excel_columns(sheet_name: str, source_filter: Optional[str] = None):
    """Retourne les colonnes d'une feuille Excel et le nombre de lignes"""
    global _current_analyzer, _current_import

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        sheet_name = unquote(sheet_name)
        if source_filter:
            source_filter = unquote(source_filter)

        print(f"[LAZY COLUMNS] Récupération colonnes pour: '{sheet_name}' (filter: {source_filter})")

        # Récupérer colonnes non vides via LazyAnalyzer (exclut les colonnes avec uniquement N/C ou vide)
        columns = _current_analyzer.get_non_empty_columns(sheet_name, source_filter)
        all_columns = _current_analyzer.get_columns(sheet_name)
        filtered_out = len(all_columns) - len(columns)
        if filtered_out > 0:
            print(f"[LAZY COLUMNS] {filtered_out} colonnes vides masquées")

        # Nombre de lignes - filtré par source si demandé
        if source_filter:
            total_rows = _current_analyzer.count_rows_filtered(sheet_name, source_filter)
            print(f"[LAZY COLUMNS] {len(columns)} colonnes, {total_rows} lignes filtrées par '{source_filter}'")
        else:
            total_rows = _current_import.row_counts.get(sheet_name, 0)
            print(f"[LAZY COLUMNS] {len(columns)} colonnes, {total_rows} lignes")

        return JSONResponse({
            "success": True,
            "sheet": sheet_name,
            "columns": columns,
            "total_rows": total_rows,
            "source_filter": source_filter
        })
    except Exception as e:
        print(f"[LAZY COLUMNS ERROR] {type(e).__name__}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


@app.post("/api/excel-table-data", dependencies=[Depends(require_api_key)])
async def get_excel_table_data(request: Request):
    """
    Retourne les données d'une feuille avec les colonnes sélectionnées.
    Utilise POLARS LAZY sur Parquet (scan_parquet).
    """
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        body = await request.json()
        sheet_name = body.get("sheet")
        columns = body.get("columns", [])
        max_rows = body.get("max_rows", 100)
        source_filter = body.get("source_filter")
        filter_column = body.get("filter_column")
        filter_exclude = body.get("filter_exclude")
        connectivity_filter = body.get("connectivity_filter")  # Pour Bluetooth
        pre_filters = body.get("pre_filters", [])  # Pré-filtres par colonne

        if not sheet_name or not columns:
            raise HTTPException(status_code=400, detail="sheet et columns sont requis")

        start_time = time.perf_counter()

        # Pour le filtre Bluetooth, on doit charger la colonne Connectivity Method même si non sélectionnée
        columns_to_load = list(columns)
        connectivity_col_actual = None
        needs_connectivity_filter = False

        if connectivity_filter:
            # Trouver le nom exact de la colonne (insensible à la casse)
            all_sheet_cols = _current_analyzer.get_columns(sheet_name)
            for col in all_sheet_cols:
                if col.lower() == "connectivity method":
                    connectivity_col_actual = col
                    break

            if connectivity_col_actual:
                # Vérifier si la colonne est déjà sélectionnée
                needs_connectivity_filter = connectivity_col_actual.lower() not in [c.lower() for c in columns_to_load]
                if needs_connectivity_filter:
                    columns_to_load.append(connectivity_col_actual)

        # Utiliser LazyAnalyzer (scan_parquet)
        # Charger plus de données si des filtres sont actifs
        has_filters = (filter_column and filter_exclude) or connectivity_filter or pre_filters
        load_max = max_rows * 10 if has_filters else max_rows
        data = _current_analyzer.get_sheet_data(
            sheet_name,
            columns_to_load,
            load_max,
            source_filter=source_filter
        )

        # Filtre Connectivity Method (pour Bluetooth)
        if connectivity_filter and connectivity_col_actual:
            filter_val = connectivity_filter.lower()
            data = [row for row in data if filter_val in str(row.get(connectivity_col_actual, "")).lower()]
            # Supprimer la colonne si elle n'était pas dans les colonnes d'origine
            if needs_connectivity_filter:
                for row in data:
                    row.pop(connectivity_col_actual, None)

        # Appliquer les pré-filtres (filtre "in" sur des valeurs spécifiques ou "contains" pour texte partiel)
        if pre_filters:
            for pf in pre_filters:
                col = pf.get("column")
                op = pf.get("operator")
                if col and op == "in":
                    values = pf.get("values", [])
                    if values:
                        values_set = set(str(v) for v in values)
                        data = [row for row in data if str(row.get(col, "")) in values_set]
                elif col and op == "contains":
                    text = pf.get("text", "").lower()
                    if text:
                        data = [row for row in data if text in str(row.get(col, "")).lower()]

        # Filtrage côté serveur si demandé
        if filter_column and filter_exclude and filter_column in columns:
            exclude_lower = filter_exclude.lower()
            data = [row for row in data if exclude_lower not in str(row.get(filter_column, "")).lower()]

        # Limiter au max_rows demandé
        data = data[:max_rows]

        elapsed_ms = (time.perf_counter() - start_time) * 1000
        print(f"[LAZY TABLE] {sheet_name}: {len(data)} lignes en {elapsed_ms:.1f}ms")

        # Décoder les entités HTML dans les valeurs texte (&amp; → &)
        # Et nettoyer le HTML de la colonne Body pour les Emails
        is_emails = "email" in sheet_name.lower()
        for row in data:
            for key in row:
                if isinstance(row[key], str):
                    val = html.unescape(row[key])
                    # Nettoyer HTML pour colonne Body des Emails
                    if is_emails and key == "Body" and val:
                        val = strip_html_tags(val)
                    row[key] = val

        return JSONResponse({
            "success": True,
            "sheet": sheet_name,
            "columns": columns,
            "data": data,
            "row_count": len(data),
            "source_filter": source_filter
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


def parse_entries_field(entries: str) -> dict:
    """
    Parse la colonne Entries pour extraire User ID, Email, Phone.
    Format: "User ID-Instagram Id: 12345 Email-Email: test@mail.com Phone-Phone: +33612345678"
    """
    import re
    result = {
        "user_id": "",
        "email": "",
        "phone": ""
    }

    if not entries:
        return result

    # Patterns pour extraire les valeurs
    # User ID: cherche "User ID-XXX:" suivi de la valeur
    user_id_match = re.search(r'User ID-[^:]*:\s*([^\s]+(?:\s+[^\s]+)*?)(?=\s+[A-Z][a-z]+-|\s*$)', entries)
    if user_id_match:
        result["user_id"] = user_id_match.group(1).strip()

    # Email: cherche "Email-XXX:" suivi de la valeur
    email_match = re.search(r'Email-[^:]*:\s*([^\s]+@[^\s]+)', entries)
    if email_match:
        result["email"] = email_match.group(1).strip()

    # Phone: cherche "Phone-XXX:" suivi de la valeur
    phone_match = re.search(r'Phone-[^:]*:\s*([+\d\s-]+)', entries)
    if phone_match:
        result["phone"] = phone_match.group(1).strip()

    return result


@app.get("/api/user-account-data/{source}", dependencies=[Depends(require_api_key)])
async def get_user_account_data(source: str, service_type: Optional[str] = None):
    """
    Récupère les données de compte utilisateur pour une source donnée.
    Utilisé pour auto-remplir les phrases dans la section Comptes associés.
    """
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        source = unquote(source)
        if service_type:
            service_type = unquote(service_type)

        # Chercher la feuille User Accounts
        user_accounts_sheet = None
        for sheet in _current_import.sheets if _current_import else []:
            if "user" in sheet.lower() and "account" in sheet.lower():
                user_accounts_sheet = sheet
                break

        if not user_accounts_sheet:
            return JSONResponse({
                "success": False,
                "error": "Feuille User Accounts introuvable"
            })

        # Récupérer les colonnes disponibles
        cols = _current_analyzer.get_columns(user_accounts_sheet)

        # Colonnes Ã  récupérer (incluant Entries pour parsing)
        needed_cols = ["Username", "Service Type", "Creation time", "Account Name", "Source", "Entries"]
        available_cols = [c for c in needed_cols if c in cols]

        if not available_cols:
            return JSONResponse({
                "success": False,
                "error": "Colonnes requises non trouvées"
            })

        # Récupérer TOUTES les données filtrées par source (pas de limite)
        data = _current_analyzer.get_sheet_data(
            user_accounts_sheet,
            available_cols,
            max_rows=500,  # Augmenté pour avoir tous les comptes
            source_filter=source
        )

        # Filtrer par Service Type si demandé
        if service_type and data:
            data = [row for row in data if service_type.lower() in str(row.get("Service Type", "")).lower()]

        # Parser les données et formater les comptes
        formatted_accounts = []
        for idx, row in enumerate(data):
            entries_parsed = parse_entries_field(row.get("Entries", ""))
            username = row.get("Username", "") or ""
            account_name = row.get("Account Name", "") or ""

            # Pour l'email: utiliser Entries d'abord, puis Username si c'est un email
            email = entries_parsed["email"]
            if not email and "@" in username:
                email = username

            formatted_account = {
                "nb": idx + 1,
                "pseudouser": account_name or "N/C",  # Account Name = Pseudonyme (carnet d'adresses)
                "usernameuser": username or "N/C",    # Username = Nom d'utilisateur (username app)
                "identifiantuser": entries_parsed["user_id"] or "N/C",
                "mailuser": email or "N/C",
                "teluser": entries_parsed["phone"] or "N/C",
                "datesynchuser": row.get("Creation time", "") or "N/C",
                "datenaissanceuser": "N/C",  # Non disponible dans les données
                "serviceType": row.get("Service Type", ""),  # Ajout du Service Type
                # Données brutes pour debug
                "raw": row
            }
            formatted_accounts.append(formatted_account)

        result = {
            "success": True,
            "source": source,
            "count": len(formatted_accounts),
            "accounts": formatted_accounts
        }

        # Si source "Accounts", chercher le compte Apple ID principal
        if source.lower() == "accounts":
            apple_id_accounts = [r for r in data if str(r.get("Service Type", "")).strip() == "Apple ID"]
            apple_accounts = [r for r in data if "apple" in str(r.get("Service Type", "")).lower() or
                             str(r.get("Service Type", "")).strip() in ["iCloud", "iTunes Store", "Game Center"]]

            if apple_id_accounts:
                entries_parsed = parse_entries_field(apple_id_accounts[0].get("Entries", ""))
                result["primary_account"] = {
                    "Username": apple_id_accounts[0].get("Username", ""),
                    "Account Name": apple_id_accounts[0].get("Account Name", ""),
                    "Creation time": apple_id_accounts[0].get("Creation time", ""),
                    "email": entries_parsed["email"] or apple_id_accounts[0].get("Username", ""),
                    "Service Type": "Apple ID"
                }
            elif apple_accounts:
                entries_parsed = parse_entries_field(apple_accounts[0].get("Entries", ""))
                result["primary_account"] = {
                    "Username": apple_accounts[0].get("Username", ""),
                    "Account Name": apple_accounts[0].get("Account Name", ""),
                    "Creation time": apple_accounts[0].get("Creation time", ""),
                    "email": entries_parsed["email"] or apple_accounts[0].get("Username", ""),
                    "Service Type": apple_accounts[0].get("Service Type", "Apple")
                }
            else:
                # Priorité 3: Comptes Google (com.google)
                google_accounts = [r for r in data if str(r.get("Service Type", "")).strip() == "com.google"]

                if google_accounts:
                    # Supprimer les doublons basés sur Username
                    seen_usernames = set()
                    unique_google_accounts = []
                    for acc in google_accounts:
                        username = str(acc.get("Username", "")).strip()
                        if username and username not in seen_usernames:
                            seen_usernames.add(username)
                            unique_google_accounts.append(acc)

                    if unique_google_accounts:
                        # Si plusieurs comptes, concaténer les usernames
                        if len(unique_google_accounts) > 1:
                            usernames_list = [acc.get("Username", "") for acc in unique_google_accounts]
                            combined_username = ", ".join(usernames_list)
                        else:
                            combined_username = unique_google_accounts[0].get("Username", "")

                        entries_parsed = parse_entries_field(unique_google_accounts[0].get("Entries", ""))
                        result["primary_account"] = {
                            "Username": combined_username,
                            "Account Name": unique_google_accounts[0].get("Account Name", ""),
                            "Creation time": unique_google_accounts[0].get("Creation time", ""),
                            "email": entries_parsed["email"] or combined_username,
                            "Service Type": "com.google"
                        }
                elif data:
                    # Fallback: premier compte disponible
                    entries_parsed = parse_entries_field(data[0].get("Entries", ""))
                    result["primary_account"] = {
                        "Username": data[0].get("Username", ""),
                        "Account Name": data[0].get("Account Name", ""),
                        "Creation time": data[0].get("Creation time", ""),
                        "email": entries_parsed["email"] or data[0].get("Username", ""),
                        "Service Type": data[0].get("Service Type", "")
                    }

        return JSONResponse(result)

    except Exception as e:
        print(f"[USER ACCOUNT DATA ERROR] {type(e).__name__}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


@app.get("/api/call-log-data/{source}", dependencies=[Depends(require_api_key)])
async def get_call_log_data(source: str, generate_charts: bool = True):
    """
    Récupère les données du journal d'appels pour une source donnée.
    Retourne:
    - KPIs: nbappelentrant, nbappelsortant, timedebut, timefin
    - top15_count: top 15 contacts par nombre d'appels
    - top15_duration: top 15 contacts par durée
    - chart_count_path: chemin du graphique top 15 nombre (si generate_charts=True)
    - chart_duration_path: chemin du graphique top 15 durée (si generate_charts=True)
    """
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        source = unquote(source)

        # Récupérer les KPIs
        kpis = _current_analyzer.get_call_log_details_by_source(source)

        # Récupérer le top 15 par nombre d'appels
        top15_count = _current_analyzer.get_call_log_top15_by_count(source)

        # Récupérer le top 15 par durée
        top15_duration = _current_analyzer.get_call_log_top15_by_duration(source)

        result = {
            "success": True,
            "source": source,
            "kpis": kpis,
            "top15_count": top15_count,
            "top15_duration": top15_duration,
            "chart_count_path": None,
            "chart_duration_path": None
        }

        # Générer les graphiques si demandé
        if generate_charts:
            if top15_count:
                result["chart_count_path"] = generate_calls_top15_count_chart(top15_count, source)
            if top15_duration:
                result["chart_duration_path"] = generate_calls_top15_duration_chart(top15_duration, source)

        return JSONResponse(result)

    except Exception as e:
        print(f"[CALL LOG DATA ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


# ENDPOINTS GESTION DU STOCKAGE


@app.get("/api/call-log-summary/{source}", dependencies=[Depends(require_api_key)])
async def get_call_log_summary(source: str):
    """
    Récupère le tableau résumé des appels pour une source.
    Colonnes: Contact, Identifiant, Émis, Reçus, Appel vidéos, Appel manqué, Appel supprimé
    """
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        source = unquote(source)

        summary = _current_analyzer.get_call_log_summary(source)

        # Décoder les entités HTML dans les valeurs texte (&amp; → &)
        for row in summary:
            for key in row:
                if isinstance(row[key], str):
                    row[key] = html.unescape(row[key])

        return JSONResponse({
            "success": True,
            "source": source,
            "summary": summary,
            "count": len(summary)
        })

    except Exception as e:
        print(f"[CALL LOG SUMMARY ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


# ENDPOINTS CHATS / MESSAGERIES

@app.get("/api/chats-top15/{source}", dependencies=[Depends(require_api_key)])
async def get_chats_top15(source: str, generate_charts: bool = True):
    """
    Récupère les top 15 contacts par nombre de messages pour une source de chat.
    Retourne aussi le graphique si generate_charts=True.
    """
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        source = unquote(source)

        top15 = _current_analyzer.get_chats_top15_by_count(source)

        # Décoder les entités HTML dans les valeurs texte (&amp; → &)
        for row in top15:
            for key in row:
                if isinstance(row[key], str):
                    row[key] = html.unescape(row[key])

        result = {
            "success": True,
            "source": source,
            "top15": top15,
            "count": len(top15),
            "chart_count_path": None
        }

        # Générer le graphique si demandé
        if generate_charts and top15:
            result["chart_count_path"] = generate_chats_top15_count_chart(top15, source)

        return JSONResponse(result)

    except Exception as e:
        print(f"[CHATS TOP15 ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


@app.get("/api/chats-summary/{source}", dependencies=[Depends(require_api_key)])
async def get_chats_summary(source: str):
    """
    Récupère le tableau résumé des chats pour une source.
    Colonnes: Contact, Identifiant, Émis, Reçus, Total
    """
    global _current_analyzer

    if _current_analyzer is None:
        raise HTTPException(status_code=400, detail="Aucun fichier Excel importé.")

    try:
        from urllib.parse import unquote
        source = unquote(source)

        summary = _current_analyzer.get_chats_summary(source)

        # Décoder les entités HTML dans les valeurs texte (&amp; → &)
        for row in summary:
            for key in row:
                if isinstance(row[key], str):
                    row[key] = html.unescape(row[key])

        return JSONResponse({
            "success": True,
            "source": source,
            "summary": summary,
            "count": len(summary)
        })

    except Exception as e:
        print(f"[CHATS SUMMARY ERROR] {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erreur: {str(e)}")


@app.get("/api/imports", dependencies=[Depends(require_api_key)])
async def get_imports():
    """Liste tous les imports Parquet disponibles"""
    imports = list_imports()
    return JSONResponse({
        "success": True,
        "imports": imports,
        "current_import_id": _current_import.import_id if _current_import else None
    })


@app.delete("/api/imports/{import_id}", dependencies=[Depends(require_api_key)])
async def delete_import(import_id: str):
    """Supprime un import et ses fichiers Parquet"""
    global _current_import, _current_analyzer

    if cleanup_import(import_id):
        # Si c'est l'import courant, le réinitialiser
        if _current_import and _current_import.import_id == import_id:
            _current_import = None
            _current_analyzer = None

        return JSONResponse({
            "success": True,
            "message": f"Import {import_id} supprimé"
        })
    else:
        raise HTTPException(status_code=404, detail=f"Import {import_id} introuvable")


@app.post("/api/imports/cleanup", dependencies=[Depends(require_api_key)])
async def cleanup_old_imports(keep_current: bool = True):
    """Nettoie tous les imports sauf l'actuel (optionnel)"""
    global _current_import

    imports = list_imports()
    deleted = 0

    for imp in imports:
        if keep_current and _current_import and imp["import_id"] == _current_import.import_id:
            continue
        if cleanup_import(imp["import_id"]):
            deleted += 1

    return JSONResponse({
        "success": True,
        "deleted_count": deleted,
        "message": f"{deleted} import(s) supprimé(s)"
    })


@app.get("/api/import/current", dependencies=[Depends(require_api_key)])
async def get_current_import():
    """Retourne les informations sur l'import courant"""
    if _current_import is None:
        raise HTTPException(status_code=404, detail="Aucun import actif")

    return JSONResponse({
        "success": True,
        "import_id": _current_import.import_id,
        "path": str(_current_import.import_path),
        "sheets": _current_import.sheets,
        "row_counts": _current_import.row_counts,
        "device_info": _current_import.device_info
    })


# SYSTÈME TAGGED FOLDER (Dossiers tagués avec éléments sélectionnés)

TAGGED_DIR = Path("tagged_data")
TAGGED_DIR.mkdir(exist_ok=True)

TAGGED_PARQUET_DIR = Path("data/tagged_parquet")

# Stockage global des données tagguées
_tagged_data: Dict[str, Any] = {
    "images": [],
    "videos": [],
    "audios": [],
    "chats": [],
    "call_log": [],
    "conversations": {},  # {conversation_id: {info, messages}}
    "parquet_dir": None,  # Chemin vers le dossier des parquets
    "sheets_info": {},  # {sheet_name: {"columns": [...], "row_count": int, "parquet_path": str}}
}


def convert_tagged_excel_to_parquet(excel_path: Path) -> Dict[str, Any]:
    """
    Convertit les feuilles d'un fichier Excel tagué en fichiers Parquet.
    Utilise la même approche XML que parse_sheet_to_parquet_fast (excel_processor.py).
    Colonnes de la ligne 2 (header=1).
    Retourne: {"parquet_dir": Path, "sheets_info": {sheet_name: {...}}}
    """
    import shutil
    import uuid
    import polars as pl
    from datetime import datetime, timedelta

    # Regex pour parser le XML Excel (mêmes que excel_processor.py)
    RE_ROW = re.compile(r'<row[^>]*r="(\d+)"[^>]*>(.*?)</row>', re.DOTALL)
    RE_CELL = re.compile(r'<c\s+r="([A-Z]+)\d+"([^>]*)>(?:<v>([^<]*)</v>)?</c>')
    RE_CELL_TYPE_S = re.compile(r't="s"')
    RE_CLEAN = re.compile(r'_x[0-9A-Fa-f]{4}_')

    def clean_value(value: str) -> str:
        if not value:
            return ""
        value = RE_CLEAN.sub('', value)
        return value.strip()

    def col_letter_to_index(letter: str) -> int:
        result = 0
        for char in letter:
            result = result * 26 + (ord(char.upper()) - ord('A') + 1)
        return result - 1

    def excel_date_to_string(excel_date: float) -> str:
        """Convertit un numéro de série Excel en date formatée"""
        try:
            # Excel date: nombre de jours depuis 1899-12-30
            base_date = datetime(1899, 12, 30)
            delta = timedelta(days=float(excel_date))
            result_date = base_date + delta
            return result_date.strftime("%d/%m/%Y %H:%M:%S")
        except:
            return str(excel_date)

    def parse_shared_strings(zf) -> list:
        """Parse sharedStrings.xml"""
        try:
            ss_xml = zf.read('xl/sharedStrings.xml').decode('utf-8')
            strings = []
            for match in re.finditer(r'<si[^>]*>(.*?)</si>', ss_xml, re.DOTALL):
                texts = re.findall(r'<t[^>]*>([^<]*)</t>', match.group(1))
                strings.append(''.join(texts))
            return strings
        except:
            return []

    # Créer un dossier unique pour cet import
    import_id = str(uuid.uuid4())[:8]
    parquet_dir = TAGGED_PARQUET_DIR / f"tagged_{import_id}"

    # Nettoyer les anciens parquets des tags
    if TAGGED_PARQUET_DIR.exists():
        shutil.rmtree(TAGGED_PARQUET_DIR)
    parquet_dir.mkdir(parents=True, exist_ok=True)

    sheets_info = {}

    try:
        with open(excel_path, 'rb') as f:
            file_bytes = f.read()

        with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as zf:
            # Parser sharedStrings
            shared_strings = parse_shared_strings(zf)

            # Récupérer les feuilles
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')

            # Parser les relations pour trouver les fichiers des feuilles
            rels_xml = zf.read('xl/_rels/workbook.xml.rels').decode('utf-8')
            rels = {}
            for match in re.finditer(r'<Relationship[^>]*Id="(rId\d+)"[^>]*Target="([^"]+)"', rels_xml):
                rels[match.group(1)] = match.group(2)

            # Parser chaque feuille
            for sheet_match in re.finditer(r'<sheet[^>]*name="([^"]+)"[^>]*r:id="(rId\d+)"', workbook_xml):
                sheet_name = sheet_match.group(1)
                rel_id = sheet_match.group(2)

                if rel_id not in rels:
                    continue

                target = rels[rel_id]
                sheet_path = 'xl/' + target if not target.startswith('xl/') else target
                sheet_path = sheet_path.replace('xl/xl/', 'xl/')

                try:
                    sheet_xml = zf.read(sheet_path).decode('utf-8')

                    # Extraire headers (ligne 2)
                    headers = {}
                    header_match = re.search(r'<row[^>]*r="2"[^>]*>(.*?)</row>', sheet_xml, re.DOTALL)
                    if not header_match:
                        continue

                    for cell_match in RE_CELL.finditer(header_match.group(1)):
                        col_letter = cell_match.group(1)
                        attrs = cell_match.group(2) or ""
                        v_val = cell_match.group(3)

                        if v_val is None:
                            continue

                        if RE_CELL_TYPE_S.search(attrs) and shared_strings:
                            try:
                                idx = int(v_val)
                                if idx < len(shared_strings):
                                    headers[col_letter] = clean_value(shared_strings[idx])
                            except:
                                pass
                        else:
                            headers[col_letter] = clean_value(v_val)

                    if not headers:
                        continue

                    # Trier les colonnes
                    sorted_cols = sorted(headers.items(), key=lambda x: col_letter_to_index(x[0]))
                    col_names = [name if name else f"Col_{i}" for i, (_, name) in enumerate(sorted_cols)]
                    col_letters = [letter for letter, _ in sorted_cols]
                    letter_to_idx = {letter: i for i, letter in enumerate(col_letters)}

                    # Identifier les colonnes de type Date/Time
                    date_col_indices = set()
                    for i, col_name in enumerate(col_names):
                        col_lower = col_name.lower()
                        if 'date' in col_lower or 'time' in col_lower or 'timestamp' in col_lower:
                            date_col_indices.add(i)

                    # Collecter les données (lignes >= 3)
                    all_rows = []
                    for row_match in RE_ROW.finditer(sheet_xml):
                        row_num = int(row_match.group(1))
                        if row_num <= 2:
                            continue

                        row_data = [None] * len(col_names)
                        row_content = row_match.group(2)

                        for cell_match in RE_CELL.finditer(row_content):
                            col_letter = cell_match.group(1)
                            if col_letter not in letter_to_idx:
                                continue

                            col_idx = letter_to_idx[col_letter]
                            attrs = cell_match.group(2) or ""
                            v_val = cell_match.group(3)

                            if v_val is None:
                                continue

                            if RE_CELL_TYPE_S.search(attrs) and shared_strings:
                                try:
                                    idx = int(v_val)
                                    if idx < len(shared_strings):
                                        row_data[col_idx] = clean_value(shared_strings[idx])
                                except:
                                    pass
                            else:
                                # Convertir les dates si c'est une colonne Date/Time
                                if col_idx in date_col_indices:
                                    try:
                                        float_val = float(v_val)
                                        if 1 < float_val < 100000:  # Plage raisonnable pour les dates Excel
                                            row_data[col_idx] = excel_date_to_string(float_val)
                                        else:
                                            row_data[col_idx] = clean_value(v_val)
                                    except:
                                        row_data[col_idx] = clean_value(v_val)
                                else:
                                    row_data[col_idx] = clean_value(v_val)

                        all_rows.append(row_data)

                    if not all_rows:
                        continue

                    # Créer le DataFrame Polars
                    data_dict = {col_names[i]: [row[i] for row in all_rows] for i in range(len(col_names))}
                    schema = {col: pl.Utf8 for col in col_names}
                    df = pl.DataFrame(data_dict, schema=schema)

                    # Écrire en Parquet avec préfixe "tagged_" pour éviter conflits
                    parquet_filename = f"tagged_{sheet_name.replace(' ', '_').replace('/', '_')}.parquet"
                    parquet_path = parquet_dir / parquet_filename
                    df.write_parquet(parquet_path, compression="zstd", compression_level=1)

                    sheets_info[sheet_name] = {
                        "columns": col_names,
                        "row_count": len(df),
                        "parquet_path": str(parquet_path)
                    }

                    print(f"[TAGGED PARQUET] {sheet_name}: {len(df)} lignes -> {parquet_filename}")

                except Exception as e:
                    print(f"[TAGGED PARQUET] Erreur feuille {sheet_name}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue

    except Exception as e:
        print(f"[TAGGED PARQUET] Erreur globale: {e}")
        import traceback
        traceback.print_exc()

    return {
        "parquet_dir": str(parquet_dir),
        "sheets_info": sheets_info
    }


def detect_file_type(file_path: Path) -> str:
    """Détecte le type de fichier via magic bytes"""
    try:
        with open(file_path, 'rb') as f:
            header = f.read(12)

        # JPEG
        if header[:3] == b'\xff\xd8\xff':
            return 'image/jpeg'
        # PNG
        if header[:8] == b'\x89PNG\r\n\x1a\n':
            return 'image/png'
        # GIF
        if header[:6] in (b'GIF87a', b'GIF89a'):
            return 'image/gif'
        # WebP
        if header[:4] == b'RIFF' and header[8:12] == b'WEBP':
            return 'image/webp'
        # MP4/M4V
        if header[4:8] == b'ftyp':
            return 'video/mp4'
        # MOV
        if header[4:8] in (b'moov', b'mdat', b'wide', b'free'):
            return 'video/quicktime'
        # MP3
        if header[:3] == b'ID3' or header[:2] == b'\xff\xfb':
            return 'audio/mpeg'
        # AAC
        if header[:2] == b'\xff\xf1' or header[:2] == b'\xff\xf9':
            return 'audio/aac'

        return 'application/octet-stream'
    except:
        return 'application/octet-stream'


def get_file_extension(mime_type: str) -> str:
    """Retourne l'extension appropriée pour un type MIME"""
    extensions = {
        'image/jpeg': '.jpg',
        'image/png': '.png',
        'image/gif': '.gif',
        'image/webp': '.webp',
        'video/mp4': '.mp4',
        'video/quicktime': '.mov',
        'audio/mpeg': '.mp3',
        'audio/aac': '.aac',
    }
    return extensions.get(mime_type, '')


def parse_tagged_excel(excel_path: Path) -> Dict[str, Any]:
    """Parse le fichier Excel tagué et extrait les métadonnées avec Polars"""
    import polars as pl
    import io

    result = {
        "images": [],
        "videos": [],
        "audios": [],
        "chats": [],
        "call_log": [],
    }

    try:
        # Lire le fichier en mémoire d'abord pour éviter les problèmes de verrouillage Windows
        with open(excel_path, 'rb') as f:
            file_bytes = f.read()

        # Créer un BytesIO pour Polars
        file_buffer = io.BytesIO(file_bytes)

        # Lire les noms des feuilles avec openpyxl (juste pour la liste)
        import zipfile
        with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as zf:
            workbook_xml = zf.read('xl/workbook.xml').decode('utf-8')
            # Extraire les noms des feuilles
            sheet_names = re.findall(r'<sheet[^>]+name="([^"]+)"', workbook_xml)

        print(f"[TAGGED] Feuilles trouvées: {sheet_names}")

        # Parser chaque feuille pertinente
        for sheet_name in sheet_names:
            if sheet_name == 'Summary':
                continue

            try:
                # Reset buffer position
                file_buffer.seek(0)

                # Lire avec Polars (header à la ligne 1 = 2ème ligne, 0-indexed)
                df = pl.read_excel(
                    file_buffer,
                    sheet_name=sheet_name,
                    read_options={"header_row": 1}
                )

                # Convertir les noms de colonnes en strings
                columns = [str(c) for c in df.columns]

                def safe_get(row_dict, key, default=''):
                    """Récupère une valeur de façon sûre"""
                    val = row_dict.get(key, default)
                    if val is None:
                        return default
                    return str(val) if val != '' else default

                if sheet_name == 'Images':
                    for row in df.iter_rows(named=True):
                        result["images"].append({
                            "name": safe_get(row, 'Name'),
                            "path": safe_get(row, 'Path'),
                            "created_date": safe_get(row, 'Created-Date'),
                            "created_time": safe_get(row, 'Created-Time'),
                            "size": row.get('Size (bytes)', 0) or 0,
                            "meta_data": safe_get(row, 'Meta Data'),
                            "tag": safe_get(row, 'Tag'),
                        })

                elif sheet_name == 'Videos' or sheet_name == 'Video':
                    for row in df.iter_rows(named=True):
                        result["videos"].append({
                            "name": safe_get(row, 'Name'),
                            "path": safe_get(row, 'Path'),
                            "created_date": safe_get(row, 'Created-Date'),
                            "created_time": safe_get(row, 'Created-Time'),
                            "size": row.get('Size (bytes)', 0) or 0,
                            "meta_data": safe_get(row, 'Meta Data'),
                            "tag": safe_get(row, 'Tag'),
                        })

                elif sheet_name == 'Audios' or sheet_name == 'Audio':
                    for row in df.iter_rows(named=True):
                        result["audios"].append({
                            "name": safe_get(row, 'Name'),
                            "path": safe_get(row, 'Path'),
                            "created_date": safe_get(row, 'Created-Date'),
                            "created_time": safe_get(row, 'Created-Time'),
                            "size": row.get('Size (bytes)', 0) or 0,
                            "meta_data": safe_get(row, 'Meta Data'),
                            "tag": safe_get(row, 'Tag'),
                            "transcription": safe_get(row, 'Transcription'),  # Transcription audio
                        })

                elif sheet_name == 'Chats':
                    # Set pour détecter les doublons (basé sur chat_id + from + body + timestamp)
                    seen_chats = set()

                    for row in df.iter_rows(named=True):
                        # Extraire les infos de conversation
                        participants = safe_get(row, 'Participants')
                        source = safe_get(row, 'Source')
                        from_user = safe_get(row, 'From')
                        to_user = safe_get(row, 'To')
                        body = safe_get(row, 'Body')
                        timestamp_date = safe_get(row, 'Timestamp: Date')
                        timestamp_time = safe_get(row, 'Timestamp: Time')
                        chat_id = safe_get(row, 'Chat #')

                        # FILTRE 1: Ignorer les messages avec Body vide ET Attachment #1 - Details rempli
                        # MAIS garder si Attachment #1 (fichier) est rempli
                        body_is_empty = not body or str(body).strip() == '' or body == 'None'

                        # Vérifier uniquement Attachment #1 - Details
                        details_val = row.get('Attachment #1 - Details')
                        has_attachment_details = details_val is not None and str(details_val).strip() and str(details_val) != 'None'

                        # Vérifier uniquement Attachment #1 (fichier)
                        att_val_check = row.get('Attachment #1')
                        has_attachment_file = att_val_check is not None and str(att_val_check).strip() and str(att_val_check) != 'None'

                        # Ignorer SEULEMENT si: Body vide + Details rempli + PAS de fichier attachment
                        if body_is_empty and has_attachment_details and not has_attachment_file:
                            continue

                        # Chercher uniquement Attachment #1 (images et vidéos)
                        attachments = []
                        attachment_transcripts = {}  # {nom_attachment: transcription}

                        # Récupérer la colonne Transcript (transcription unique pour le message)
                        transcript_val = safe_get(row, 'Transcript')
                        has_transcript = transcript_val and str(transcript_val).strip() and str(transcript_val).lower() not in ('none', 'nan', '')

                        # Uniquement Attachment #1
                        att_val = row.get('Attachment #1')
                        if att_val is not None and str(att_val).strip() and str(att_val) != 'None':
                            attachments.append(str(att_val))
                            att_name = str(att_val).split('/')[-1].split('\\')[-1]
                            # Si c'est un audio et qu'il y a une transcription, l'associer
                            if has_transcript:
                                att_lower = att_name.lower()
                                if any(att_lower.endswith(ext) for ext in ['.opus', '.ogg', '.aac', '.m4a', '.wav', '.mp3', '.amr', '.wma']):
                                    attachment_transcripts[att_name] = str(transcript_val).strip()

                        # Pour la déduplication, utiliser seulement Attachment #1
                        attachment_for_dedup = str(att_val or '')

                        # FILTRE 2: Déduplication basée sur From, To, Body, Timestamp, Attachment #1
                        dedup_key = (
                            str(from_user or ''),
                            str(to_user or ''),
                            str(body or ''),
                            str(timestamp_time or ''),
                            str(timestamp_date or ''),
                            attachment_for_dedup  # Uniquement Attachment #1
                        )
                        if dedup_key in seen_chats:
                            # Doublon, ignorer
                            continue
                        seen_chats.add(dedup_key)

                        result["chats"].append({
                            "chat_id": chat_id,
                            "participants": participants,
                            "source": source,
                            "from": from_user,
                            "to": to_user,
                            "body": body if body != 'None' else '',
                            "attachments": attachments,
                            "attachment_transcripts": attachment_transcripts,  # Transcriptions des audios
                            "timestamp_date": timestamp_date,
                            "timestamp_time": timestamp_time,
                            "account": safe_get(row, 'Account'),
                        })

                elif sheet_name == 'Call Log':
                    # Collecter toutes les lignes d'abord
                    all_rows = []
                    for row in df.iter_rows(named=True):
                        # Support both column name formats
                        date_val = safe_get(row, 'Date') or safe_get(row, 'Timestamp: Date')
                        time_val = safe_get(row, 'Time') or safe_get(row, 'Timestamp: Time')
                        all_rows.append({
                            "parties": safe_get(row, 'Parties'),
                            "source": safe_get(row, 'Source'),
                            "type": safe_get(row, 'Type'),
                            "duration": safe_get(row, 'Duration'),
                            "direction": safe_get(row, 'Direction'),
                            "status": safe_get(row, 'Status'),
                            "video_call": safe_get(row, 'Video call'),
                            "account": safe_get(row, 'Account'),
                            "timestamp_date": date_val,
                            "timestamp_time": time_val,
                        })

                    # Déduplication: supprimer les entrées où seul Time diffère de 0-1 seconde
                    # MAIS garder les entrées où Duration = 0
                    def parse_time_to_seconds(time_str):
                        """Parse time string en secondes - supporte plusieurs formats:
                        - HH:MM:SS
                        - HH:MM
                        - DD/MM/YYYY HH:MM:SS(UTC+X)
                        - etc.
                        """
                        try:
                            if not time_str:
                                return 0
                            time_str = str(time_str)

                            # Extraire l'heure du format "17/11/2025 00:30:10(UTC+1)"
                            # Chercher le pattern HH:MM:SS
                            match = re.search(r'(\d{1,2}):(\d{2}):(\d{2})', time_str)
                            if match:
                                h, m, s = int(match.group(1)), int(match.group(2)), int(match.group(3))
                                return h * 3600 + m * 60 + s

                            # Chercher le pattern HH:MM
                            match = re.search(r'(\d{1,2}):(\d{2})', time_str)
                            if match:
                                h, m = int(match.group(1)), int(match.group(2))
                                return h * 3600 + m * 60

                            return 0
                        except:
                            return 0

                    def get_duration_seconds(duration_str):
                        """Parse duration pour vérifier si c'est 0"""
                        try:
                            if not duration_str or duration_str == '' or duration_str == 'None':
                                return 0
                            # Duration peut être "0:00:00" ou "00:00:05" etc
                            if ':' in str(duration_str):
                                parts = str(duration_str).split(':')
                                total = 0
                                for i, p in enumerate(parts):
                                    total += int(p) * (60 ** (len(parts) - 1 - i))
                                return total
                            return int(re.sub(r'[^0-9]', '', str(duration_str)) or 0)
                        except:
                            return 0

                    # Déduplication
                    seen = {}
                    deduplicated = []

                    for row in all_rows:
                        duration_sec = get_duration_seconds(row['duration'])

                        # Si Duration = 0, toujours garder (impossible de déterminer si doublon)
                        if duration_sec == 0:
                            deduplicated.append(row)
                            continue

                        # Créer clé de comparaison (colonnes: Parties, Date, Duration, Direction, Status, Video call, Source, Account)
                        key = (
                            row['parties'],
                            row['timestamp_date'],
                            row['duration'],
                            row['direction'],
                            row['status'],
                            row['video_call'],
                            row['source'],
                            row['account']
                        )

                        time_seconds = parse_time_to_seconds(row['timestamp_time'])

                        if key in seen:
                            # Vérifier si Time diffère de 0 ou 1 seconde
                            existing_time = seen[key]['time_seconds']
                            time_diff = abs(time_seconds - existing_time)
                            print(f"[CALL_LOG] Comparaison: {row['timestamp_time']} ({time_seconds}s) vs existant ({existing_time}s) = diff {time_diff}s")
                            if time_diff <= 1:
                                # Doublon (même clé + Time diffère de 0-1s) - ignorer cette entrée
                                print(f"[CALL_LOG] DOUBLON SUPPRIME: {row['parties']} - Time diff: {time_diff}s")
                                continue
                            else:
                                # Time différent de plus d'1 seconde - garder les deux
                                deduplicated.append(row)
                        else:
                            seen[key] = {'row': row, 'time_seconds': time_seconds}
                            deduplicated.append(row)

                    result["call_log"] = deduplicated
                    print(f"[CALL_LOG] Déduplication: {len(all_rows)} -> {len(deduplicated)} entrées (supprimé {len(all_rows) - len(deduplicated)} doublons)")

                print(f"[TAGGED] Feuille {sheet_name}: {len(df)} lignes")

            except Exception as sheet_error:
                print(f"[TAGGED] Erreur feuille {sheet_name}: {sheet_error}")
                continue

    except Exception as e:
        print(f"[TAGGED] Erreur parsing Excel: {e}")
        import traceback
        traceback.print_exc()

    return result


def extract_conversations(chats: List[Dict]) -> Dict[str, Dict]:
    """Extrait et groupe les conversations par participant (hors owner)"""
    global _owner_usernames
    conversations = {}

    for chat in chats:
        participants = chat.get('participants', '')
        source = chat.get('source', '')
        from_user = chat.get('from', '')
        account = chat.get('account', '')

        # Identifier le owner (utilisateur du smartphone) → bulle VERTE à droite
        # PRIORITÉ 1: Vérifier si from_user correspond à un username de User Accounts
        # PRIORITÉ 2: Si account est défini et contenu dans from → c'est le owner
        # PRIORITÉ 3: Sinon chercher "(owner)" dans participants
        is_owner = False

        # Nettoyer from_user pour comparaison
        from_clean = re.sub(r'[\s\-\.\+]', '', from_user).lower()
        from_clean = re.sub(r'@s\.whatsapp\.net$', '', from_clean, flags=re.IGNORECASE)

        # Méthode 1: Comparer avec les usernames de User Accounts (dump Excel)
        if _owner_usernames:
            # Vérifier si from_clean correspond à un owner username
            for owner_id in _owner_usernames:
                if owner_id in from_clean or from_clean in owner_id:
                    is_owner = True
                    break

        # Méthode 2: Utiliser la colonne Account si pas trouvé
        if not is_owner and account and account.strip():
            account_clean = re.sub(r'[\s\-\.\+]', '', account).lower()
            account_clean = re.sub(r'@s\.whatsapp\.net$', '', account_clean, flags=re.IGNORECASE)
            is_owner = account_clean in from_clean or from_clean in account_clean

        # Méthode 3: Chercher "(owner)" dans participants
        if not is_owner and participants:
            owner_name = ""
            parts = re.split(r'[,\n\r_x000d_]+', participants)
            for p in parts:
                p = p.strip()
                if "(owner)" in p.lower():
                    owner_name = re.sub(r'\s*\(owner\)\s*', '', p, flags=re.IGNORECASE).strip()
                    break
            if owner_name:
                owner_clean = re.sub(r'[\s\-\.\+]', '', owner_name).lower()
                is_owner = owner_clean in from_clean or from_clean in owner_clean

        # Extraire le nom du contact (non-owner) depuis participants
        contact_name = ""
        if participants:
            # Nettoyer _x000D_ (retour chariot Windows encodé) puis split par newline
            cleaned_participants = re.sub(r'_x000[dD]_', '\n', participants)
            parts = re.split(r'[\n\r,]+', cleaned_participants)
            for p in parts:
                p = p.strip()
                if p and "(owner)" not in p.lower():
                    contact_name = p
                    break

        if not contact_name:
            contact_name = "Inconnu"

        # Décoder les entités HTML (&amp; → &, etc.)
        contact_name = html.unescape(contact_name)

        # Créer une clé unique pour la conversation
        conv_key = f"{source}_{contact_name}".replace(" ", "_")

        if conv_key not in conversations:
            conversations[conv_key] = {
                "contact_name": contact_name,
                "source": source,
                "account": account,
                "messages": []
            }

        conversations[conv_key]["messages"].append({
            "from": from_user,
            "is_owner": is_owner,
            "body": chat.get('body', ''),
            "attachments": chat.get('attachments', []),
            "attachment_urls": chat.get('attachment_urls', []),
            "attachment_transcripts": chat.get('attachment_transcripts', {}),  # Transcriptions des audios
            "timestamp_date": chat.get('timestamp_date', ''),
            "timestamp_time": chat.get('timestamp_time', ''),
        })

    return conversations


@app.post("/api/import-tagged-folder", dependencies=[Depends(require_api_key)])
async def import_tagged_folder(file: UploadFile = File(...)):
    """
    Import d'un dossier tagué (ZIP contenant Excel + files/Image + files/Video + chats/)
    """
    global _tagged_data

    if not file.filename:
        raise HTTPException(status_code=400, detail="Aucun fichier fourni")

    ext = Path(file.filename).suffix.lower()
    if ext != '.zip':
        raise HTTPException(status_code=400, detail="Format non supporté. Utilisez un fichier ZIP")

    try:
        content = await file.read()

        # Nettoyer le dossier tagged_data
        import shutil
        if TAGGED_DIR.exists():
            shutil.rmtree(TAGGED_DIR)
        TAGGED_DIR.mkdir(exist_ok=True)

        # Extraire le ZIP - FILTRAGE: ne garder que chats/, files/ et *.xlsx
        # Ignore: AccountPackage/, FileUploads/, mailmessage/, Logs/, etc.
        with zipfile.ZipFile(io.BytesIO(content), 'r') as zf:
            extracted_count = 0
            skipped_count = 0

            for member in zf.namelist():
                member_lower = member.lower()

                # 1. Garder tous les fichiers Excel valides (hors fichiers verrou ~$)
                # NB: le filtre historique "nom contient 'Report'" excluait des ZIP
                # legitimes (ex: Conv_Whatsapp_Lakub.xlsx). Retire pour compatibilite
                # avec les exports WhatsApp / autres outils. Les exports Cellebrite
                # ne contiennent generalement qu'un seul xlsx, donc pas de risque.
                if member.endswith('.xlsx') and not member.split('/')[-1].startswith('~$'):
                    filename = member.split('/')[-1]
                    zf.extract(member, TAGGED_DIR)
                    extracted_count += 1
                    print(f"[IMPORT] Excel extrait: {filename}")
                    continue

                # 2. Garder uniquement chats/ et files/ (et leur contenu)
                is_chats = '/chats/' in member_lower or member_lower.endswith('/chats')
                is_files = '/files/' in member_lower or member_lower.endswith('/files')

                if is_chats or is_files:
                    zf.extract(member, TAGGED_DIR)
                    extracted_count += 1
                else:
                    skipped_count += 1

            print(f"[IMPORT] Extraction filtrée: {extracted_count} fichiers gardés, {skipped_count} ignorés")

        # Trouver le fichier Excel et les dossiers
        excel_file = None
        files_dir = None
        chats_dir = None

        for root, dirs, files in os.walk(TAGGED_DIR):
            for f in files:
                if f.endswith('.xlsx') and not f.startswith('~$'):
                    excel_file = Path(root) / f
            if 'files' in dirs:
                files_dir = Path(root) / 'files'
            if 'chats' in dirs:
                chats_dir = Path(root) / 'chats'

        if not excel_file:
            raise HTTPException(status_code=400, detail="Fichier Excel non trouvé dans le ZIP")

        # Parser le fichier Excel
        parsed_data = parse_tagged_excel(excel_file)

        # Copier les fichiers médias vers uploads avec conversion si nécessaire
        media_mapping = {}  # {original_name: new_path}
        # Mapping multi-candidats: {nom_base: [{url, source_path, preview_url}]}
        attachment_candidates = {}

        if files_dir:
            for media_type in ['Image', 'Video', 'Audio']:
                media_dir = files_dir / media_type
                if media_dir.exists():
                    for f in media_dir.iterdir():
                        if f.is_file():
                            # Détecter le type réel
                            mime_type = detect_file_type(f)
                            new_ext = get_file_extension(mime_type)

                            # Nom du fichier de destination
                            dest_name = f.name
                            if not f.suffix and new_ext:
                                dest_name = f.name + new_ext

                            dest_path = UPLOAD_DIR / dest_name
                            shutil.copy2(f, dest_path)
                            media_mapping[f.name] = f"/uploads/{dest_name}"

        # Copier les attachments depuis chats/[source]/[attachments]/
        # Avec gestion des doublons (même nom dans différents dossiers)
        if chats_dir and chats_dir.exists():
            print(f"[IMPORT] Recherche d'attachments dans chats: {chats_dir}")
            attachment_counter = {}  # Pour renommer les doublons

            for source_dir in chats_dir.iterdir():
                if source_dir.is_dir():
                    # Chercher les dossiers contenant "attachments" dans le nom
                    for sub_dir in source_dir.iterdir():
                        if sub_dir.is_dir() and 'attachment' in sub_dir.name.lower():
                            print(f"[IMPORT] Dossier attachments trouvé: {sub_dir}")
                            for f in sub_dir.iterdir():
                                if f.is_file():
                                    # Détecter le type réel
                                    mime_type = detect_file_type(f)
                                    new_ext = get_file_extension(mime_type)

                                    # Nom du fichier original (clé pour les candidats)
                                    original_name = f.name

                                    # Nom du fichier de destination avec extension
                                    base_dest_name = f.name
                                    if not f.suffix and new_ext:
                                        base_dest_name = f.name + new_ext

                                    # Gérer les doublons: ajouter un suffixe unique si le fichier existe
                                    dest_name = base_dest_name
                                    dest_path = UPLOAD_DIR / dest_name
                                    if dest_path.exists():
                                        # Fichier existe déjà, ajouter un suffixe unique
                                        if base_dest_name not in attachment_counter:
                                            attachment_counter[base_dest_name] = 1
                                        else:
                                            attachment_counter[base_dest_name] += 1

                                        # Créer un nom unique: nom_source_counter.ext
                                        name_part = dest_path.stem
                                        ext_part = dest_path.suffix
                                        source_id = source_dir.name[:20].replace(" ", "_")
                                        dest_name = f"{name_part}_{source_id}_{attachment_counter[base_dest_name]}{ext_part}"
                                        dest_path = UPLOAD_DIR / dest_name

                                    shutil.copy2(f, dest_path)
                                    url = f"/uploads/{dest_name}"
                                    media_mapping[f.name] = url

                                    # Ajouter aux candidats (pour sélection utilisateur si plusieurs)
                                    if original_name not in attachment_candidates:
                                        attachment_candidates[original_name] = []
                                    attachment_candidates[original_name].append({
                                        "url": url,
                                        "source_folder": source_dir.name,
                                        "original_path": str(f.relative_to(chats_dir))
                                    })

                                    print(f"[IMPORT] Attachment copié: {f.name} -> {dest_name} (source: {source_dir.name})")

        # Mettre à jour les chemins dans les données parsées
        for img in parsed_data["images"]:
            if img["name"] in media_mapping:
                img["url"] = media_mapping[img["name"]]
        for vid in parsed_data["videos"]:
            if vid["name"] in media_mapping:
                vid["url"] = media_mapping[vid["name"]]
        for aud in parsed_data["audios"]:
            if aud["name"] in media_mapping:
                aud["url"] = media_mapping[aud["name"]]

        # Mettre à jour les attachments dans les chats avec leurs URLs
        # Inclure tous les candidats si plusieurs fichiers ont le même nom
        print(f"[IMPORT] media_mapping contient {len(media_mapping)} fichiers")
        print(f"[IMPORT] attachment_candidates contient {len(attachment_candidates)} fichiers")
        print(f"[IMPORT] Traitement de {len(parsed_data['chats'])} chats")

        attachment_found_count = 0
        attachment_missing_count = 0

        for chat in parsed_data["chats"]:
            attachments = chat.get("attachments", [])
            attachment_urls = []
            for att in attachments:
                att_name = str(att).split("/")[-1].split("\\")[-1]
                found_url = None
                candidates = []

                # Vérifier s'il y a plusieurs candidats pour ce nom
                if att_name in attachment_candidates:
                    candidates = attachment_candidates[att_name]
                    if len(candidates) == 1:
                        found_url = candidates[0]["url"]
                    elif len(candidates) > 1:
                        # Plusieurs candidats: l'utilisateur devra choisir
                        found_url = candidates[0]["url"]  # URL par défaut (premier trouvé)
                        print(f"[IMPORT] Attachment avec {len(candidates)} candidats: {att_name}")

                # Fallback: chercher dans le mapping simple
                if not found_url and att_name in media_mapping:
                    found_url = media_mapping[att_name]

                # Fallback: chercher directement dans uploads
                if not found_url:
                    direct_path = UPLOAD_DIR / att_name
                    if direct_path.exists():
                        found_url = f"/uploads/{att_name}"
                        print(f"[IMPORT] Fichier trouvé directement: {att_name}")

                # Fallback: recherche partielle
                if not found_url:
                    for mapped_name, mapped_url in media_mapping.items():
                        if att_name in mapped_name or mapped_name in att_name:
                            found_url = mapped_url
                            print(f"[IMPORT] Correspondance partielle: {att_name} -> {mapped_name}")
                            break

                if found_url:
                    attachment_found_count += 1
                    attachment_urls.append({
                        "name": att_name,
                        "url": found_url,
                        "candidates": candidates if len(candidates) > 1 else None,  # Liste si plusieurs choix
                        "has_multiple": len(candidates) > 1
                    })
                else:
                    attachment_missing_count += 1
                    print(f"[IMPORT] Attachment NON TROUVÉ: {att_name}")
                    attachment_urls.append({
                        "name": att_name,
                        "url": None,
                        "candidates": None,
                        "has_multiple": False
                    })
            chat["attachment_urls"] = attachment_urls

        print(f"[IMPORT] Attachments: {attachment_found_count} trouvés, {attachment_missing_count} non trouvés")

        # Extraire les conversations
        conversations = extract_conversations(parsed_data["chats"])

        # Convertir les feuilles Excel en Parquet pour Tableau Tags+
        parquet_result = convert_tagged_excel_to_parquet(excel_file)

        # Stocker globalement
        _tagged_data = {
            "images": parsed_data["images"],
            "videos": parsed_data["videos"],
            "audios": parsed_data["audios"],
            "chats": parsed_data["chats"],
            "call_log": parsed_data["call_log"],
            "conversations": conversations,
            "attachment_candidates": attachment_candidates,  # Pour sélection si plusieurs fichiers même nom
            "excel_path": str(excel_file),  # Chemin vers l'Excel pour Tableau Tags+
            "parquet_dir": parquet_result["parquet_dir"],
            "sheets_info": parquet_result["sheets_info"],
        }

        # Sauvegarder la session pour persistance
        save_session()

        return JSONResponse({
            "success": True,
            "message": "Dossier tagué importé avec succès",
            "stats": {
                "images": len(parsed_data["images"]),
                "videos": len(parsed_data["videos"]),
                "audios": len(parsed_data["audios"]),
                "chats": len(parsed_data["chats"]),
                "call_log": len(parsed_data["call_log"]),
                "conversations": len(conversations),
            },
            "media_mapping": media_mapping,
        })

    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="Fichier ZIP invalide")
    except Exception as e:
        print(f"[TAGGED] Erreur import: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erreur lors de l'import: {str(e)}")


@app.get("/api/tagged-data", dependencies=[Depends(require_api_key)])
async def get_tagged_data():
    """Retourne toutes les données tagguées"""
    return JSONResponse({
        "success": True,
        "data": _tagged_data
    })


# Cache pour les miniatures vidéo des candidats (évite régénération à chaque appel)
_video_thumbnail_cache = {}

@app.get("/api/attachment-candidates/{attachment_name}", dependencies=[Depends(require_api_key)])
async def get_attachment_candidates(attachment_name: str):
    """Retourne tous les fichiers candidats pour un nom d'attachment donné.
    Utilisé quand plusieurs fichiers ont le même nom dans différents dossiers.
    Pour les vidéos, utilise un cache pour les miniatures."""
    from urllib.parse import unquote
    attachment_name = unquote(attachment_name)

    candidates = _tagged_data.get("attachment_candidates", {}).get(attachment_name, [])

    if not candidates:
        # Chercher par correspondance partielle
        for name, cands in _tagged_data.get("attachment_candidates", {}).items():
            if attachment_name in name or name in attachment_name:
                candidates = cands
                break

    # Vérifier si c'est une vidéo
    video_extensions = ['.mp4', '.mov', '.avi', '.mkv', '.webm', '.m4v', '.3gp']
    is_video = any(attachment_name.lower().endswith(ext) for ext in video_extensions)

    if is_video and candidates:
        # Utiliser le cache pour les miniatures vidéo
        for cand in candidates:
            try:
                video_url = cand.get("url", "")

                # Vérifier si déjà en cache
                if video_url in _video_thumbnail_cache:
                    cand["thumbnail"] = _video_thumbnail_cache[video_url]
                    continue

                # Vérifier si le candidat a déjà une miniature
                if cand.get("thumbnail"):
                    _video_thumbnail_cache[video_url] = cand["thumbnail"]
                    continue

                if video_url.startswith("/uploads/"):
                    video_path = UPLOAD_DIR / video_url.replace("/uploads/", "")
                    if video_path.exists():
                        # Générer 1 miniature au milieu de la vidéo
                        thumbnails = extract_video_thumbnails(video_path, num_thumbnails=1)
                        if thumbnails:
                            cand["thumbnail"] = thumbnails[0]
                            _video_thumbnail_cache[video_url] = thumbnails[0]
                            print(f"[CANDIDATES] Miniature générée et mise en cache: {video_path.name}")
            except Exception as e:
                print(f"[CANDIDATES] Erreur miniature pour {cand.get('url', '')}: {e}")

    return JSONResponse({
        "success": True,
        "attachment_name": attachment_name,
        "candidates": candidates,
        "count": len(candidates),
        "is_video": is_video
    })


@app.post("/api/update-message-attachment", dependencies=[Depends(require_api_key)])
async def update_message_attachment(request: Request):
    """
    Met à jour l'URL d'un attachment dans un message et régénère l'image de la bulle.
    Utilisé quand l'utilisateur sélectionne un fichier parmi plusieurs candidats.
    """
    from urllib.parse import unquote

    try:
        body = await request.json()
        conv_id = unquote(body.get("conv_id", ""))
        msg_idx = body.get("msg_idx", 0)
        attachment_name = body.get("attachment_name", "")
        new_url = body.get("new_url", "")

        if not conv_id or not attachment_name or not new_url:
            raise HTTPException(status_code=400, detail="Paramètres manquants")

        # Vérifier que la conversation existe
        if conv_id not in _tagged_data.get("conversations", {}):
            raise HTTPException(status_code=404, detail="Conversation non trouvée")

        conv = _tagged_data["conversations"][conv_id]
        messages = conv.get("messages", [])

        if msg_idx < 0 or msg_idx >= len(messages):
            raise HTTPException(status_code=400, detail="Index de message invalide")

        # Mettre à jour l'URL de l'attachment dans le message
        msg = messages[msg_idx]
        attachment_urls = msg.get("attachment_urls", [])

        print(f"[UPDATE ATTACH] conv_id={conv_id}, msg_idx={msg_idx}")
        print(f"[UPDATE ATTACH] attachment_name recherché: '{attachment_name}'")
        print(f"[UPDATE ATTACH] new_url: '{new_url}'")
        print(f"[UPDATE ATTACH] Nombre d'attachments dans le message: {len(attachment_urls)}")

        updated = False
        for i, att in enumerate(attachment_urls):
            att_current_name = att.get("name", "")
            print(f"[UPDATE ATTACH] Attachment {i}: name='{att_current_name}', match={att_current_name == attachment_name}")
            if att_current_name == attachment_name:
                old_url = att.get("url", "")
                att["url"] = new_url
                att["selected_url"] = new_url  # Marquer comme sélectionné
                updated = True
                print(f"[UPDATE ATTACH] ✓ URL mise à jour: '{old_url}' -> '{new_url}'")
                break

        if not updated:
            raise HTTPException(status_code=404, detail="Attachment non trouvé dans le message")

        # Régénérer l'image du message avec le nouvel attachment
        contact_name = conv.get("contact_name", "")
        source = conv.get("source", "")

        # Construire le dictionnaire des transcriptions audio
        audio_transcriptions = {}
        for audio in _tagged_data.get("audios", []):
            audio_name = audio.get("name", "")
            transcription = audio.get("transcription", "")
            if audio_name and transcription and str(transcription).strip() and str(transcription).lower() != 'nan':
                audio_transcriptions[audio_name] = str(transcription).strip()

        # Régénérer l'image pour ce message uniquement
        single_msg_result = generate_message_images([msg], contact_name, source, audio_transcriptions)

        if single_msg_result and len(single_msg_result) > 0:
            result = single_msg_result[0]
            new_image_url = result["image_url"]
            new_image_path = result["image_path"]
            image_attachments = result.get("image_attachments", [])
            pdf_attachments = result.get("pdf_attachments", [])
            video_attachments = result.get("video_attachments", [])

            print(f"[UPDATE ATTACH] Message {msg_idx} mis à jour avec {attachment_name} -> {new_url}")
            print(f"[UPDATE ATTACH] Nouvelle image: {new_image_url}")
            print(f"[UPDATE ATTACH] image_attachments: {len(image_attachments)}, pdf_attachments: {len(pdf_attachments)}")

            return JSONResponse({
                "success": True,
                "new_image_url": new_image_url,
                "new_image_path": new_image_path,
                "image_attachments": image_attachments,
                "pdf_attachments": pdf_attachments,
                "video_attachments": video_attachments
            })
        else:
            raise HTTPException(status_code=500, detail="Erreur lors de la régénération de l'image")

    except HTTPException:
        raise
    except Exception as e:
        print(f"[UPDATE ATTACH] Erreur: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/tagged-conversations", dependencies=[Depends(require_api_key)])
async def get_tagged_conversations():
    """Retourne la liste des conversations avec les pseudos des contacts"""
    conversations_list = []

    for conv_key, conv_data in _tagged_data.get("conversations", {}).items():
        # Décoder les entités HTML (&amp; → &)
        contact_name = html.unescape(conv_data["contact_name"]) if conv_data.get("contact_name") else conv_data.get("contact_name", "")
        conversations_list.append({
            "id": conv_key,
            "contact_name": contact_name,
            "source": conv_data["source"],
            "message_count": len(conv_data["messages"]),
        })

    return JSONResponse({
        "success": True,
        "conversations": conversations_list
    })


@app.get("/api/tagged-conversation/{conv_id}", dependencies=[Depends(require_api_key)])
async def get_tagged_conversation(conv_id: str):
    """Retourne les messages d'une conversation spécifique"""
    from urllib.parse import unquote
    conv_id = unquote(conv_id)

    if conv_id not in _tagged_data.get("conversations", {}):
        raise HTTPException(status_code=404, detail="Conversation non trouvée")

    return JSONResponse({
        "success": True,
        "conversation": _tagged_data["conversations"][conv_id]
    })


@app.get("/api/tagged-conversation-images/{conv_id}", dependencies=[Depends(require_api_key)])
async def get_tagged_conversation_images(conv_id: str):
    """Génère une image par message et retourne les infos pour ajout de commentaires"""
    from urllib.parse import unquote
    from datetime import datetime
    conv_id = unquote(conv_id)

    if conv_id not in _tagged_data.get("conversations", {}):
        raise HTTPException(status_code=404, detail="Conversation non trouvée")

    conv = _tagged_data["conversations"][conv_id]
    messages = conv.get("messages", [])
    contact_name = conv.get("contact_name", "")
    source = conv.get("source", "")

    # Construire le dictionnaire des transcriptions audio
    audio_transcriptions = {}
    for audio in _tagged_data.get("audios", []):
        audio_name = audio.get("name", "")
        transcription = audio.get("transcription", "")
        if audio_name and transcription and str(transcription).strip() and str(transcription).lower() != 'nan':
            audio_transcriptions[audio_name] = str(transcription).strip()
    print(f"[AUDIO TRANS] {len(audio_transcriptions)} transcriptions disponibles")

    # Générer les images individuelles
    message_images = generate_message_images(messages, contact_name, source, audio_transcriptions)

    # Fonction pour nettoyer les timestamps (enlever les nombres Excel comme 45842.85617064815)
    def clean_timestamp(ts):
        if not ts:
            return ts
        # Enlever les nombres décimaux Excel au début (ex: "45842.85617064815 04/07/2025")
        cleaned = re.sub(r'^\d+\.\d+\s*', '', str(ts)).strip()
        return cleaned if cleaned else ts

    # Calculer la période de communication (dates min et max)
    period_start = ""
    period_end = ""
    timestamps = []
    print(f"[CONV PERIOD] Nombre de messages: {len(messages)}")
    for i, msg in enumerate(messages):
        date_str = str(msg.get("timestamp_date", "") or "").strip()
        time_str = str(msg.get("timestamp_time", "") or "").strip()
        if i < 3:  # Debug premiers messages
            print(f"[CONV PERIOD] Msg {i}: date='{date_str}', time='{time_str}'")

        # Vérifier que ce n'est pas vide ou "nan"
        valid_date = date_str and date_str.lower() not in ("nan", "none", "")
        valid_time = time_str and time_str.lower() not in ("nan", "none", "")

        if valid_date and valid_time:
            # Les deux sont présents: combiner
            full_ts = f"{date_str} {time_str}"
            timestamps.append(full_ts)
        elif valid_date:
            # Seulement la date
            timestamps.append(date_str)
        elif valid_time:
            # Parfois le timestamp complet est dans time_str
            timestamps.append(time_str)

    print(f"[CONV PERIOD] Timestamps trouvés: {len(timestamps)}")
    if timestamps:
        print(f"[CONV PERIOD] Premier: '{timestamps[0]}', Dernier: '{timestamps[-1]}'")

    if timestamps:
        # Trier les timestamps pour trouver min et max
        try:
            # Essayer plusieurs formats de date
            parsed_dates = []
            for ts in timestamps:
                # Nettoyer le timestamp (enlever nombres Excel)
                clean_ts = clean_timestamp(ts)
                # Enlever aussi les infos UTC pour le parsing
                parse_ts = re.sub(r'\(UTC[+-]?\d*\)', '', clean_ts).strip()

                for fmt in ["%d/%m/%Y %H:%M:%S", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y", "%Y-%m-%d",
                           "%Y-%m-%d %H:%M", "%d/%m/%Y %H:%M"]:
                    try:
                        dt = datetime.strptime(parse_ts, fmt)
                        parsed_dates.append((dt, clean_ts))  # Garder le timestamp nettoyé
                        break
                    except:
                        continue

            if parsed_dates:
                parsed_dates.sort(key=lambda x: x[0])
                period_start = parsed_dates[0][1]
                period_end = parsed_dates[-1][1]
                print(f"[CONV PERIOD] Parsing réussi: {period_start} -> {period_end}")
            else:
                # Fallback: utiliser premier et dernier nettoyés
                period_start = clean_timestamp(timestamps[0])
                period_end = clean_timestamp(timestamps[-1]) if len(timestamps) > 1 else ""
                print(f"[CONV PERIOD] Fallback (pas de parsing): {period_start} -> {period_end}")
        except Exception as e:
            print(f"[CONV] Erreur parsing dates: {e}")
            # Fallback en cas d'erreur
            period_start = clean_timestamp(timestamps[0])
            period_end = clean_timestamp(timestamps[-1]) if len(timestamps) > 1 else ""

    # Si period_start == period_end, ne garder que period_start (un seul message ou même date)
    if period_start and period_end and period_start == period_end:
        period_end = ""

    # Lookup des informations de contact depuis la feuille Contacts (Excel)
    source_lower = source.lower() if source else ""
    # Signal est traité comme une source sociale (Name → nom_utilisateur, identifier → identifiant_utilisateur)
    is_phone_source = source_lower in ("native messages", "natif", "whatsapp", "whatsapp business")
    # Toutes les sources non-téléphoniques (Snapchat, Instagram, TikTok, Telegram, Signal, etc.) = format social
    is_social_source = not is_phone_source

    contact_info = {
        "pseudonyme": contact_name,
        "nom_utilisateur": "",
        "identifiant_utilisateur": "",
        "numero_telephone": "",
        "source": source,  # Garder la source originale
        "is_social_source": is_social_source,
        "is_phone_source": is_phone_source
    }

    # Essayer de trouver les infos du contact dans le fichier Excel importé
    is_signal = source_lower == "signal"

    print(f"[CONV CONTACT] Source: {source}, contact_name: '{contact_name}', is_social: {is_social_source}, is_phone: {is_phone_source}")

    if _current_analyzer:
        try:
            # Construire le lookup depuis Contacts pour TOUTES les sources (comme Excel get_chats_summary)
            contacts_lookup = _current_analyzer._build_contacts_user_id_lookup(source)
            print(f"[CONV LOOKUP] {len(contacts_lookup)} entrées dans le lookup pour {source}")

            # Parser contact_name pour extraire current_name et current_identifier (comme Excel)
            # Formats possibles: "+590690527166 Fabio", "12345 username Nom", "username Nom", "Nom", etc.
            contact_parts = contact_name.split() if contact_name else []
            current_name = ""
            current_identifier = ""

            if len(contact_parts) >= 2 and re.match(r'^(\+?\d{8,}|[\d@\.]+@\S+)$', contact_parts[0]):
                # Format: "téléphone/ID Nom" → identifier=phone, name=rest
                current_identifier = contact_parts[0]
                current_name = " ".join(contact_parts[1:])
            elif len(contact_parts) >= 3 and re.match(r'^\d{5,}$', contact_parts[0]):
                # Format Instagram: "ID_numérique username Nom..." → identifier=ID, name=username (premier mot après ID)
                current_identifier = contact_parts[0]
                current_name = contact_parts[1]  # Le username
            elif len(contact_parts) >= 2 and re.match(r'^\d{5,}$', contact_parts[0]):
                # Format: "ID_numérique username" → identifier=ID, name=username
                current_identifier = contact_parts[0]
                current_name = contact_parts[1]
            elif len(contact_parts) >= 2:
                # Format: "username Nom Prénom" ou "Nom Prénom" → identifier=premier, name=premier
                # Le premier pourrait être un username ou un nom
                current_identifier = contact_parts[0]
                current_name = contact_parts[0]  # Utiliser le premier comme les deux
            elif contact_name:
                # Format: "Nom" ou "username" seul
                current_name = contact_name
                current_identifier = contact_name

            # Nettoyer l'identifier (supprimer @s.whatsapp.net etc.)
            if current_identifier and "@" in current_identifier:
                current_identifier = current_identifier.split("@")[0]

            print(f"[CONV PARSED] current_name='{current_name}', current_identifier='{current_identifier}'")

            # Chercher dans le lookup (même logique que Excel get_chats_summary)
            lookup_info = None
            if current_identifier and current_identifier.lower() in contacts_lookup:
                lookup_info = contacts_lookup[current_identifier.lower()]
                print(f"[CONV LOOKUP] Trouvé par identifier: '{current_identifier}'")
            elif current_name and current_name.lower() in contacts_lookup:
                lookup_info = contacts_lookup[current_name.lower()]
                print(f"[CONV LOOKUP] Trouvé par name: '{current_name}'")
            else:
                # Essayer chaque partie du nom séparément
                for part in contact_parts:
                    part_clean = part.strip().lower()
                    if part_clean and part_clean in contacts_lookup:
                        lookup_info = contacts_lookup[part_clean]
                        print(f"[CONV LOOKUP] Trouvé par partie: '{part_clean}'")
                        break

            # Appliquer les résultats du lookup (même logique pour TOUTES les sources, comme Excel)
            if lookup_info:
                print(f"[CONV LOOKUP FOUND] Lookup info: {lookup_info}")
                contact_info["pseudonyme"] = lookup_info.get("pseudonyme", "") or "-"
                contact_info["nom_utilisateur"] = lookup_info.get("nom_utilisateur", "") or "-"
                contact_info["identifiant_utilisateur"] = lookup_info.get("identifiant_utilisateur", "")

                # Pour sources téléphoniques: utiliser numero_telephone si disponible
                if is_phone_source:
                    phone = lookup_info.get("numero_telephone") or lookup_info.get("identifiant_utilisateur") or current_identifier
                    # Nettoyer et formater le numéro
                    if phone and "@" in phone:
                        phone = phone.split("@")[0]
                    if phone and phone[0].isdigit():
                        phone = "+" + phone
                    contact_info["numero_telephone"] = phone

                # Pour Signal: pas de nom_utilisateur
                if is_signal:
                    contact_info["nom_utilisateur"] = "-"
            else:
                print(f"[CONV LOOKUP NOT FOUND] Pas trouvé dans Contacts pour '{contact_name}'")
                # Pas trouvé dans Contacts: utiliser current_name et current_identifier (comme Excel)

                if is_phone_source:
                    # Sources téléphoniques: Pseudonyme = nom, Numéro = identifier
                    contact_info["pseudonyme"] = current_name if current_name != current_identifier else "-"
                    phone = current_identifier
                    if phone and phone[0].isdigit():
                        phone = "+" + phone
                    contact_info["numero_telephone"] = phone
                elif is_signal:
                    # Signal: Pseudonyme = nom parsé, Identifiant = UUID/phone
                    signal_match = re.match(r'^([A-Fa-f0-9\-]{20,50}|\+?\d[\d\s\-\.]{8,})\s*(.*)$', contact_name)
                    if signal_match:
                        contact_info["identifiant_utilisateur"] = signal_match.group(1).strip()
                        contact_info["pseudonyme"] = signal_match.group(2).strip() or "-"
                    else:
                        contact_info["identifiant_utilisateur"] = current_identifier
                        contact_info["pseudonyme"] = current_name if current_name != current_identifier else "-"
                    contact_info["nom_utilisateur"] = "-"
                else:
                    # Sources sociales (comme Excel get_chats_summary):
                    # Pseudonyme = "-" (pas dans carnet d'adresses)
                    # Nom utilisateur = identifier (peut être username ou ID)
                    # Si identifier ressemble à un ID (UUID ou numérique long) → le mettre dans identifiant_utilisateur
                    contact_info["pseudonyme"] = "-"

                    # Fonction helper pour vérifier si c'est un ID utilisateur (comme Excel)
                    def looks_like_user_id(value):
                        if not value:
                            return False
                        # UUID pattern (avec ou sans tirets)
                        if re.match(r'^[0-9a-f]{8}-?[0-9a-f]{4}-?[0-9a-f]{4}-?[0-9a-f]{4}-?[0-9a-f]{12}$', value, re.IGNORECASE):
                            return True
                        # Numérique long (>= 10 chiffres)
                        if re.match(r'^\d{10,}$', value):
                            return True
                        return False

                    # Si current_identifier ressemble à un ID, le mettre dans identifiant_utilisateur
                    if looks_like_user_id(current_identifier):
                        contact_info["identifiant_utilisateur"] = current_identifier
                        contact_info["nom_utilisateur"] = current_name if current_name != current_identifier else "-"
                    else:
                        # Sinon, current_identifier est le nom_utilisateur (username)
                        contact_info["nom_utilisateur"] = current_identifier or current_name or "-"
                        contact_info["identifiant_utilisateur"] = ""
        except Exception as e:
            print(f"[CONV] Erreur lookup contacts: {e}")
            import traceback
            traceback.print_exc()

    # Décoder les entités HTML dans toutes les valeurs texte (&amp; → &, etc.)
    contact_name = html.unescape(contact_name) if contact_name else contact_name
    for key in contact_info:
        if isinstance(contact_info[key], str):
            contact_info[key] = html.unescape(contact_info[key])

    # Log final des valeurs de contact_info
    print(f"[CONV RESULT] contact_name='{contact_name}' -> pseudonyme='{contact_info.get('pseudonyme')}', "
          f"nom_utilisateur='{contact_info.get('nom_utilisateur')}', "
          f"identifiant_utilisateur='{contact_info.get('identifiant_utilisateur')}'")

    return JSONResponse({
        "success": True,
        "contact_name": contact_name,
        "source": source,
        "message_count": len(messages),
        "messages": message_images,
        "period_start": period_start,
        "period_end": period_end,
        "contact_info": contact_info
    })


@app.get("/api/tagged-images", dependencies=[Depends(require_api_key)])
async def get_tagged_images():
    """Retourne les images tagguées"""
    return JSONResponse({
        "success": True,
        "images": _tagged_data.get("images", [])
    })


@app.get("/api/tagged-videos", dependencies=[Depends(require_api_key)])
async def get_tagged_videos():
    """Retourne les vidéos tagguées"""
    return JSONResponse({
        "success": True,
        "videos": _tagged_data.get("videos", [])
    })


@app.get("/api/tagged-audios", dependencies=[Depends(require_api_key)])
async def get_tagged_audios():
    """Retourne les audios tagguées"""
    return JSONResponse({
        "success": True,
        "audios": _tagged_data.get("audios", [])
    })


@app.get("/api/tagged-sheets", dependencies=[Depends(require_api_key)])
async def get_tagged_sheets():
    """Retourne la liste des feuilles Excel et leurs colonnes depuis les données taguées (via Parquet)"""
    import polars as pl

    sheets_info = _tagged_data.get("sheets_info")

    if not sheets_info:
        return JSONResponse({
            "success": False,
            "error": "Aucun fichier Excel tagué disponible"
        })

    # Formater pour le frontend (sans le chemin parquet)
    # Filtrer les colonnes vides (uniquement null, "", ou "N/C")
    result = {}
    for sheet_name, info in sheets_info.items():
        all_columns = info.get("columns", [])
        parquet_path = info.get("parquet_path")

        # Filtrer les colonnes non vides si le fichier parquet existe
        non_empty_columns = []
        if parquet_path and Path(parquet_path).exists():
            try:
                df = pl.read_parquet(parquet_path)
                for col in all_columns:
                    if col not in df.columns:
                        continue
                    # Vérifier si la colonne a au moins une valeur non vide et non "N/C"
                    col_data = df[col]
                    has_valid_value = False
                    for val in col_data:
                        if val is not None:
                            str_val = str(val).strip()
                            if str_val and str_val.lower() not in ("n/c", "n\\c", "nan", "") and str_val != "None":
                                has_valid_value = True
                                break
                    if has_valid_value:
                        non_empty_columns.append(col)

                filtered_count = len(all_columns) - len(non_empty_columns)
                if filtered_count > 0:
                    print(f"[TAGGED SHEETS] {sheet_name}: {filtered_count} colonnes vides masquées")
            except Exception as e:
                print(f"[TAGGED SHEETS] Erreur lecture parquet pour {sheet_name}: {e}")
                non_empty_columns = all_columns
        else:
            non_empty_columns = all_columns

        result[sheet_name] = {
            "columns": non_empty_columns,
            "row_count": info.get("row_count", 0)
        }

    return JSONResponse({
        "success": True,
        "sheets": result
    })


@app.post("/api/tagged-table-data", dependencies=[Depends(require_api_key)])
async def get_tagged_table_data(request: Request):
    """
    Retourne les données d'une feuille Excel taguée avec filtrage (via Parquet).
    Body JSON: {sheet_name, columns, limit, filters: [{column, operator, value}]}
    """
    import polars as pl

    sheets_info = _tagged_data.get("sheets_info", {})

    if not sheets_info:
        return JSONResponse({
            "success": False,
            "error": "Aucun fichier Excel tagué disponible"
        })

    try:
        body = await request.json()
        sheet_name = body.get("sheet_name")
        columns = body.get("columns", [])
        limit = body.get("limit", 50)
        filters = body.get("filters", [])
        source_filter = body.get("source_filter")  # Filtre par source (ex: "WhatsApp")

        if not sheet_name:
            return JSONResponse({
                "success": False,
                "error": "Nom de feuille requis"
            })

        # Vérifier que la feuille existe
        if sheet_name not in sheets_info:
            return JSONResponse({
                "success": False,
                "error": f"Feuille '{sheet_name}' non trouvée"
            })

        parquet_path = sheets_info[sheet_name].get("parquet_path")
        if not parquet_path or not Path(parquet_path).exists():
            return JSONResponse({
                "success": False,
                "error": f"Fichier Parquet non trouvé pour '{sheet_name}'"
            })

        # Lire le Parquet avec Polars (lazy mode pour performance)
        lf = pl.scan_parquet(parquet_path)
        all_columns = sheets_info[sheet_name].get("columns", [])

        # Filtrer par source si spécifié (ex: "WhatsApp", "Signal", etc.) - AVANT la sélection de colonnes
        if source_filter:
            # Chercher une colonne "Source" ou similaire (plus de variations)
            source_col = None
            source_keywords = ["source", "application", "app", "platform", "plateforme", "origine"]
            for col in all_columns:
                col_lower = col.lower().replace(" ", "").replace("_", "")
                for keyword in source_keywords:
                    if keyword in col_lower:
                        source_col = col
                        break
                if source_col:
                    break

            if source_col:
                # Filtrer par source (comparaison exacte, insensible à la casse)
                source_lower = source_filter.lower()
                print(f"[TAGGED-TABLE] Filtrage par source: {source_filter} (colonne: {source_col})")
                lf = lf.filter(pl.col(source_col).cast(pl.Utf8).str.to_lowercase() == source_lower)
            else:
                print(f"[TAGGED-TABLE] ATTENTION: Aucune colonne source trouvée parmi: {all_columns}")

        # Sélectionner les colonnes demandées (APRÈS le filtre source)
        if columns:
            available_cols = [c for c in columns if c in all_columns]
            if available_cols:
                lf = lf.select(available_cols)

        # Appliquer les filtres
        for f in filters:
            col = f.get("column")
            op = f.get("operator", "contains")
            val = f.get("value", "") or f.get("text", "")  # Support "value" et "text" pour contains
            values = f.get("values", [])  # Pour le filtre "in"

            col_exists = col in (columns if columns else sheets_info[sheet_name].get("columns", []))
            if not col_exists:
                continue

            if op == "in" and values:
                # Filtre "in" : la valeur doit être dans la liste
                lf = lf.filter(pl.col(col).cast(pl.Utf8).is_in(values))
            elif col and val:
                if op == "contains":
                    lf = lf.filter(pl.col(col).cast(pl.Utf8).str.to_lowercase().str.contains(str(val).lower()))
                elif op == "equals":
                    lf = lf.filter(pl.col(col).cast(pl.Utf8) == str(val))
                elif op == "starts_with":
                    lf = lf.filter(pl.col(col).cast(pl.Utf8).str.starts_with(str(val)))
                elif op == "ends_with":
                    lf = lf.filter(pl.col(col).cast(pl.Utf8).str.ends_with(str(val)))

        # Compter le total de lignes filtrées AVANT d'appliquer la limite
        total_filtered_rows = lf.select(pl.count()).collect().item()

        # Limiter le nombre de lignes
        if limit and limit > 0:
            lf = lf.head(limit)

        # Collecter les résultats
        df = lf.collect()

        # Convertir en liste de dicts
        data = df.to_dicts()

        # Décoder les entités HTML dans les valeurs texte (&amp; → &)
        # Et nettoyer le HTML de la colonne Body pour les Emails
        is_emails = "email" in sheet_name.lower()
        for row in data:
            for key in row:
                if row[key] is None:
                    row[key] = ""
                elif isinstance(row[key], str):
                    val = html.unescape(row[key])
                    # Nettoyer HTML pour colonne Body des Emails
                    if is_emails and key == "Body" and val:
                        val = strip_html_tags(val)
                    row[key] = val

        return JSONResponse({
            "success": True,
            "columns": df.columns,
            "data": data,
            "total_rows": total_filtered_rows,  # Total filtré (avant limite)
            "returned_rows": len(data)  # Nombre de lignes retournées (après limite)
        })
    except Exception as e:
        print(f"[TAGGED-TABLE] Erreur: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({
            "success": False,
            "error": str(e)
        })


@app.post("/api/tagged-column-values", dependencies=[Depends(require_api_key)])
async def get_tagged_column_values(request: Request):
    """
    Retourne les valeurs uniques d'une colonne (max 10 000 valeurs).
    Body JSON: {sheet_name, column}
    """
    import polars as pl

    sheets_info = _tagged_data.get("sheets_info", {})

    if not sheets_info:
        return JSONResponse({
            "success": False,
            "error": "Aucun fichier Excel tagué disponible"
        })

    try:
        body = await request.json()
        sheet_name = body.get("sheet_name")
        column = body.get("column")

        if not sheet_name or not column:
            return JSONResponse({
                "success": False,
                "error": "Nom de feuille et colonne requis"
            })

        if sheet_name not in sheets_info:
            return JSONResponse({
                "success": False,
                "error": f"Feuille '{sheet_name}' non trouvée"
            })

        parquet_path = sheets_info[sheet_name].get("parquet_path")
        if not parquet_path or not Path(parquet_path).exists():
            return JSONResponse({
                "success": False,
                "error": f"Fichier Parquet non trouvé pour '{sheet_name}'"
            })

        # Lire le Parquet et récupérer les valeurs uniques
        lf = pl.scan_parquet(parquet_path)

        # Vérifier que la colonne existe
        available_cols = sheets_info[sheet_name].get("columns", [])
        if column not in available_cols:
            return JSONResponse({
                "success": False,
                "error": f"Colonne '{column}' non trouvée"
            })

        # Récupérer les valeurs uniques (max 10 000)
        unique_values = (
            lf.select(pl.col(column).cast(pl.Utf8))
            .collect()
            .get_column(column)
            .unique()
            .to_list()
        )

        # Filtrer les valeurs None/vides et limiter à 10 000
        unique_values = [v for v in unique_values if v is not None and str(v).strip() != ""]
        unique_values = sorted(unique_values, key=lambda x: str(x).lower())[:10000]

        # Décoder les entités HTML
        unique_values = [html.unescape(str(v)) for v in unique_values]

        return JSONResponse({
            "success": True,
            "column": column,
            "values": unique_values,
            "total_count": len(unique_values),
            "truncated": len(unique_values) >= 10000
        })

    except Exception as e:
        print(f"[TAGGED-COLUMN-VALUES] Erreur: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({
            "success": False,
            "error": str(e)
        })


@app.post("/api/excel-column-values", dependencies=[Depends(require_api_key)])
async def get_excel_column_values(request: Request):
    """
    Retourne les valeurs uniques d'une colonne pour les fichiers Excel réguliers (max 10 000 valeurs).
    Body JSON: {sheet_name, column, source_filter (optionnel)}
    """
    import polars as pl
    global _current_analyzer

    if _current_analyzer is None:
        return JSONResponse({
            "success": False,
            "error": "Aucun fichier Excel importé"
        })

    try:
        body = await request.json()
        sheet_name = body.get("sheet_name")
        column = body.get("column")
        source_filter = body.get("source_filter")

        if not sheet_name or not column:
            return JSONResponse({
                "success": False,
                "error": "Nom de feuille et colonne requis"
            })

        # Vérifier que la colonne existe
        available_cols = _current_analyzer.get_columns(sheet_name)
        if column not in available_cols:
            return JSONResponse({
                "success": False,
                "error": f"Colonne '{column}' non trouvée"
            })

        # Scanner le Parquet
        lf = _current_analyzer._scan(sheet_name)
        if lf is None:
            return JSONResponse({
                "success": False,
                "error": f"Feuille '{sheet_name}' non trouvée"
            })

        # Appliquer le filtre source si demandé
        if source_filter and "Source" in available_cols:
            lf = lf.filter(pl.col("Source") == source_filter)

        # Récupérer les valeurs uniques (max 10 000)
        unique_values = (
            lf.select(pl.col(column).cast(pl.Utf8))
            .collect()
            .get_column(column)
            .unique()
            .to_list()
        )

        # Filtrer les valeurs None/vides et limiter à 10 000
        unique_values = [v for v in unique_values if v is not None and str(v).strip() != ""]
        unique_values = sorted(unique_values, key=lambda x: str(x).lower())[:10000]

        # Décoder les entités HTML
        unique_values = [html.unescape(str(v)) for v in unique_values]

        return JSONResponse({
            "success": True,
            "column": column,
            "values": unique_values,
            "total_count": len(unique_values),
            "truncated": len(unique_values) >= 10000
        })

    except Exception as e:
        print(f"[EXCEL-COLUMN-VALUES] Erreur: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse({
            "success": False,
            "error": str(e)
        })


def parse_time_to_seconds(time_str: str) -> float:
    """Convertit un temps au format mm:ss ou hh:mm:ss en secondes"""
    if not time_str:
        return 0
    parts = time_str.strip().split(":")
    try:
        if len(parts) == 2:  # mm:ss
            return int(parts[0]) * 60 + float(parts[1])
        elif len(parts) == 3:  # hh:mm:ss
            return int(parts[0]) * 3600 + int(parts[1]) * 60 + float(parts[2])
        else:
            return float(time_str)
    except ValueError:
        return 0


def extract_video_thumbnails_ffmpeg(video_path: Path, num_thumbnails: int = 8) -> List[str]:
    """Fallback: Extrait des miniatures avec FFmpeg quand OpenCV échoue."""
    try:
        import subprocess
        import shutil

        # Vérifier si ffmpeg est disponible
        ffmpeg_path = shutil.which("ffmpeg")
        if not ffmpeg_path:
            print("[VIDEO FFMPEG] ffmpeg non trouvé")
            return []

        # Obtenir la durée de la vidéo avec ffprobe
        ffprobe_path = shutil.which("ffprobe")
        if ffprobe_path:
            result = subprocess.run(
                [ffprobe_path, "-v", "error", "-show_entries", "format=duration",
                 "-of", "default=noprint_wrappers=1:nokey=1", str(video_path)],
                capture_output=True, text=True, timeout=10
            )
            try:
                duration = float(result.stdout.strip())
            except:
                duration = 10.0  # Défaut 10 secondes
        else:
            duration = 10.0

        thumbnails = []
        video_name = video_path.stem

        for i in range(num_thumbnails):
            # Position dans la vidéo
            time_pos = (i + 0.5) * duration / num_thumbnails
            thumb_name = f"thumb_{video_name}_{i + 1}.jpg"
            thumb_path = UPLOAD_DIR / thumb_name

            # Extraire la frame avec ffmpeg
            cmd = [
                ffmpeg_path, "-y", "-ss", str(time_pos), "-i", str(video_path),
                "-vframes", "1", "-q:v", "2", "-vf", "scale=400:-1",
                str(thumb_path)
            ]
            result = subprocess.run(cmd, capture_output=True, timeout=15)

            if thumb_path.exists() and thumb_path.stat().st_size > 0:
                thumbnails.append(f"/uploads/{thumb_name}")
                print(f"[VIDEO FFMPEG] Miniature créée: {thumb_name}")

        return thumbnails
    except Exception as e:
        print(f"[VIDEO FFMPEG] Erreur: {e}")
        return []


def extract_video_thumbnails(video_path: Path, num_thumbnails: int = 8, start_time: str = None, end_time: str = None) -> List[str]:
    """
    Extrait des miniatures d'une vidéo à intervalles réguliers.
    start_time et end_time: format "mm:ss" ou "hh:mm:ss" pour définir l'intervalle
    Retourne une liste de chemins vers les images générées.
    Utilise FFmpeg comme fallback si OpenCV échoue.
    """
    try:
        import cv2
        import numpy as np

        cap = cv2.VideoCapture(str(video_path))
        if not cap.isOpened():
            print(f"[VIDEO] Impossible d'ouvrir avec OpenCV: {video_path}, essai avec FFmpeg...")
            cap.release()
            return extract_video_thumbnails_ffmpeg(video_path, num_thumbnails)

        # Obtenir les infos de la vidéo
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        if total_frames <= 0 or fps <= 0:
            print(f"[VIDEO] Pas de frames avec OpenCV: {video_path}, essai avec FFmpeg...")
            cap.release()
            return extract_video_thumbnails_ffmpeg(video_path, num_thumbnails)

        duration = total_frames / fps  # Durée totale en secondes
        print(f"[VIDEO] Durée: {duration:.1f}s, FPS: {fps}, Frames: {total_frames}")

        # Calculer les frames de début et fin selon l'intervalle
        start_seconds = parse_time_to_seconds(start_time) if start_time else 0
        end_seconds = parse_time_to_seconds(end_time) if end_time else duration

        # Valider les bornes
        start_seconds = max(0, min(start_seconds, duration))
        end_seconds = max(start_seconds, min(end_seconds, duration))

        start_frame = int(start_seconds * fps)
        end_frame = int(end_seconds * fps)
        interval_frames = end_frame - start_frame

        print(f"[VIDEO] Intervalle: {start_seconds:.1f}s - {end_seconds:.1f}s (frames {start_frame} - {end_frame})")

        if interval_frames <= 0:
            print(f"[VIDEO] Intervalle invalide")
            cap.release()
            return []

        # Calculer les positions des frames à extraire (réparties uniformément dans l'intervalle)
        frame_positions = []
        for i in range(num_thumbnails):
            pos = start_frame + int((i + 0.5) * interval_frames / num_thumbnails)
            frame_positions.append(min(pos, end_frame - 1))

        thumbnails = []
        video_name = video_path.stem

        for idx, frame_pos in enumerate(frame_positions):
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_pos)
            ret, frame = cap.read()

            if ret:
                # Générer un nom unique pour la miniature
                thumb_name = f"thumb_{video_name}_{idx + 1}.jpg"
                thumb_path = UPLOAD_DIR / thumb_name

                # Redimensionner pour miniatures haute qualité
                # Dimensions plus grandes pour conserver la qualité
                height, width = frame.shape[:2]
                max_width = 400
                max_height = 250

                # Calculer le ratio pour respecter les deux contraintes
                scale_w = max_width / width if width > max_width else 1
                scale_h = max_height / height if height > max_height else 1
                scale = min(scale_w, scale_h)  # Prendre le plus restrictif

                if scale < 1:
                    new_width = int(width * scale)
                    new_height = int(height * scale)
                    # INTER_LANCZOS4 pour une meilleure qualité de réduction
                    frame = cv2.resize(frame, (new_width, new_height), interpolation=cv2.INTER_LANCZOS4)

                # Sauvegarder avec qualité maximale (98% JPEG)
                cv2.imwrite(str(thumb_path), frame, [cv2.IMWRITE_JPEG_QUALITY, 98])
                thumbnails.append(f"/uploads/{thumb_name}")
                print(f"[VIDEO] Miniature créée: {thumb_name}")

        cap.release()
        return thumbnails

    except ImportError:
        print("[VIDEO] opencv-python non installé, impossible d'extraire les miniatures")
        return []
    except Exception as e:
        print(f"[VIDEO] Erreur extraction miniatures: {e}")
        return []


def format_seconds_to_time(seconds: float) -> str:
    """Convertit des secondes en format mm:ss ou hh:mm:ss"""
    if seconds < 0:
        return "00:00"
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


@app.post("/api/video-info", dependencies=[Depends(require_api_key)])
async def get_video_info(request: Request):
    """Récupère les informations d'une vidéo (durée, etc.)"""
    try:
        import cv2

        data = await request.json()
        video_url = data.get("video_url", "")

        if not video_url:
            raise HTTPException(status_code=400, detail="URL vidéo manquante")

        # Convertir l'URL en chemin local
        if video_url.startswith("/uploads/"):
            video_path = UPLOAD_DIR / video_url.replace("/uploads/", "")
        else:
            video_path = Path(video_url)

        if not video_path.exists():
            raise HTTPException(status_code=404, detail=f"Vidéo non trouvée: {video_path}")

        cap = cv2.VideoCapture(str(video_path))
        if not cap.isOpened():
            raise HTTPException(status_code=500, detail="Impossible d'ouvrir la vidéo")

        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        cap.release()

        if fps <= 0:
            raise HTTPException(status_code=500, detail="Impossible de lire les métadonnées vidéo")

        duration_seconds = total_frames / fps
        duration_formatted = format_seconds_to_time(duration_seconds)

        return JSONResponse({
            "success": True,
            "duration_seconds": duration_seconds,
            "duration_formatted": duration_formatted,
            "fps": fps,
            "total_frames": total_frames
        })

    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(status_code=500, detail="opencv-python non installé")
    except Exception as e:
        print(f"[VIDEO INFO] Erreur: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/video-thumbnails", dependencies=[Depends(require_api_key)])
async def generate_video_thumbnails(request: Request):
    """Génère les miniatures pour une vidéo spécifique avec intervalle de temps optionnel"""
    try:
        data = await request.json()
        video_url = data.get("video_url", "")
        start_time = data.get("start_time", None)  # Format: "mm:ss" ou "hh:mm:ss"
        end_time = data.get("end_time", None)  # Format: "mm:ss" ou "hh:mm:ss"
        num_thumbnails = data.get("num_thumbnails", 8)  # Nombre de miniatures (défaut: 8)

        if not video_url:
            raise HTTPException(status_code=400, detail="URL vidéo manquante")

        # Convertir l'URL en chemin local
        if video_url.startswith("/uploads/"):
            video_path = UPLOAD_DIR / video_url.replace("/uploads/", "")
        else:
            video_path = Path(video_url)

        if not video_path.exists():
            raise HTTPException(status_code=404, detail=f"Vidéo non trouvée: {video_path}")

        # Extraire les miniatures avec l'intervalle de temps
        thumbnails = extract_video_thumbnails(video_path, num_thumbnails=num_thumbnails, start_time=start_time, end_time=end_time)

        return JSONResponse({
            "success": True,
            "thumbnails": thumbnails,
            "count": len(thumbnails)
        })

    except HTTPException:
        raise
    except Exception as e:
        print(f"[VIDEO] Erreur: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════════════
# SYNCHRONISATION AUTOMATIQUE DE SESSION (survie aux refresh / cache wipe)
# ═══════════════════════════════════════════════════════════════════════
# SESSION_DIR et SESSION_FILE sont définis au début du fichier (lignes ~104-106)
_session_lock = threading.Lock()

# Les uploads non référencés doivent avoir au moins cet âge avant d'être déplacés,
# pour éviter de poubelliser un fichier tout juste uploadé avant son premier sync.
_CLEANUP_GRACE_SECONDS = 24 * 3600
_TRASH_RETENTION_SECONDS = 30 * 24 * 3600

_UPLOAD_URL_RE = re.compile(r"/?uploads/([^/\"'?#\s]+)")

def _extract_upload_filenames(obj: Any, out: set) -> None:
    """Parcours récursif du state pour collecter tous les fichiers /uploads/ référencés."""
    if isinstance(obj, str):
        for m in _UPLOAD_URL_RE.finditer(obj):
            out.add(m.group(1))
    elif isinstance(obj, dict):
        for v in obj.values():
            _extract_upload_filenames(v, out)
    elif isinstance(obj, list):
        for v in obj:
            _extract_upload_filenames(v, out)


class SyncSessionPayload(BaseModel):
    templates: Dict[str, Any] = Field(default_factory=dict)
    lastTemplate: Optional[str] = None


@app.post("/api/sync-session", dependencies=[Depends(require_api_key)])
async def sync_session(payload: SyncSessionPayload):
    """Sauvegarde la session complète côté serveur (appelé automatiquement par le client)."""
    try:
        data = {
            "version": 1,
            "saved_at": time.time(),
            "lastTemplate": payload.lastTemplate,
            "templates": payload.templates,
        }
        serialized = json.dumps(data, ensure_ascii=False)
        # Nom temporaire unique + verrou pour éviter les collisions entre requêtes concurrentes
        tmp = SESSION_DIR / f".sync_{uuid.uuid4().hex}.tmp"
        with _session_lock:
            SESSION_DIR.mkdir(parents=True, exist_ok=True)
            tmp.write_text(serialized, encoding="utf-8")
            try:
                tmp.replace(SESSION_FILE)
            except OSError:
                # Fallback: si replace échoue, on nettoie
                if tmp.exists():
                    tmp.unlink(missing_ok=True)
                raise
        return JSONResponse({"success": True, "saved_at": data["saved_at"]})
    except Exception as e:
        print(f"[SYNC-SESSION] Erreur: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/session-snapshot", dependencies=[Depends(require_api_key)])
async def session_snapshot():
    """Retourne le dernier snapshot serveur (pour restauration après wipe du navigateur)."""
    if not SESSION_FILE.exists():
        return JSONResponse({"success": True, "exists": False})
    try:
        data = json.loads(SESSION_FILE.read_text(encoding="utf-8"))
        return JSONResponse({"success": True, "exists": True, **data})
    except Exception as e:
        print(f"[SESSION-SNAPSHOT] Erreur: {e}")
        return JSONResponse({"success": False, "exists": False, "error": str(e)})


def _auto_cleanup_uploads() -> None:
    """Au démarrage: déplace dans uploads/_trash/ les fichiers non référencés (avec grâce 24h)
    et purge uploads/_trash/ au-delà de 30 jours."""
    try:
        now = time.time()
        trash_dir = UPLOAD_DIR / "_trash"

        # Purger la corbeille au-delà de la rétention
        if trash_dir.is_dir():
            for p in trash_dir.iterdir():
                try:
                    if p.is_file() and (now - p.stat().st_mtime) > _TRASH_RETENTION_SECONDS:
                        p.unlink()
                except OSError:
                    pass

        # Références actuelles depuis le snapshot
        referenced: set = set()
        if SESSION_FILE.exists():
            try:
                data = json.loads(SESSION_FILE.read_text(encoding="utf-8"))
                _extract_upload_filenames(data.get("templates", {}), referenced)
            except Exception as e:
                print(f"[STARTUP-CLEANUP] Snapshot illisible, skip: {e}")
                return
        else:
            # Pas de snapshot → on ne risque rien, on ne touche à rien
            print("[STARTUP-CLEANUP] Aucun snapshot, cleanup sauté")
            return

        moved = 0
        candidates = []
        for p in UPLOAD_DIR.iterdir():
            if not p.is_file():
                continue
            if p.name in referenced:
                continue
            # Grâce: on ne touche pas aux fichiers récents (potentiellement pas encore synced)
            try:
                age = now - p.stat().st_mtime
            except OSError:
                continue
            if age < _CLEANUP_GRACE_SECONDS:
                continue
            candidates.append(p)

        if not candidates:
            print(f"[STARTUP-CLEANUP] 0 orphelin (références: {len(referenced)})")
            return

        trash_dir.mkdir(exist_ok=True)
        ts = time.strftime("%Y%m%d_%H%M%S")
        for p in candidates:
            dest = trash_dir / f"{ts}__{p.name}"
            i = 1
            while dest.exists():
                dest = trash_dir / f"{ts}__{i}__{p.name}"
                i += 1
            try:
                p.rename(dest)
                moved += 1
            except OSError as e:
                print(f"[STARTUP-CLEANUP] Impossible de déplacer {p.name}: {e}")

        print(f"[STARTUP-CLEANUP] {moved} orphelin(s) déplacé(s) dans {trash_dir} (références: {len(referenced)})")
    except Exception as e:
        print(f"[STARTUP-CLEANUP] Erreur: {e}")


def _restore_referenced_from_trash() -> None:
    """
    Au demarrage: restaure dans uploads/ tous les fichiers references par
    le session.json (templates) qui auraient ete deplaces dans uploads/_trash/
    par une ancienne version du cleanup. Plus AUCUN deplacement automatique
    de uploads/ vers _trash/ — c'etait la source des pertes de donnees signalees
    par l'utilisateur lorsqu'il revenait apres plusieurs jours.

    Conserve la purge des fichiers du _trash plus vieux que _TRASH_RETENTION_SECONDS.
    """
    try:
        trash_dir = UPLOAD_DIR / "_trash"
        # 1) Purger les vieux fichiers du trash (>30j par defaut)
        if trash_dir.is_dir():
            now = time.time()
            purged = 0
            for p in trash_dir.iterdir():
                try:
                    if p.is_file() and (now - p.stat().st_mtime) > _TRASH_RETENTION_SECONDS:
                        p.unlink()
                        purged += 1
                except OSError:
                    pass
            if purged:
                print(f"[STARTUP-RESTORE] {purged} fichier(s) du trash purge(s) (>{_TRASH_RETENTION_SECONDS//86400}j)")

        if not trash_dir.is_dir():
            return
        if not SESSION_FILE.exists():
            print("[STARTUP-RESTORE] Pas de session.json, restauration skipped")
            return

        try:
            data = json.loads(SESSION_FILE.read_text(encoding="utf-8"))
        except Exception as e:
            print(f"[STARTUP-RESTORE] Session illisible: {e}")
            return

        referenced: set = set()
        _extract_upload_filenames(data.get("templates", {}), referenced)
        if not referenced:
            print("[STARTUP-RESTORE] Aucune reference dans session.json")
            return

        # Index des fichiers du trash par nom original (sans prefixe YYYYMMDD_HHMMSS__ ou ..__N__).
        trash_index: Dict[str, Path] = {}
        prefix_re = re.compile(r"^\d{8}_\d{6}__(?:\d+__)?(.+)$")
        for p in trash_dir.iterdir():
            if not p.is_file():
                continue
            m = prefix_re.match(p.name)
            original = m.group(1) if m else p.name
            existing = trash_index.get(original)
            if existing is None or p.stat().st_mtime > existing.stat().st_mtime:
                trash_index[original] = p

        restored = 0
        for fname in referenced:
            target = UPLOAD_DIR / fname
            if target.exists():
                continue
            src = trash_index.get(fname)
            if src is None:
                continue
            try:
                src.rename(target)
                restored += 1
            except OSError as e:
                print(f"[STARTUP-RESTORE] Echec restauration {fname}: {e}")

        if restored:
            print(f"[STARTUP-RESTORE] {restored} fichier(s) restaure(s) depuis _trash (sur {len(referenced)} references)")
        else:
            print(f"[STARTUP-RESTORE] Aucun fichier a restaurer (references: {len(referenced)})")
    except Exception as e:
        print(f"[STARTUP-RESTORE] Erreur: {e}")


@app.on_event("startup")
async def _on_startup_cleanup():
    # IMPORTANT: l'ancien _auto_cleanup_uploads() deplacait dans _trash/ tous les
    # fichiers non references depuis 24h, ce qui provoquait des pertes apparentes
    # de donnees quand le user revenait apres plusieurs jours et que session.json
    # ne referencait pas (encore) ses fichiers (sync incomplet, template cache, etc).
    # On remplace par une RESTAURATION: si des refs pointent vers _trash, on les
    # ramene dans uploads/. Le user peut nettoyer manuellement quand il le decide.
    _restore_referenced_from_trash()


if FRONTEND_DIR.exists():
    app.mount("/ui", StaticFiles(directory=str(FRONTEND_DIR), html=True), name="ui")
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR), html=False), name="uploads")



